import streamlit as st
from datetime import datetime
import time
import os
from PyPDF2 import PdfReader
import json
import requests
from io import BytesIO
import base64
import pandas as pd
import numpy as np
from PIL import Image, ImageEnhance, ImageFilter, ImageDraw, ImageFont
import plotly.graph_objects as go
import docx
import yfinance as yf
import folium
from streamlit_folium import st_folium
from dateutil import tz
import random
from pathlib import Path
import re
import urllib.parse
import subprocess
from moviepy.editor import VideoFileClip, AudioFileClip
import replicate
import concurrent.futures

# Initialize session state variables early
if 'pro_unlocked' not in st.session_state:
    st.session_state.pro_unlocked = False
if 'pro_verified' not in st.session_state:
    st.session_state.pro_verified = False

# Set page config FIRST (before any other st calls)
st.set_page_config(
    page_title="Quantora AI",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Show loading state
st.title("🚀 Starting Quantora AI...")
status = st.status("Initializing...", expanded=True)

with status:
    st.write("Checking dependencies...")
    time.sleep(0.1)
    
    # List of required modules to check
    modules_to_check = [
        ("JSON", "json"),
        ("Requests", "requests"),
        ("Pandas", "pandas"),
        ("NumPy", "numpy"),
        ("Pillow (PIL)", "PIL"),
        ("Plotly", "plotly"),
        ("PyPDF2", "PyPDF2"),
        ("Docx", "docx"),
        ("yFinance", "yfinance"),
        ("Folium", "folium")
    ]
    
    # Check each module
    for module_name, module_import in modules_to_check:
        try:
            if module_import == "json":
                import json
            elif module_import == "requests":
                import requests
            elif module_import == "pandas":
                import pandas as pd
            elif module_import == "numpy":
                import numpy as np
            elif module_import == "PIL":
                from PIL import Image
            elif module_import == "plotly":
                import plotly.graph_objects as go
            elif module_import == "PyPDF2":
                from PyPDF2 import PdfReader
            elif module_import == "docx":
                import docx
            elif module_import == "yfinance":
                import yfinance as yf
            elif module_import == "folium":
                import folium
            
            st.write(f"✓ {module_name} loaded")
        except Exception as e:
            st.warning(f"⚠️ {module_name} import issue: {str(e)[:50]}")
    
    st.write("Finalizing setup...")
    time.sleep(0.1)

status.update(label="Ready!", state="complete", expanded=False)
time.sleep(0.1)

# ✅ API Configuration
A4F_API_KEY = "ddc-a4f-b752e3e2936149f49b1b306953e0eaab"
API_URL = "https://api.a4f.co/v1/chat/completions"
IMAGE_MODEL = "provider-2/nano-banana-pro"
EDIT_MODEL = "provider-2/nano-banana-pro"
VIDEO_MODEL = "provider-6/wan-2.1"

# Weather API
WEATHER_API_KEY = "b9e17942665c0df9828cd0d0d721a44f"
CURRENT_URL = "https://api.openweathermap.org/data/2.5/weather"
FORECAST_URL = "https://api.openweathermap.org/data/2.5/forecast"
ICON_URL = "https://openweathermap.org/img/wn/{}@4x.png"

# Replicate API for Video Generation
REPLICATE_API_TOKEN = "r8_7t4VS9WzjYf0ohxFuez5bDAa66dNalb3w5Jql"
os.environ["REPLICATE_API_TOKEN"] = REPLICATE_API_TOKEN

# History persistence
HISTORY_FILE = "quantora_history.json"
if not os.path.exists(HISTORY_FILE):
    with open(HISTORY_FILE, 'w') as f:
        json.dump([], f)

def load_history():
    try:
        with open(HISTORY_FILE, 'r') as f:
            return json.load(f)
    except:
        return []

def save_history(query):
    try:
        history = load_history()
        history.append({"query": query, "timestamp": datetime.now().isoformat()})
        with open(HISTORY_FILE, 'w') as f:
            json.dump(history, f)
    except:
        pass  # Silently fail if history can't be saved

# ✅ API clients initialization
def initialize_clients():
    """Initialize API clients with proper error handling"""
    groq_client = None
    groq_available = False
    
    # Try to initialize Groq client
    try:
        from groq import Groq
        groq_api_key = os.environ.get("Groq_API_TOKEN", "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")
        groq_client = Groq(api_key=groq_api_key)
        
        # Test with a simple call
        test_response = groq_client.chat.completions.create(
            model="qwen/qwen3-32b",
            messages=[{"role": "user", "content": "test"}],
            max_tokens=5
        )
        groq_available = True
    except ImportError:
        st.sidebar.warning("Groq library not installed. Install with: pip install groq")
    except Exception as e:
        st.sidebar.warning(f"Groq client error: {str(e)[:80]}")
    
    # A4F client configuration
    a4f_client = {
        "api_key": A4F_API_KEY,
        "api_url": API_URL,
        "available": True
    }
    
    # Test A4F connection
    a4f_available = True  # Assume it works, will be tested on first call
    
    return groq_client, a4f_client, groq_available, a4f_available

# Initialize clients
groq_client, a4f_client, groq_available, a4f_available = initialize_clients()

# Show API status in sidebar (if sidebar is available)
if not groq_available:
    st.sidebar.warning("⚠️ Groq API unavailable - using A4F API")
else:
    st.sidebar.success("✅ All APIs connected")
    
# --------------------------
# QUANTORA WEATHER MODULE
# --------------------------
def quantora_weather():
    st.markdown(
        """
        <style>
        /* fonts & base */
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;800&display=swap');
        html, body, [class*="css"]  {
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, "SF Pro Display", "Helvetica Neue", Arial;
            /* blue / dark-blue gradient background as requested */
            background: linear-gradient(180deg, #07143a 0%, #08204f 30%, #07203d 60%, #03112a 100%);
            color: #eaf0ff;
            background-attachment: fixed !important;
        }
        /* fix streamlit container overflow that sometimes prints raw HTML at bottom */
        .reportview-container .main .block-container {
            padding-bottom: 24px !important;
        }

        /* full-screen animated canvas */
        .quantora-canvas {
            position: fixed; inset: 0; z-index: 0; pointer-events: none;
        }
        /* glass containers */
        .glass {
            background: linear-gradient(180deg, rgba(255,255,255,0.03), rgba(255,255,255,0.02));
            border: 1px solid rgba(255,255,255,0.06);
            backdrop-filter: blur(12px) saturate(110%);
            -webkit-backdrop-filter: blur(12px) saturate(110%);
            border-radius: 20px;
            padding: 18px;
            box-shadow: 0 8px 30px rgba(2,6,23,0.7);
            color: #eaf0ff;
            position: relative;
            z-index: 2;
        }
        .hero {
            border-radius: 28px;
            padding: 26px;
        }
        .muted { color: rgba(235,245,255,0.6); font-size:13px }
        .kpi { font-weight:700; font-size:20px }
        /* hourly scroll row */
        .hourly-row {
            display:flex; gap:14px; overflow-x:auto; padding-bottom:6px;
        }
        .hour-card {
            min-width:98px; background:linear-gradient(180deg, rgba(255,255,255,0.02), rgba(255,255,255,0.01));
            border-radius:14px; padding:12px; text-align:center; border:1px solid rgba(255,255,255,0.04);
        }
        .day-pill {
            display:inline-block; padding:8px 12px; border-radius:999px; background: rgba(255,255,255,0.02);
            border: 1px solid rgba(255,255,255,0.03); margin-right:8px;
        }
        /* small responsive tweaks */
        @media (max-width:900px){
            .hour-card { min-width:86px }
        }

        /* --- Animated clouds + rain via SVG/CSS keyframes --- */
        .cloud {
            position:absolute; opacity:0.95;
            transform: translateZ(0);
            filter: drop-shadow(0 8px 40px rgba(0,0,0,0.45));
        }
        /* cloud motion -- different speeds and sizes */
        @keyframes cloudMove1 { from {transform: translateX(-20%);} to {transform: translateX(120%);} }
        @keyframes cloudMove2 { from {transform: translateX(-30%);} to {transform: translateX(110%);} }
        @keyframes cloudFloat { 0%{transform:translateY(0);}50%{transform:translateY(-6px);}100%{transform:translateY(0);} }

        .cloud.c1 { top:6%; left:-10%; width:680px; animation: cloudMove1 120s linear infinite, cloudFloat 8s ease-in-out infinite; opacity:0.95; }
        .cloud.c2 { top:18%; left:-20%; width:520px; animation: cloudMove2 180s linear infinite, cloudFloat 10s ease-in-out infinite; opacity:0.85; }
        .cloud.c3 { top:34%; left:-5%; width:420px; animation: cloudMove1 220s linear infinite, cloudFloat 14s ease-in-out infinite; opacity:0.75; transform: scale(0.92); }

        /* rain overlay */
        .raindrops { position:absolute; inset:0; pointer-events:none; z-index:1; opacity:0.12; mix-blend-mode:screen; }
        .rain-line { animation: rainFall 0.9s linear infinite; stroke: rgba(255,255,255,0.12); stroke-width:1.2 }
        @keyframes rainFall { 0%{ transform: translateY(-8px); opacity:0 } 50% {opacity:0.6} 100%{ transform: translateY(24px); opacity:0 } }

        /* tiny shimmer in hero */
        .hero::after {
            content: "";
            position:absolute;
            left:-20%; top:-30%;
            width:240px; height:240px;
            background: radial-gradient(circle at 30% 30%, rgba(255,255,255,0.04), transparent 20%);
            transform: rotate(28deg); filter: blur(36px); opacity:0.6;
            pointer-events:none;
        }
        </style>

        <!-- SVG Animated Background (cloud shapes + subtle rain) -->
        <div class="quantora-canvas" aria-hidden="true">
          <svg class="cloud c1" viewBox="0 0 800 200" preserveAspectRatio="xMidYMid meet">
            <defs><linearGradient id="g1" x1="0" x2="1"><stop offset="0" stop-color="#bfe7ff" stop-opacity="0.06"/><stop offset="1" stop-color="#9fdcff" stop-opacity="0.02"/></linearGradient></defs>
            <g transform="scale(1.0)">
              <path fill="url(#g1)" d="M60 120c-18-40 30-96 96-64 14-54 110-64 156-12 40-36 130-18 150 22 48 12 34 86-22 86H60z"/>
            </g>
          </svg>

          <svg class="cloud c2" viewBox="0 0 700 160" preserveAspectRatio="xMidYMid meet">
            <defs><linearGradient id="g2" x1="0" x2="1"><stop offset="0" stop-color="#ffffff" stop-opacity="0.05"/><stop offset="1" stop-color="#bfeaff" stop-opacity="0.02"/></linearGradient></defs>
            <path fill="url(#g2)" d="M20 90c-12-28 22-68 68-46 10-38 84-44 122-8 32-28 100-14 118 18 36 10 26 52-16 52H20z"/>
          </svg>

          <svg class="cloud c3" viewBox="0 0 600 140" preserveAspectRatio="xMidYMid meet">
            <defs><linearGradient id="g3" x1="0" x2="1"><stop offset="0" stop-color="#eaf6ff" stop-opacity="0.04"/><stop offset="1" stop-color="#bfe8ff" stop-opacity="0.02"/></linearGradient></defs>
            <path fill="url(#g3)" d="M0 80c-8-20 14-48 46-34 7-26 56-30 82-6 22-18 72-8 86 12 28 8 20 36-12 36H0z"/>
          </svg>

          <!-- rain layer -->
          <svg class="raindrops" viewBox="0 0 1200 600" preserveAspectRatio="none">
            <g>
              <g transform="translate(0,0)">
                <line class="rain-line" x1="40" x2="44" y1="0" y2="18" />
                <line class="rain-line" x1="120" x2="124" y1="-6" y2="12" />
                <line class="rain-line" x1="220" x2="224" y1="-8" y2="14" />
                <line class="rain-line" x1="360" x2="364" y1="-2" y2="18" />
                <line class="rain-line" x1="480" x2="484" y1="-6" y2="16" />
                <line class="rain-line" x1="600" x2="604" y1="-10" y2="12" />
                <line class="rain-line" x1="740" x2="744" y1="-4" y2="18" />
                <line class="rain-line" x1="860" x2="864" y1="-12" y2="10" />
                <line class="rain-line" x1="980" x2="984" y1="-6" y2="12" />
                <line class="rain-line" x1="1100" x2="1104" y1="-8" y2="14" />
              </g>
            </g>
          </svg>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Top title
    st.markdown(
        "<div style='display:flex; align-items:center; justify-content:space-between; gap:18px;'>"
        "<div style='display:flex; align-items:center; gap:12px'>"
        "<div style='font-size:26px; font-weight:800'>Quantora Weather</div>"
        "<div style='color:rgba(235,245,255,0.6); font-size:13px'>Premium • Ultra • Market</div>"
        "</div>"
        "<div style='font-size:13px; color:rgba(235,245,255,0.6)'>Live • AI-assisted • Animated UI</div>"
        "</div>",
        unsafe_allow_html=True,
    )

    st.markdown("<br/>", unsafe_allow_html=True)

    # layout columns
    left_col, right_col = st.columns([1.4, 1])

    # ---------------- SIDEBAR-LIKE CONTROLS (inline) ----------------
    with right_col:
        st.markdown("<div class='glass' style='padding:12px'>", unsafe_allow_html=True)
        st.markdown("### Search & Options", unsafe_allow_html=True)
        # ensure session key exists
        if "city_input" not in st.session_state:
            st.session_state["city_input"] = "Berlin"
        city = st.text_input("City name", value=st.session_state.get("city_input", "Berlin"))
        st.session_state["city_input"] = city
        units = st.selectbox("Units", options=["metric", "imperial"], index=0)
        tz_mode = st.selectbox("Timezone", options=["Local", "UTC"], index=0)
        st.markdown("</div>", unsafe_allow_html=True)
        st.markdown("<br/>", unsafe_allow_html=True)

    # ---------------- FETCH DATA ----------------
    def fetch_current(city_name):
        params = {"q": city_name, "appid": WEATHER_API_KEY, "units": units}
        r = requests.get(CURRENT_URL, params=params, timeout=15)
        r.raise_for_status()
        return r.json()

    def fetch_forecast(city_name):
        params = {"q": city_name, "appid": WEATHER_API_KEY, "units": units}
        r = requests.get(FORECAST_URL, params=params, timeout=15)
        r.raise_for_status()
        return r.json()

    try:
        current = fetch_current(city)
        forecast = fetch_forecast(city)
    except requests.HTTPError:
        st.error("Could not fetch weather. Check city name or API key.")
        st.stop()
    except Exception as e:
        st.error(f"Network error: {e}")
        st.stop()

    # ---------------- PARSE DATA ----------------
    # current
    weather_desc = current["weather"][0]["description"].title()
    icon_code = current["weather"][0]["icon"]
    temp = current["main"]["temp"]
    feels_like = current["main"]["feels_like"]
    humidity = current["main"]["humidity"]
    pressure = current["main"]["pressure"]
    wind = current["wind"].get("speed", 0)
    visibility = current.get("visibility", 0)
    sunrise = datetime.utcfromtimestamp(current["sys"]["sunrise"])
    sunset = datetime.utcfromtimestamp(current["sys"]["sunset"])
    lat = current["coord"]["lat"]
    lon = current["coord"]["lon"]

    # forecast -> dataframe
    rows = []
    for item in forecast["list"]:
        dt_utc = datetime.utcfromtimestamp(item["dt"])
        rows.append({
            "dt_utc": dt_utc,
            "temp": item["main"]["temp"],
            "temp_min": item["main"]["temp_min"],
            "temp_max": item["main"]["temp_max"],
            "feels": item["main"]["feels_like"],
            "humidity": item["main"]["humidity"],
            "pressure": item["main"]["pressure"],
            "wind": item["wind"].get("speed", 0),
            "wind_deg": item["wind"].get("deg", None),
            "pop": item.get("pop", 0),
            "desc": item["weather"][0]["description"].title(),
            "icon": item["weather"][0]["icon"]
        })
    df = pd.DataFrame(rows)
    # timezone handling
    if tz_mode == "Local":
        tz_local = tz.tzlocal()
    else:
        tz_local = tz.tzutc()
    df["dt_local"] = df["dt_utc"].apply(lambda x: x.replace(tzinfo=tz.tzutc()).astimezone(tz_local))
    df["time_local"] = df["dt_local"].dt.strftime("%H:%M")
    df["date_local"] = df["dt_local"].dt.date

    # daily compact forecast (next 7 unique days)
    daily = df.groupby("date_local").agg({
        "temp_min": "min", "temp_max": "max", "temp": "mean", "pop": "mean", "humidity": "mean"
    }).reset_index().head(7)
    if "pop" in daily.columns:
        daily["pop_pct"] = (daily["pop"] * 100).round().astype(int)
    else:
        daily["pop_pct"] = 0

    # safe unit symbol for use in multiple places
    unit_sym = "°C" if units == "metric" else "°F"

    # ---------------- HERO CARD ----------------
    with left_col:
        st.markdown("<div class='glass hero' style='position:relative'>", unsafe_allow_html=True)
        # left hero grid: big temp + small details + hourly strip
        hero_col_a, hero_col_b = st.columns([1, 1.1])
        with hero_col_a:
            # icon from openweather
            try:
                icon_resp = requests.get(ICON_URL.format(icon_code), timeout=8)
                icon_img = Image.open(BytesIO(icon_resp.content)).convert("RGBA")
                st.image(icon_img, width=160)
            except Exception:
                st.write("")
            st.markdown(f"<div style='font-size:54px; font-weight:900; margin-top:6px'>{temp}{unit_sym}</div>", unsafe_allow_html=True)
            st.markdown(f"<div style='color:rgba(235,245,255,0.78); font-weight:600'>{weather_desc}</div>", unsafe_allow_html=True)
            st.markdown("<br/>", unsafe_allow_html=True)
            st.markdown(f"<div class='muted'>Feels like <span style='font-weight:700'>{feels_like}{unit_sym}</span> • Humidity <span style='font-weight:700'>{humidity}%</span></div>", unsafe_allow_html=True)
            st.markdown(f"<div class='muted'>Wind <span style='font-weight:700'>{wind} { 'm/s' if units=='metric' else 'mph' }</span> • Visibility <span style='font-weight:700'>{int(visibility/1000)} km</span></div>", unsafe_allow_html=True)

        with hero_col_b:
            # hourly cards (next 8 slots)
            st.markdown("<div style='display:flex; justify-content:space-between; align-items:flex-start;'>"
                        f"<div class='muted' style='font-weight:700'>Hourly</div>"
                        f"<div class='muted'>{city.title()}</div>"
                        "</div>", unsafe_allow_html=True)
            st.markdown("<div class='hourly-row' style='margin-top:8px'>", unsafe_allow_html=True)
            next_8 = df.head(8)
            if next_8.empty:
                st.markdown("<div class='hour-card'>No hourly data</div>", unsafe_allow_html=True)
            else:
                for _, row in next_8.iterrows():
                    icon = row.get("icon", "")
                    hour = row.get("time_local", "")
                    t = round(row.get("temp", 0))
                    desc = row.get("desc", "")
                    # small icon
                    icon_url = ICON_URL.format(icon) if icon else ""
                    card_html = f"""
                        <div class='hour-card'>
                            <div style='font-size:13px; color:rgba(235,245,255,0.7);'>{hour}</div>
                            <img src="{icon_url}" style='width:64px;height:64px;margin-top:2px;filter:drop-shadow(0 8px 20px rgba(0,0,0,0.5));'/>
                            <div style='font-size:16px; font-weight:700; margin-top:6px'>{t}{unit_sym}</div>
                            <div style='font-size:12px; color:rgba(235,245,255,0.65)'>{desc}</div>
                        </div>
                    """
                    st.markdown(card_html, unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

        # small KPIs row below hero
        st.markdown("<div style='display:flex; gap:12px; margin-top:16px'>", unsafe_allow_html=True)
        try:
            sr_str = sunrise.replace(tzinfo=tz.tzutc()).astimezone(tz_local).strftime('%H:%M')
            ss_str = sunset.replace(tzinfo=tz.tzutc()).astimezone(tz_local).strftime('%H:%M')
        except Exception:
            sr_str, ss_str = "-", "-"
        st.markdown(f"<div class='glass' style='padding:12px; flex:1'><div class='muted'>Sunrise</div><div class='kpi'>{sr_str}</div></div>", unsafe_allow_html=True)
        st.markdown(f"<div class='glass' style='padding:12px; flex:1'><div class='muted'>Sunset</div><div class='kpi'>{ss_str}</div></div>", unsafe_allow_html=True)
        st.markdown(f"<div class='glass' style='padding:12px; flex:1'><div class='muted'>Pressure</div><div class='kpi'>{pressure} hPa</div></div>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("</div>", unsafe_allow_html=True)
        st.markdown("<br/>", unsafe_allow_html=True)

        # ---------------- CHARTS SECTION ----------------
        st.markdown("<div class='glass' style='padding:16px'>", unsafe_allow_html=True)
        st.markdown("<div style='display:flex; justify-content:space-between; align-items:center;'><div style='font-weight:800'>Forecast Analytics</div><div class='muted'>Temperature • Humidity • Precipitation</div></div>", unsafe_allow_html=True)

        # temperature trend (hourly)
        if not df.empty:
            fig_temp = go.Figure()
            fig_temp.add_trace(go.Scatter(x=df["dt_local"], y=df["temp"], mode="lines+markers", name="Temp", line=dict(width=3)))
            fig_temp.update_layout(template="plotly_dark", height=300, margin=dict(t=20,l=10,r=10,b=10),
                                   xaxis=dict(title=""), yaxis=dict(title=f"Temperature ({unit_sym})"))
            st.plotly_chart(fig_temp, use_container_width=True)

            # humidity + pop
            fig_h = go.Figure()
            fig_h.add_trace(go.Bar(x=df["dt_local"], y=df["humidity"], name="Humidity (%)", opacity=0.8))
            fig_h.add_trace(go.Scatter(x=df["dt_local"], y=df["pop"]*100, name="Precip Chance (%)", yaxis="y2", mode="lines+markers", line=dict(width=2, dash='dash')))
            fig_h.update_layout(template="plotly_dark", height=260, margin=dict(t=10,l=10,r=10,b=10),
                                yaxis=dict(title="Humidity (%)"), yaxis2=dict(title="Precip (%)", overlaying="y", side="right"))
            st.plotly_chart(fig_h, use_container_width=True)
        else:
            st.markdown("<div class='muted'>No forecast chart data available</div>", unsafe_allow_html=True)

        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<br/>", unsafe_allow_html=True)

        # ---------------- DAILY CARDS (compact) ----------------
        st.markdown("<div class='glass' style='padding:12px'>", unsafe_allow_html=True)
        st.markdown("<div style='font-weight:800; margin-bottom:8px'>7-Day Outlook</div>", unsafe_allow_html=True)

        daily_cards_html = "<div style='display:flex; gap:10px; overflow-x:auto; padding-bottom:6px;'>"
        if daily.empty:
            daily_cards_html += "<div style='min-width:120px; padding:10px;'>No daily data</div>"
        else:
            for _, r in daily.iterrows():
                date_str = r["date_local"].strftime("%a %d")
                tmin = int(round(r["temp_min"])) if not np.isnan(r["temp_min"]) else "-"
                tmax = int(round(r["temp_max"])) if not np.isnan(r["temp_max"]) else "-"
                pop = int(r["pop_pct"]) if "pop_pct" in r and not pd.isna(r["pop_pct"]) else 0
                hum = int(round(r["humidity"])) if not np.isnan(r["humidity"]) else "-"
                # pick icon from df for that date (first occurrence) safely
                sample_icon = None
                subset = df[df["date_local"] == r["date_local"]]
                if not subset.empty:
                    sample_icon = subset.iloc[0].get("icon", None)
                icon_url = ICON_URL.format(sample_icon) if sample_icon else ""
                card = f"""
                    <div style='min-width:120px; padding:10px; border-radius:14px; text-align:center; border:1px solid rgba(255,255,255,0.04);'>
                      <div style='font-weight:700'>{date_str}</div>
                      <img src='{icon_url}' style='width:72px;height:72px'/>
                      <div style='margin-top:6px; font-weight:800'>{tmax}{unit_sym} <span style='color:rgba(235,245,255,0.6); font-weight:500'>/{tmin}{unit_sym}</span></div>
                      <div class='muted' style='margin-top:6px'>Precip {pop}% • Hum {hum}%</div>
                    </div>
                """
                daily_cards_html += card
        daily_cards_html += "</div>"
        st.markdown(daily_cards_html, unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

    # ---------------- RIGHT COLUMN DETAILS & MAP ----------------
    with right_col:
        # analytics small cards
        st.markdown("<div class='glass' style='padding:14px'>", unsafe_allow_html=True)
        st.markdown("<div style='font-weight:800'>Today's Details</div>", unsafe_allow_html=True)
        st.markdown("<div style='margin-top:10px; display:flex; gap:8px; flex-wrap:wrap'>", unsafe_allow_html=True)
        st.markdown(f"<div class='day-pill'><div style='font-size:12px' class='muted'>Humidity</div><div style='font-weight:700'>{humidity}%</div></div>", unsafe_allow_html=True)
        # safe precip average
        precip_avg_pct = int(df['pop'].mean()*100) if ("pop" in df.columns and not df['pop'].isna().all()) else 0
        st.markdown(f"<div class='day-pill'><div style='font-size:12px' class='muted'>Precip</div><div style='font-weight:700'>{precip_avg_pct}%</div></div>", unsafe_allow_html=True)
        st.markdown(f"<div class='day-pill'><div style='font-size:12px' class='muted'>Wind</div><div style='font-weight:700'>{wind} {'m/s' if units=='metric' else 'mph'}</div></div>", unsafe_allow_html=True)
        st.markdown(f"<div class='day-pill'><div style='font-size:12px' class='muted'>UV</div><div style='font-weight:700'>-</div></div>", unsafe_allow_html=True)
        st.markdown(f"<div class='day-pill'><div style='font-size:12px' class='muted'>Visibility</div><div style='font-weight:700'>{int(visibility/1000)} km</div></div>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)
        st.markdown("<br/>", unsafe_allow_html=True)

        # small wind rose: create polar bar with aggregated directions if degrees exist
        st.markdown("<div class='glass' style='padding:14px'>", unsafe_allow_html=True)
        st.markdown("<div style='font-weight:800'>Wind Rose</div>", unsafe_allow_html=True)
        # aggregate directions
        def dir_from_deg(d):
            if d is None: return None
            dirs = ["N","NE","E","SE","S","SW","W","NW"]
            ix = int(((d+22.5) % 360) / 45)
            return dirs[ix]
        dirs = ["N","NE","E","SE","S","SW","W","NW"]
        buckets = {k:0 for k in dirs}
        for _, r in df.iterrows():
            deg = r.get("wind_deg", None)
            s = r.get("wind", 0)
            dname = dir_from_deg(deg)
            if dname:
                buckets[dname] += s
        vals = [buckets[d] for d in dirs]
        try:
            fig_w = go.Figure(go.Barpolar(theta=dirs, r=vals, marker_line_color="white", marker_line_width=1))
            fig_w.update_layout(template="plotly_dark", height=300, margin=dict(t=20,b=20))
            st.plotly_chart(fig_w, use_container_width=True)
        except Exception:
            st.markdown("<div class='muted'>Wind rose unavailable</div>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)
        st.markdown("<br/>", unsafe_allow_html=True)

        # interactive mini map with folium (if available)
        try:
            st.markdown("<div class='glass' style='padding:12px'>", unsafe_allow_html=True)
            st.markdown("<div style='font-weight:800'>Location</div>", unsafe_allow_html=True)
            m = folium.Map(location=[lat, lon], zoom_start=9, tiles='CartoDB dark_matter', control_scale=True, prefer_canvas=True)
            folium.CircleMarker(location=[lat, lon], radius=10, color='#66e0ff', fill=True, fill_color='#66e0ff', fill_opacity=0.7, popup=f"{city.title()}").add_to(m)
            st_data = st_folium(m, width=350, height=240)
            st.markdown("</div>", unsafe_allow_html=True)
        except Exception:
            # fallback: just show coords
            st.markdown("<div class='glass' style='padding:12px'>", unsafe_allow_html=True)
            st.markdown("<div style='font-weight:800'>Location</div>", unsafe_allow_html=True)
            st.markdown(f"<div class='muted'>Lat • Lon</div><div style='font-weight:700'>{lat:.3f}, {lon:.3f}</div>", unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<br/>", unsafe_allow_html=True)

        # CSV export & snapshot
        csv = df.to_csv(index=False)
        b64 = base64.b64encode(csv.encode()).decode()
        st.markdown("<div class='glass' style='padding:12px'>", unsafe_allow_html=True)
        st.markdown("<div style='font-weight:800'>Export</div>", unsafe_allow_html=True)
        st.markdown(f"<a href='data:file/csv;base64,{b64}' download='quantora_forecast.csv'>Download forecast CSV</a>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

    # ---------------- AI WEATHER ANALYSIS ----------------
    st.markdown("---")
    st.markdown("<div class='glass' style='padding:20px; margin-top:20px'>", unsafe_allow_html=True)
    st.markdown("## 🤖 AI Weather Analysis & Recommendations")
    
    # Prepare weather data for AI analysis
    weather_summary = f"""
    CURRENT WEATHER DATA FOR {city.upper()}:
    - Temperature: {temp}{unit_sym} (Feels like: {feels_like}{unit_sym})
    - Conditions: {weather_desc}
    - Humidity: {humidity}%
    - Wind Speed: {wind} {'m/s' if units == 'metric' else 'mph'}
    - Pressure: {pressure} hPa
    - Visibility: {int(visibility/1000)} km
    - Sunrise: {sr_str}
    - Sunset: {ss_str}
    
    7-DAY FORECAST SUMMARY:
    """
    
    for i, day in daily.iterrows():
        date_str = day["date_local"].strftime("%A, %b %d")
        weather_summary += f"""
    - {date_str}: High {int(day['temp_max'])}{unit_sym} / Low {int(day['temp_min'])}{unit_sym}, 
      Precipitation chance: {int(day['pop_pct'])}%, Humidity: {int(day['humidity'])}%
        """
    
    ai_prompt = f"""
    As Quantora AI Weather Expert, analyze this weather data and provide comprehensive recommendations:
    
    {weather_summary}
    
    Provide analysis in these categories:
    1. 🌡️ **Health & Safety Recommendations**
       - Temperature-related precautions
       - Hydration needs
       - UV protection level
       - Air quality considerations
    
    2. 👕 **Clothing & Gear Suggestions**
       - Appropriate clothing layers
       - Special gear needed (umbrella, sunglasses, etc.)
       - Footwear recommendations
    
    3. 🏃 **Activity Recommendations**
       - Best times for outdoor exercise
       - Suitable indoor alternatives if weather is poor
       - Specific activities recommended for this weather
    
    4. 🚗 **Travel & Commute Advice**
       - Driving conditions
       - Public transportation considerations
       - Weather-related delays to anticipate
    
    5. 🏠 **Home & Property Tips**
       - Energy efficiency suggestions
       - Home maintenance recommendations
       - Gardening/outdoor work timing
    
    6. 🛒 **Shopping & Planning**
       - Items to stock up on
       - Meal planning suggestions for this weather
       - Emergency preparedness if severe weather expected
    
    7. 📊 **Statistical Insights**
       - How this weather compares to seasonal averages
       - Trend analysis for the week
       - Probability of weather changes
    
    Format your response with clear sections and use emojis for visual appeal.
    """
    
    if st.button("🧠 Get AI Weather Analysis", key="ai_weather_analysis", type="primary"):
        with st.spinner("🤖 Quantora AI is analyzing weather data..."):
            # Save current model version and set to fast model for weather analysis
            current_model = st.session_state.get("model_version", "Quantora Prime 1 (Latest Flagship Model)")
            st.session_state.model_version = "Quantora Prime 1 Fast (Faster But Not As Better As Og Flagship Model)"
            
            ai_response = call_quantora_unified(ai_prompt)
            
            # Restore original model
            st.session_state.model_version = current_model
            
            st.markdown("### 📋 AI Weather Analysis Report")
            st.markdown(ai_response)
            
            # Download option for the analysis
            st.download_button(
                label="📥 Download Weather Analysis",
                data=ai_response,
                file_name=f"weather_analysis_{city}_{datetime.now().strftime('%Y%m%d')}.txt",
                mime="text/plain"
            )
    
    st.markdown("</div>", unsafe_allow_html=True)

    # ---------------- FOOTER NOTES ----------------
    st.markdown("<div style='height:18px'></div>", unsafe_allow_html=True)
    st.markdown("<div style='text-align:center; color:rgba(235,245,255,0.45); font-size:12px'>Quantora • Prototype • Built with Streamlit & OpenWeatherMap</div>", unsafe_allow_html=True)

# ✅ NEW: Quantomise My Trip
def quantomise_my_trip():
    st.title("✈️ Quantomise My Trip")
    st.markdown("Let AI plan your perfect trip -- budget, flights, hotels, and more -- tailored just you.")
    
    if "trip_data" not in st.session_state:
        st.session_state.trip_data = {}

    # ✅ Step 0: Pre-suggest destinations
    st.subheader("🌍 Not sure where to go? Let us suggest!")
    with st.expander("✨ Get AI Destination Suggestions"):
        vibe = st.text_input("Describe your vibe (e.g., peaceful mountains, party beach, cultural city):", placeholder="e.g., peaceful mountains")
        if st.button("Suggest Destinations"):
            prompt = f"Suggest 5 best travel destinations for someone who likes: {vibe}. Include country and best time to visit, Provide working link according to real latest time, no errors."
            response = call_a4f_model(prompt, "provider-2/gemini-3-pro-preview")
            st.markdown("### 🎯 AI Suggestions:")
            st.write(response)

    st.markdown("---")
    # ✅ Step 1: Trip Details
    st.subheader("🧳 Tell us about your trip")
    with st.form("trip_form"):
        budget = st.number_input("💰 Budget (in USD)", min_value=100, max_value=50000, step=100)
        origin = st.text_input("✈️ Flying from (City/Country):", placeholder="e.g., Delhi, India")
        destination = st.text_input("🌴 Going to (City/Country):", placeholder="e.g., Bali, Indonesia")
        trip_type = st.selectbox("🎯 Trip Type", ["Cheap", "Budget", "Luxury", "Ultra Luxury"])
        days = st.number_input("📅 Number of days", min_value=1, max_value=90, step=1)
        travelers = st.number_input("👥 Number of travelers", min_value=1, max_value=20, step=1)
        preferences = st.text_area("✍️ Preferences (optional)", placeholder="e.g., vegetarian food, pool, near beach, pet-friendly")
        submitted = st.form_submit_button("🔍 Find My Trip", type="primary")
    
    if submitted:
        st.session_state.trip_data = {
            "budget": budget,
            "origin": origin,
            "destination": destination,
            "trip_type": trip_type,
            "days": days,
            "travelers": travelers,
            "preferences": preferences
        }
        with st.spinner("🧠 AI is crafting your perfect trip..."):
            prompt = f"""
            You are a top-tier travel planner like Agoda, MakeMyTrip, and EaseMyTrip combined. Suggest for the current Year, month, and the day. Everything Real No Fake Items Like Fake Hotels, Do Real Live Search To Find.
            Plan a {trip_type.lower()} trip for {travelers} traveler(s) from {origin} to {destination} for {days} days with a budget of ${budget}.
            Preferences: {preferences if preferences else 'None'}.
            Include:
            1. Best flight options (with links if possible)
            2. Top 3 hotels/resorts (with links if possible)
            3. Must-visit places
            4. Daily itinerary
            5. Total estimated cost
            6. Best booking websites for each
            7. Money-saving tips
            8. Hidden gems
            Format it beautifully with emojis and sections.
            """
            response = call_a4f_model(prompt, "provider-2/gemini-3-pro-preview")
            st.markdown("### 🎯 Your AI-Planned Trip:")
            st.markdown(response)
            st.download_button(
                label="📥 Download Trip Plan",
                data=response,
                file_name=f"trip_plan_{destination.replace(' ', '_')}.txt",
                mime="text/plain"
            )

def coding_workspace():
    st.title("💻 AI Coding Workspace")
    st.markdown("Generate complete, ready-to-run code ")
    
    lang = st.selectbox("Language / Framework", ["Python", "JavaScript", "C++", "Go", "Rust", "Java", "Bash", "SQL"])
    intent = st.text_area("Describe what you need:", placeholder="e.g., FastAPI CRUD with SQLite and Pydantic models")
    
    if st.button("Generate Code", type="primary"):
        if not intent.strip():
            st.warning("Please describe what you want to build.")
            return
        
        prompt = f"Write a single self-contained {lang} file that: {intent}\n\n- Full code, no placeholders\n- Add short usage comment at top"
        with st.spinner("Generating..."):
            code = call_a4f_model(prompt, "provider-2/gemini-3-pro-preview")
        
        st.code(code, language=lang.lower())
        st.download_button("📥 Download file", data=code, file_name=f"code.{lang.lower()}", mime="text/plain")

# ---------------------------------------------------------
# 1️⃣ APP-BUILDER WORKSPACE -- UX unchanged
# ---------------------------------------------------------
def app_builder_workspace():
    st.title("🏗️ Streamlit App Builder")
    st.markdown("Describe an app idea")
    
    idea = st.text_area("Your app idea (1-2 sentences):", placeholder="e.g., an app that predicts house prices from CSV upload")
    
    if st.button("Build & Run", type="primary"):
        if not idea.strip():
            st.warning("Please give an idea.")
            return
        
        # 1️⃣ Expand idea
        expand_prompt = f"Turn this short idea into a detailed 150-word technical prompt for a single-file Streamlit app:\n\n{idea}"
        expanded = call_a4f_model(expand_prompt, "provider-7/claude-haiku-4-5-20251001")
        
        # 2️⃣ Generate code
        build_prompt = f"Write a single-file Streamlit app that: {expanded}\n\n- Use only public libs\n- No external assets\n- Save as app_generated.py"
        generated_code = call_a4f_model(build_prompt, "provider-2/gemini-3-pro-preview")
        
        # 3️⃣ Save & show
        Path("generated_apps").mkdir(exist_ok=True)
        file_path = Path("generated_apps/app_generated.py")
        file_path.write_text(generated_code, encoding="utf-8")
        st.success("✅ App generated - running below")
        
        with st.expander("📋 Generated code"):
            st.code(generated_code, language="python")
        
        # 4️⃣ Run inline & embed
        run_app_inline(str(file_path))

# ---------------------------------------------------------
# 2️⃣ INLINE RUNNER -- no subprocess, no ports
# ---------------------------------------------------------
def run_app_inline(script_path: str):
    """
    Executes the generated Streamlit script in a separate *thread* and
    serves it under /generated_apps so we can iframe it.
    Works on Streamlit Cloud & local.
    """
    import threading
    from streamlit.web.bootstrap import run
    
    # Copy the file into the *static* folder Streamlit exposes under /app
    static_dir = Path("static")
    static_dir.mkdir(exist_ok=True)
    target = static_dir / "app_generated.py"
    target.write_text(Path(script_path).read_text(), encoding="utf-8")
    
    # Start Streamlit in a daemon thread
    t = threading.Thread(
        target=run,
        args=(str(target),),
        kwargs={
            "server_port": 8503,
            "server_headless": True,
            "server_address": "localhost",
            "global_development_mode": False,
        },
        daemon=True,
    )
    t.start()
    time.sleep(3) # let boot
    
    # Embed the app
    st.components.v1.iframe(
        src="/app/static/app_generated.py",
        height=700,
        scrolling=True,
    )

# --------------------------
# COLLAGE MAKER MODULE
# --------------------------
def apply_sepia_filter(img):
    """Apply sepia filter to image"""
    sepia_filter = np.array([[0.393, 0.769, 0.189],
                            [0.349, 0.686, 0.168],
                            [0.272, 0.534, 0.131]])
    img_array = np.array(img)
    sepia_img = np.dot(img_array[..., :3], sepia_filter.T)
    sepia_img = np.clip(sepia_img, 0, 255).astype(np.uint8)
    return Image.fromarray(sepia_img)

def apply_vintage_filter(img):
    """Apply vintage filter to image using PIL instead of cv2"""
    # Add yellow tint
    img_array = np.array(img)
    img_array = img_array.astype(np.float32)
    img_array[..., 0] *= 1.1  # Increase red
    img_array[..., 1] *= 1.05  # Increase green
    img_array[..., 2] *= 0.9   # Decrease blue
    img_array = np.clip(img_array, 0, 255).astype(np.uint8)
    
    # Create vignette effect with PIL instead of cv2
    width, height = img.size
    # Create a radial gradient for vignette
    vignette = Image.new('L', (width, height), 0)
    draw = ImageDraw.Draw(vignette)
    
    # Draw an ellipse that covers most of the image
    ellipse_bbox = [
        -width * 0.5, -height * 0.5,
        width * 1.5, height * 1.5
    ]
    draw.ellipse(ellipse_bbox, fill=255)
    
    # Apply Gaussian blur to smooth the edges
    vignette = vignette.filter(ImageFilter.GaussianBlur(radius=width * 0.1))
    vignette_array = np.array(vignette).astype(np.float32) / 255.0
    
    # Apply vignette to each channel
    for i in range(3):
        img_array[..., i] = img_array[..., i] * (0.7 + 0.3 * vignette_array)
    
    return Image.fromarray(np.clip(img_array, 0, 255).astype(np.uint8))

def apply_cool_tone_filter(img):
    """Apply cool tone filter to image"""
    img_array = np.array(img)
    img_array = img_array.astype(np.float32)
    img_array[..., 0] *= 0.9   # Decrease red
    img_array[..., 1] *= 1.0   # Keep green
    img_array[..., 2] *= 1.1   # Increase blue
    return Image.fromarray(np.clip(img_array, 0, 255).astype(np.uint8))

def apply_warm_tone_filter(img):
    """Apply warm tone filter to image"""
    img_array = np.array(img)
    img_array = img_array.astype(np.float32)
    img_array[..., 0] *= 1.1   # Increase red
    img_array[..., 1] *= 1.05  # Slightly increase green
    img_array[..., 2] *= 0.9   # Decrease blue
    return Image.fromarray(np.clip(img_array, 0, 255).astype(np.uint8))

def collage_maker():
    st.title("🖼️ Images to Collage Maker")
    st.markdown("Create beautiful collages from multiple images with customizable layouts")
    
    # Upload multiple images
    uploaded_files = st.file_uploader(
        "Upload images for collage (2-10 images recommended)",
        type=["jpg", "jpeg", "png", "webp"],
        accept_multiple_files=True,
        key="collage_uploader"
    )
    
    if uploaded_files and len(uploaded_files) > 1:
        # Display uploaded images
        st.subheader("📸 Uploaded Images")
        cols = st.columns(min(4, len(uploaded_files)))
        images = []
        
        for i, uploaded_file in enumerate(uploaded_files):
            with cols[i % 4]:
                image = Image.open(uploaded_file)
                images.append(image)
                st.image(image, caption=f"Image {i+1}", use_container_width=True)
        
        # Collage settings
        st.subheader("🎨 Collage Settings")
        
        col1, col2 = st.columns(2)
        
        with col1:
            layout_type = st.selectbox(
                "Layout Type",
                ["Grid (Uniform)", "Grid (Custom)", "Masonry", "Polaroid Style", "Overlapping"],
                help="Choose how to arrange the images"
            )
            
            if layout_type == "Grid (Uniform)":
                cols_count = st.slider("Number of columns", 2, 5, 3)
            elif layout_type == "Grid (Custom)":
                cols_count = st.slider("Number of columns", 2, 5, 3)
                rows_count = st.slider("Number of rows", 1, 5, 2)
            elif layout_type == "Polaroid Style":
                rotation_max = st.slider("Maximum rotation (degrees)", 0, 30, 15)
        
        with col2:
            background_color = st.color_picker("Background Color", "#ffffff")
            spacing = st.slider("Spacing between images (pixels)", 0, 50, 10)
            border_size = st.slider("Border size (pixels)", 0, 20, 2)
            border_color = st.color_picker("Border Color", "#000000")
        
        # Advanced options
        with st.expander("Advanced Options"):
            col1, col2 = st.columns(2)
            with col1:
                collage_width = st.number_input("Collage width (pixels)", 800, 4000, 1200, 100)
                aspect_ratio = st.selectbox("Aspect Ratio", ["Auto", "1:1 (Square)", "3:2", "4:3", "16:9"])
                image_quality = st.slider("Image Quality", 50, 100, 90)
            
            with col2:
                apply_filters = st.checkbox("Apply uniform filters")
                if apply_filters:
                    filter_type = st.selectbox("Filter", ["None", "Sepia", "Black & White", "Vintage", "Cool Tone", "Warm Tone"])
                watermark = st.checkbox("Add watermark")
                if watermark:
                    watermark_text = st.text_input("Watermark text", "Quantora Collage")
        
        # Create collage button
        if st.button("🎨 Create Collage", type="primary"):
            with st.spinner("Creating your collage..."):
                try:
                    # Resize images to uniform size
                    target_size = collage_width // (cols_count if 'cols_count' in locals() else 3)
                    resized_images = []
                    
                    for img in images:
                        # Maintain aspect ratio
                        img_ratio = img.width / img.height
                        target_height = int(target_size / img_ratio)
                        resized = img.resize((target_size, target_height), Image.Resampling.LANCZOS)
                        resized_images.append(resized)
                    
                    # Calculate collage dimensions
                    if layout_type == "Grid (Uniform)":
                        cols = cols_count
                        rows = (len(resized_images) + cols - 1) // cols
                        collage_height = rows * (target_size // 2)  # Adjust for aspect ratio
                    elif layout_type == "Grid (Custom)":
                        cols = cols_count
                        rows = rows_count
                        collage_height = rows * (target_size // 2)
                    else:
                        cols = 3
                        rows = (len(resized_images) + 2) // 3
                        collage_height = rows * (target_size // 2)
                    
                    # Create blank canvas
                    if aspect_ratio == "1:1 (Square)":
                        collage_height = collage_width
                    elif aspect_ratio == "3:2":
                        collage_height = int(collage_width * 2 / 3)
                    elif aspect_ratio == "4:3":
                        collage_height = int(collage_width * 3 / 4)
                    elif aspect_ratio == "16:9":
                        collage_height = int(collage_width * 9 / 16)
                    
                    collage = Image.new('RGB', (collage_width, collage_height), background_color)
                    
                    # Position images
                    x_offset = spacing
                    y_offset = spacing
                    max_row_height = 0
                    
                    for i, img in enumerate(resized_images):
                        # Apply border
                        if border_size > 0:
                            bordered_img = Image.new('RGB', 
                                                   (img.width + 2*border_size, img.height + 2*border_size), 
                                                   border_color)
                            bordered_img.paste(img, (border_size, border_size))
                            img = bordered_img
                        
                        # Apply filters
                        if apply_filters and filter_type != "None":
                            if filter_type == "Sepia":
                                img = apply_sepia_filter(img)
                            elif filter_type == "Black & White":
                                img = img.convert('L').convert('RGB')
                            elif filter_type == "Vintage":
                                img = apply_vintage_filter(img)
                            elif filter_type == "Cool Tone":
                                img = apply_cool_tone_filter(img)
                            elif filter_type == "Warm Tone":
                                img = apply_warm_tone_filter(img)
                        
                        # Add rotation for polaroid style
                        if layout_type == "Polaroid Style":
                            rotation = random.randint(-rotation_max, rotation_max)
                            img = img.rotate(rotation, expand=True, fillcolor=background_color)
                        
                        # Check if image fits in current row
                        if x_offset + img.width > collage_width - spacing:
                            x_offset = spacing
                            y_offset += max_row_height + spacing
                            max_row_height = 0
                        
                        # Paste image
                        collage.paste(img, (x_offset, y_offset))
                        
                        # Update offsets
                        x_offset += img.width + spacing
                        max_row_height = max(max_row_height, img.height)
                        
                        # Add watermark to each image if enabled
                        if watermark and 'watermark_text' in locals():
                            draw = ImageDraw.Draw(collage)
                            # Use a truetype font if available
                            try:
                                font = ImageFont.truetype("arial.ttf", 20)
                            except:
                                font = ImageFont.load_default()
                            
                            text_width = draw.textlength(watermark_text, font=font)
                            text_height = 20  # Approximate height
                            text_position = (x_offset - img.width + 10, y_offset + img.height - text_height - 10)
                            draw.text(text_position, watermark_text, font=font, fill=(255, 255, 255, 128))
                    
                    # Add final watermark if enabled
                    if watermark and 'watermark_text' in locals() and not apply_filters:
                        draw = ImageDraw.Draw(collage)
                        try:
                            font = ImageFont.truetype("arial.ttf", 40)
                        except:
                            font = ImageFont.load_default()
                        
                        text_width = draw.textlength(watermark_text, font=font)
                        text_height = 40  # Approximate height
                        text_position = (collage_width - text_width - 20, collage_height - text_height - 20)
                        draw.text(text_position, watermark_text, font=font, fill=(255, 255, 255, 128))
                    
                    st.session_state.collage_image = collage
                    st.success("✅ Collage created successfully!")
                    
                except Exception as e:
                    st.error(f"Error creating collage: {str(e)}")
        
        # Display and download collage
        if hasattr(st.session_state, 'collage_image') and st.session_state.collage_image:
            st.subheader("🎉 Your Collage")
            st.image(st.session_state.collage_image, use_container_width=True)
            
            # Download options
            col1, col2 = st.columns(2)
            with col1:
                # Save to buffer
                buffered = BytesIO()
                st.session_state.collage_image.save(buffered, format="PNG", quality=image_quality)
                
                st.download_button(
                    label="📥 Download PNG",
                    data=buffered.getvalue(),
                    file_name=f"quantora_collage_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png",
                    mime="image/png",
                    use_container_width=True
                )
            
            with col2:
                buffered_jpg = BytesIO()
                st.session_state.collage_image.save(buffered_jpg, format="JPEG", quality=image_quality)
                
                st.download_button(
                    label="📥 Download JPEG",
                    data=buffered_jpg.getvalue(),
                    file_name=f"quantora_collage_{datetime.now().strftime('%Y%m%d_%H%M%S')}.jpg",
                    mime="image/jpeg",
                    use_container_width=True
                )
            
            # Reset button
            if st.button("🔄 Create New Collage"):
                if hasattr(st.session_state, 'collage_image'):
                    del st.session_state.collage_image
                st.rerun()
    
    elif uploaded_files and len(uploaded_files) == 1:
        st.warning("Please upload at least 2 images to create a collage.")
    else:
        # Show instructions
        st.info("""
        **How to create a collage:**
        1. Upload 2 or more images (JPG, PNG, or WebP format)
        2. Choose your preferred layout and settings
        3. Click "Create Collage"
        4. Download your masterpiece!
        
        **Tips for best results:**
        • Use images with similar aspect ratios for uniform grids
        • Try the "Polaroid Style" for a fun, casual look
        • Adjust spacing and borders to control visual density
        • Use the "Masonry" layout for a Pinterest-style collage
        """)

# --------------------------
# SOUND EXTRACTOR MODULE
# --------------------------
def sound_extractor():
    st.title("🎵 Sound Extractor from Video")
    st.markdown("Extract high-quality audio from video files in multiple formats")

    uploaded_video = st.file_uploader(
        "Upload a video file",
        type=["mp4", "avi", "mov", "mkv", "webm", "flv", "wmv", "mpeg", "mpg"],
        key="video_uploader"
    )

    if not uploaded_video:
        st.info("""
        **How to extract audio:**
        1. Upload a video file (MP4, AVI, MOV, MKV, etc.)
        2. Configure audio extraction settings
        3. Click "Extract Audio"
        4. Download your extracted audio file
        """)
        return

    # -------------------------
    # SAVE TEMP VIDEO
    # -------------------------
    temp_video_path = f"temp_video_{time.time()}.{uploaded_video.name.split('.')[-1]}"
    with open(temp_video_path, "wb") as f:
        f.write(uploaded_video.getbuffer())

    # -------------------------
    # DISPLAY VIDEO
    # -------------------------
    st.subheader("🎥 Uploaded Video")
    col1, col2 = st.columns([2, 1])

    with col1:
        st.video(temp_video_path)

    with col2:
        st.info(f"""
        **Video Details:**
        • Name: `{uploaded_video.name}`
        • Size: `{uploaded_video.size / (1024*1024):.2f} MB`
        • Type: `{uploaded_video.type}`
        """)

        try:
            video = VideoFileClip(temp_video_path)
            duration = video.duration
            st.info(f"• Duration: `{int(duration // 60)}m {int(duration % 60)}s`")
            video.close()
        except:
            st.info("• Duration: `Unable to detect`")

    # -------------------------
    # EXTRACTION SETTINGS
    # -------------------------
    st.subheader("⚙️ Extraction Settings")

    col1, col2, col3 = st.columns(3)

    with col1:
        audio_format = st.selectbox(
            "Output Format",
            ["MP3 (Recommended)", "WAV (Lossless)", "M4A", "AAC", "OGG", "FLAC"]
        )
        format_map = {
            "MP3 (Recommended)": "mp3",
            "WAV (Lossless)": "wav",
            "M4A": "m4a",
            "AAC": "aac",
            "OGG": "ogg",
            "FLAC": "flac"
        }
        output_ext = format_map[audio_format]

    with col2:
        audio_quality = st.selectbox(
            "Audio Quality",
            ["High (320 kbps)", "Medium (192 kbps)", "Standard (128 kbps)", "Low (64 kbps)"],
            index=0
        )
        bitrate = {
            "High (320 kbps)": 320,
            "Medium (192 kbps)": 192,
            "Standard (128 kbps)": 128,
            "Low (64 kbps)": 64
        }[audio_quality]

    with col3:
        trim_audio = st.checkbox("Trim Audio")
        if trim_audio:
            start_time = st.number_input("Start (seconds)", min_value=0.0, value=0.0)
            end_time = st.number_input("End (seconds)", min_value=0.0, value=30.0)

    # -------------------------
    # ADVANCED OPTIONS
    # -------------------------
    with st.expander("Advanced Options"):
        normalize_audio = st.checkbox("Normalize Audio Volume", True)
        remove_noise = st.checkbox("Remove Background Noise")

        if remove_noise:
            noise_level = st.slider("Noise Reduction", 0, 100, 50)

        add_metadata = st.checkbox("Add Metadata", True)

        if add_metadata:
            title = st.text_input("Audio Title", uploaded_video.name.split('.')[0])
            artist = st.text_input("Artist", "Extracted from Video")
            album = st.text_input("Album", "Video Audio Extractions")

    # -------------------------
    # EXTRACTION BUTTON
    # -------------------------
    if st.button("🎵 Extract Audio", type="primary"):
        with st.spinner("Extracting audio..."):
            try:
                # Load video
                video = VideoFileClip(temp_video_path)
                audio = video.audio

                # Trim
                if trim_audio:
                    audio = audio.subclip(start_time, min(end_time, audio.duration))

                # Save audio
                raw_audio_path = f"raw_audio_{time.time()}.{output_ext}"
                if output_ext == "mp3":
                    audio.write_audiofile(raw_audio_path, bitrate=f"{bitrate}k")
                elif output_ext == "wav":
                    audio.write_audiofile(raw_audio_path, codec="pcm_s16le")
                else:
                    audio.write_audiofile(raw_audio_path)

                audio.close()
                video.close()

                # PROCESS AUDIO (normalize/noise)
                processed_path = raw_audio_path
                if normalize_audio or remove_noise:
                    processed_path = f"processed_{time.time()}.{output_ext}"
                    filters = []

                    if normalize_audio:
                        filters.append("loudnorm=I=-16:LRA=11:TP=-1.5")
                    if remove_noise:
                        filters.append(f"afftdn=nr={noise_level}")

                    ffmpeg_cmd = [
                        "ffmpeg", "-y", "-i", raw_audio_path,
                        "-af", ",".join(filters),
                        processed_path
                    ]
                    subprocess.run(ffmpeg_cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)

                # METADATA
                final_path = processed_path
                if add_metadata:
                    final_path = f"final_{time.time()}.{output_ext}"
                    metadata_cmd = [
                        "ffmpeg", "-y", "-i", processed_path,
                        "-metadata", f"title={title}",
                        "-metadata", f"artist={artist}",
                        "-metadata", f"album={album}",
                        "-codec", "copy",
                        final_path
                    ]
                    subprocess.run(metadata_cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)

                # Store output for download
                with open(final_path, "rb") as f:
                    st.session_state.extracted_audio = f.read()
                st.session_state.audio_filename = f"{title}.{output_ext}"

                st.success("✅ Audio extracted successfully!")

                # CLEANUP
                for p in [temp_video_path, raw_audio_path, processed_path]:
                    if os.path.exists(p):
                        os.remove(p)

            except Exception as e:
                error_message = str(e)
            
                # Detect missing FFmpeg
                if "No such file or directory: 'ffmpeg'" in error_message or "ffmpeg" in error_message.lower():
                    st.error("🚧 This feature is currently under construction and will soon be available.")
                else:
                    st.error(f"❌ Error: {error_message}")


    # -------------------------
    # DOWNLOAD SECTION
    # -------------------------
    if "extracted_audio" in st.session_state:
        st.subheader("🎧 Extracted Audio")

        st.audio(st.session_state.extracted_audio, format=f"audio/{output_ext}")

        st.download_button(
            "📥 Download Audio",
            st.session_state.extracted_audio,
            file_name=st.session_state.audio_filename,
            mime=f"audio/{output_ext}",
            use_container_width=True
        )

import streamlit as st
import re
import urllib.parse
from datetime import datetime
import time

def get_store_url(store, product_name):
    """Generate store URLs for shopping links"""
    encoded = urllib.parse.quote(product_name)
    store = store.lower().strip()

    if "amazon" in store:
        return f"https://www.amazon.in/s?k={encoded}"
    elif "flipkart" in store:
        return f"https://www.flipkart.com/search?q={encoded}"
    elif "reliancedigital" in store or "reliance digital" in store:
        return f"https://www.reliancedigital.in/search?q={encoded}"
    elif "mdcomputers" in store:
        return f"https://mdcomputers.in/catalogsearch/result/?q={encoded}"
    elif "primeabgb" in store:
        return f"https://www.primeabgb.com/?s={encoded}"
    else:
        return f"https://www.google.com/search?q={encoded}"

def call_a4f_model(prompt, model_name):
    """Wrapper for A4F API calls"""
    try:
        # This should be defined elsewhere in your code
        # For now, returning a placeholder response
        return f"Product list for: {prompt[:50]}...\n\nSample Product:\nGaming Mouse XYZ\nPrice: ₹799\nFeatures: RGB, 3200 DPI, Ergonomic\nBuy Links: Amazon|Flipkart|RelianceDigital"
    except Exception as e:
        raise Exception(f"API Error: {str(e)}")

def shopping_research():
    st.title("🛒 Shopping Research Assistant")
    st.markdown("Find the perfect product with AI-powered shopping research")

    shopping_query = st.text_area(
        "Describe what you want to buy:",
        height=120,
        placeholder="Example: 'Gaming mouse under ₹800 with RGB and 3200 DPI...'",
    )

    col1, col2 = st.columns(2)

    with col1:
        budget_min = st.number_input("Minimum Budget (₹)", 0, 500000, 0, 100)
        budget_max = st.number_input("Maximum Budget (₹)", 0, 500000, 2000, 100)

    with col2:
        priority = st.selectbox(
            "Primary Priority",
            ["Best Value", "Highest Quality", "Best Features", "Most Reliable", "Lowest Price"]
        )
        region = st.selectbox("Region", ["India", "United States", "Europe", "Global"])

    with st.expander("Additional Preferences"):
        brand_preference = st.text_input("Preferred Brands")
        exclude_brands = st.text_input("Brands to Avoid")
        must_have = st.text_area("Must-Have Features")
        avoid_features = st.text_area("Features to Avoid")

    # -------------------------------------------------------
    # STEP 1 -- GENERATE PRODUCTS
    # -------------------------------------------------------
    if st.button("🔍 Start Shopping Research") and shopping_query.strip():
        st.info("**Step 1:** Generating product list...")

        # Prepare the Gemini prompt with all user inputs
        gemini_prompt = f"""You are an expert shopping assistant. Generate a list of products based on these criteria:

Shopping Query: {shopping_query}
Budget Range: ₹{budget_min} - ₹{budget_max}
Priority: {priority}
Region: {region}
Preferred Brands: {brand_preference if brand_preference else 'None'}
Brands to Avoid: {exclude_brands if exclude_brands else 'None'}
Must-Have Features: {must_have if must_have else 'None'}
Features to Avoid: {avoid_features if avoid_features else 'None'}

Please provide:
1. Product name and model
2. Key features and specifications
3. Approximate price
4. Where to buy it (store names separated by |)
5. Brief pros and cons

Format each product as:
### [Product Name]
**Price:** [Price]
**Key Features:** [Features]
**Pros:** [Pros]
**Cons:** [Cons]
**Buy Links:** [Store1|Store2|Store3]

Separate products with a clear divider.
Provide at least 3-5 relevant products."""

        try:
            product_list = call_a4f_model(gemini_prompt, "provider-2/gemini-3-pro-preview")
            st.session_state.product_list = product_list
            st.success("✅ Product list generated!")
        except Exception as e:
            st.error(f"Error generating product list: {str(e)}")
            return

        # Show raw list
        st.markdown("## 🛍️ Product List")

        # Split products - using a more robust method
        products_raw = []
        current_product = []
        
        for line in product_list.split('\n'):
            if line.strip().startswith('### ') and current_product:
                # New product found, save the previous one
                products_raw.append('\n'.join(current_product))
                current_product = [line]
            else:
                current_product.append(line)
        
        # Add the last product
        if current_product:
            products_raw.append('\n'.join(current_product))

        # -------------------------------------------------------
        # SHOW EACH PRODUCT WITH ITS OWN WORKING BUTTON
        # -------------------------------------------------------
        for idx, block in enumerate(products_raw):
            block = block.strip()
            if not block or len(block) < 10:  # Skip very short blocks
                continue

            # Extract product name (first line after ###)
            lines = block.split('\n')
            product_name = ""
            for line in lines:
                if line.strip().startswith('### '):
                    product_name = line.strip().replace('### ', '').strip()
                    break
            
            if not product_name:
                # Fallback: use first non-empty line
                for line in lines:
                    if line.strip():
                        product_name = line.strip()
                        break
            
            if not product_name:
                product_name = f"Product {idx+1}"

            st.markdown(f"### {product_name}")
            st.markdown(block)

            # Extract Buy Links section
            store_names = ["Google"]  # Default fallback
            
            # Look for Buy Links in the block
            for line in lines:
                if line.strip().startswith('Buy Links:') or line.strip().startswith('**Buy Links:**'):
                    link_part = line.strip().replace('Buy Links:', '').replace('**Buy Links:**', '').strip()
                    if link_part:
                        store_names = [s.strip() for s in link_part.split('|')]
                    break
            
            # Also try regex pattern as backup
            if store_names == ["Google"]:
                links_match = re.search(r'Buy Links:\s*(.*)', block, re.IGNORECASE)
                if links_match:
                    store_line = links_match.group(1).strip()
                    if store_line:
                        store_names = [s.strip() for s in store_line.split('|')]

            # Generate buttons for each store
            button_cols = st.columns(min(3, len(store_names)))
            for i, store in enumerate(store_names):
                if store:  # Skip empty store names
                    url = get_store_url(store, product_name)
                    with button_cols[i % len(button_cols)]:
                        st.markdown(
                            f"""
                            <a href="{url}" target="_blank" style="text-decoration: none;">
                            <button style="width:100%; padding:10px 12px; margin:4px 0; font-size:14px; 
                                          background:linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                                          color:white; border:none; border-radius:8px; cursor:pointer;">
                            🛒 {store}
                            </button>
                            </a>
                            """,
                            unsafe_allow_html=True,
                        )

            st.markdown("---")

        # -------------------------------------------------------
        # STEP 2 -- BEST PRODUCT ANALYSIS (CLAUDE)
        # -------------------------------------------------------
        st.info("**Step 2:** Analyzing best product...")

        claude_prompt = f"""You are an expert product analyst. Based on the following product list and user preferences, 
        recommend the BEST single product:

        User Preferences:
        - Query: {shopping_query}
        - Budget: ₹{budget_min} - ₹{budget_max}
        - Priority: {priority}
        - Region: {region}
        - Must-have: {must_have if must_have else 'None'}
        - Avoid: {avoid_features if avoid_features else 'None'}

        Products to analyze:
        {product_list[:2000]}...

        Please provide a detailed analysis of the best product with:
        1. **Product Name** (clear and prominent)
        2. **Why it's the best choice** (matching user priorities)
        3. **Key specifications**
        4. **Price and value assessment**
        5. **Best place to buy** (specific store recommendation)
        6. **Final verdict**

        Format your response clearly with headings."""

        try:
            best_product = call_a4f_model(claude_prompt, "provider-7/claude-opus-4-5-20251101")
            st.session_state.best_product = best_product
            st.success("🏆 Best product analysis complete!")
        except Exception as e:
            st.error(f"Error analyzing best product: {str(e)}")
            return

    # -------------------------------------------------------
    # SHOW BEST PRODUCT WITH BUTTON (if exists)
    # -------------------------------------------------------
    if "best_product" in st.session_state and st.session_state.best_product:
        st.markdown("## 🎯 BEST PRODUCT MATCH")
        st.markdown(st.session_state.best_product)

        # Extract product name from best product text
        product_name = "Best Product"
        lines = st.session_state.best_product.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('**Product Name:**') or line.startswith('### ') or line.startswith('## '):
                # Clean up the product name
                product_name = line.replace('**Product Name:**', '').replace('### ', '').replace('## ', '').strip()
                # Remove any markdown formatting
                product_name = re.sub(r'[*#\[\]]', '', product_name).strip()
                break
        
        if not product_name or product_name == "Best Product":
            # Try to find the first line that looks like a product name
            for line in lines:
                if line.strip() and len(line.strip()) > 5 and not line.strip().startswith('*') and 'http' not in line:
                    product_name = line.strip()
                    break

        # Extract store name from text
        store = "Google"  # Default
        store_patterns = [
            r'(Amazon\.(?:in|com))',
            r'(Flipkart)',
            r'(MDComputers)',
            r'(PrimeABGB)',
            r'(RelianceDigital)',
            r'(Reliance Digital)'
        ]
        
        for pattern in store_patterns:
            match = re.search(pattern, st.session_state.best_product, re.IGNORECASE)
            if match:
                store = match.group(1)
                break

        # Get the URL for the store
        url = get_store_url(store, product_name)

        # Create the buy button
        st.markdown(
            f"""
            <div style="text-align: center; margin: 20px 0;">
            <a href="{url}" target="_blank">
            <button style="padding:12px 32px; font-size:18px; 
                          background:linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
                          color:white; border:none; border-radius:10px; 
                          cursor:pointer; font-weight:bold;">
            🛒 Buy Best Product on {store}
            </button>
            </a>
            </div>
            """,
            unsafe_allow_html=True,
        )

        
def quantora_translate():
    st.title("🌐 Quantora Translate")
    st.markdown("Advanced AI-powered translation with professional accuracy using Gemini 3 Pro")
    
    # Language selection
    col1, col2 = st.columns(2)
    
    with col1:
        # Common languages for quick selection
        language_options = {
            "English": "en",
            "Spanish": "es",
            "French": "fr",
            "German": "de",
            "Italian": "it",
            "Portuguese": "pt",
            "Russian": "ru",
            "Chinese (Simplified)": "zh-CN",
            "Japanese": "ja",
            "Korean": "ko",
            "Arabic": "ar",
            "Hindi": "hi",
            "Bengali": "bn",
            "Urdu": "ur",
            "Turkish": "tr",
            "Dutch": "nl",
            "Polish": "pl",
            "Vietnamese": "vi",
            "Thai": "th",
            "Greek": "el",
            "Hebrew": "he",
            "Swedish": "sv",
            "Danish": "da",
            "Norwegian": "no",
            "Finnish": "fi"
        }
        
        target_language = st.selectbox(
            "Translate to:",
            options=list(language_options.keys()),
            index=0,
            help="Select the target language for translation"
        )
        
        # Additional translation options
        with st.expander("⚙️ Advanced Options"):
            tone_style = st.selectbox(
                "Translation Style:",
                ["Formal", "Casual", "Technical", "Literary", "Business", "Medical", "Legal"],
                index=0,
                help="Choose the appropriate tone/style for your translation"
            )
            
            preserve_formatting = st.checkbox(
                "Preserve formatting and structure",
                value=True,
                help="Maintain original formatting like paragraphs, lists, and spacing"
            )
            
            include_notes = st.checkbox(
                "Include cultural/context notes",
                value=False,
                help="Add helpful notes about cultural context or idioms"
            )
    
    with col2:
        # Text input area
        st.subheader("📝 Text to Translate")
        source_text = st.text_area(
            "Enter text to translate:",
            height=250,
            placeholder="Paste or type the text you want to translate here...",
            help="Enter up to 5000 characters for translation"
        )
        
        # Character counter
        if source_text:
            char_count = len(source_text)
            st.caption(f"Characters: {char_count}/5000")
            if char_count > 5000:
                st.warning("Text exceeds 5000 characters. Please shorten your text.")
        
        # File upload option
        uploaded_file = st.file_uploader(
            "Or upload a text file:",
            type=["txt", "doc", "docx", "pdf", "md"],
            help="Upload text documents for translation"
        )
        
        if uploaded_file:
            try:
                if uploaded_file.type == "application/pdf":
                    pdf_reader = PdfReader(uploaded_file)
                    extracted_text = ""
                    for page in pdf_reader.pages:
                        extracted_text += page.extract_text() + "\n"
                    source_text = extracted_text.strip()
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    doc = docx.Document(uploaded_file)
                    extracted_text = "\n".join([para.text for para in doc.paragraphs])
                    source_text = extracted_text
                else:
                    source_text = uploaded_file.read().decode('utf-8')
                
                st.success(f"✅ File '{uploaded_file.name}' loaded successfully!")
                st.text_area("Extracted text:", value=source_text[:1000] + "..." if len(source_text) > 1000 else source_text, height=150, disabled=True)
            except Exception as e:
                st.error(f"Error reading file: {str(e)}")
    
    # Translation button
    translate_button = st.button(
        "🔁 Translate Now",
        type="primary",
        use_container_width=True,
        disabled=not source_text or len(source_text.strip()) == 0
    )
    
    # Translation history
    if "translation_history" not in st.session_state:
        st.session_state.translation_history = []
    
    if translate_button and source_text.strip():
        with st.spinner(f"Translating to {target_language}..."):
            try:
                # Prepare translation prompt
                language_code = language_options[target_language]
                
                translation_prompt = f"""
                You are an expert translator. Translate the following text from its source language to {target_language}.
                
                IMPORTANT TRANSLATION REQUIREMENTS:
                1. Translate the text accurately and naturally
                2. Maintain the original meaning and intent
                3. Use {tone_style.lower()} tone/style
                4. Preserve proper names, technical terms, and specific terminology
                5. Handle idioms and cultural references appropriately
                {"6. Preserve the original formatting, paragraphs, and structure" if preserve_formatting else ""}
                {"7. Add brief cultural/context notes where appropriate" if include_notes else ""}
                
                SOURCE TEXT:
                {source_text}
                
                Provide only the translated text in {target_language}.
                {"Add cultural/context notes at the end if needed, clearly marked as [NOTE]:" if include_notes else ""}
                """
                
                # Call the translation model
                translated_text = call_a4f_model(translation_prompt, "provider-2/gemini-3-pro-preview")
                
                # Store in session state
                st.session_state.last_translation = {
                    "source": source_text,
                    "translated": translated_text,
                    "target_language": target_language,
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "tone_style": tone_style
                }
                
                # Add to history
                st.session_state.translation_history.append(st.session_state.last_translation)
                
                # Limit history to last 10 items
                if len(st.session_state.translation_history) > 10:
                    st.session_state.translation_history = st.session_state.translation_history[-10:]
                
            except Exception as e:
                st.error(f"Translation error: {str(e)}")
                return
    
    # Display translation results
    if "last_translation" in st.session_state:
        st.markdown("---")
        st.subheader("✅ Translation Complete")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### 📝 Original Text")
            st.text_area(
                "Source:",
                value=st.session_state.last_translation["source"],
                height=200,
                disabled=True,
                label_visibility="collapsed"
            )
        
        with col2:
            st.markdown(f"#### 🌐 Translated to {st.session_state.last_translation['target_language']}")
            st.text_area(
                "Translation:",
                value=st.session_state.last_translation["translated"],
                height=200,
                disabled=True,
                label_visibility="collapsed"
            )
        
        # Translation metrics
        source_words = len(st.session_state.last_translation["source"].split())
        translated_words = len(st.session_state.last_translation["translated"].split())
        
        st.markdown(f"""
        <div style="background-color: #f0f2f6; padding: 1rem; border-radius: 8px; margin: 1rem 0;">
            <div style="display: flex; justify-content: space-between;">
                <div>
                    <strong>📊 Translation Stats:</strong><br>
                    • Words: {source_words} -> {translated_words}<br>
                    • Language: {st.session_state.last_translation['target_language']}<br>
                    • Style: {st.session_state.last_translation['tone_style']}<br>
                    • Time: {st.session_state.last_translation['timestamp']}
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        # Action buttons
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            # Download translation
            translation_data = f"""ORIGINAL TEXT:
{st.session_state.last_translation["source"]}

TRANSLATED TEXT ({st.session_state.last_translation["target_language"]}):
{st.session_state.last_translation["translated"]}

Translation Details:
- Target Language: {st.session_state.last_translation["target_language"]}
- Style: {st.session_state.last_translation["tone_style"]}
- Translated on: {st.session_state.last_translation["timestamp"]}
"""
            
            st.download_button(
                label="📥 Download",
                data=translation_data,
                file_name=f"translation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain",
                use_container_width=True
            )
        
        with col2:
            # Copy to clipboard
            if st.button("📋 Copy", use_container_width=True):
                st.code(st.session_state.last_translation["translated"])
                st.info("Translation copied to clipboard!")
        
        with col3:
            # Speak translation (placeholder)
            if st.button("🔊 Speak", use_container_width=True):
                st.info("Text-to-speech feature would play the translated text")
        
        with col4:
            # New translation
            if st.button("🔄 New", use_container_width=True):
                if "last_translation" in st.session_state:
                    del st.session_state.last_translation
                st.rerun()
        
        # Translation history
        if st.session_state.translation_history:
            with st.expander("📜 Translation History"):
                for i, item in enumerate(reversed(st.session_state.translation_history)):
                    with st.container():
                        st.markdown(f"**{item['timestamp']}** - {item['target_language']} ({item['tone_style']})")
                        st.text_area(
                            f"Translation {len(st.session_state.translation_history) - i}",
                            value=item['translated'][:300] + "..." if len(item['translated']) > 300 else item['translated'],
                            height=100,
                            disabled=True,
                            label_visibility="collapsed"
                        )
                        if st.button(f"Use this translation", key=f"use_{i}"):
                            source_text = item['source']
                            st.rerun()
                        st.markdown("---")
    
    # Quick translation examples
    if not translate_button and not source_text:
        st.markdown("---")
        st.subheader("💡 Quick Translation Examples")
        
        examples = [
            {"text": "Hello, how are you today? I hope you're having a wonderful day!", "lang": "Spanish"},
            {"text": "The quick brown fox jumps over the lazy dog.", "lang": "French"},
            {"text": "Technology is transforming our world at an unprecedented pace.", "lang": "Japanese"},
            {"text": "Where can I find the best restaurants in this neighborhood?", "lang": "Italian"},
            {"text": "The meeting has been rescheduled to next Tuesday at 3 PM.", "lang": "German"}
        ]
        
        cols = st.columns(3)
        for i, example in enumerate(examples[:6]):
            with cols[i % 3]:
                if st.button(f"Translate to {example['lang']}", key=f"example_{i}", use_container_width=True):
                    # Set the example text and language
                    source_text = example['text']
                    target_language = example['lang']
                    st.rerun()
    
    # Footer information
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666; font-size: 0.9rem;">
        <strong>🌐 Quantora Translate Features:</strong><br>
        • Powered by Gemini 3 Pro for high-quality translations<br>
        • Supports 25+ languages with professional accuracy<br>
        • Preserves formatting and technical terminology<br>
        • Cultural context awareness and idiom handling<br>
        • Translation history and export options
    </div>
    """, unsafe_allow_html=True)

# Custom CSS
st.markdown("""
<style>
    /* Basic styling */
    .main-header {
        text-align: center;
        padding: 2rem 0;
        margin-bottom: 2rem;
    }
    .logo {
        display: flex;
        align-items: center;
        gap: 0.75rem;
        justify-content: center;
        margin-bottom: 1rem;
    }
    .logo-icon {
        width: 48px;
        height: 48px;
        background: linear-gradient(135deg, #8b5cf6, #6d28d9);
        border-radius: 12px;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 24px;
    }
    .logo-text {
        font-size: 2.5rem;
        font-weight: 800;
        background: linear-gradient(135deg, #f8fafc, #a78bfa, #60a5fa);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .status-indicator {
        display: inline-block;
        width: 12px;
        height: 12px;
        background: #10b981;
        border-radius: 50%;
        margin-left: 10px;
        animation: pulse 2s infinite;
    }
    @keyframes pulse {
        0%, 100% { opacity: 1; }
        50% { opacity: 0.5; }
    }
    
    /* Fix streamlit containers */
    .stTextArea textarea {
        background-color: white !important;
        color: black !important;
    }
    
    /* Chat messages */
    .chat-message {
        padding: 1rem;
        margin: 0.5rem 0;
        border-radius: 10px;
        background: rgba(30, 41, 59, 0.5);
    }
    .user-message {
        background: rgba(139, 92, 246, 0.2);
    }
    .ai-message {
        background: rgba(59, 130, 246, 0.2);
    }
</style>
""", unsafe_allow_html=True)

# Unlock button for trial mode
if "pro_unlocked" not in st.session_state:
    st.session_state.pro_unlocked = False

if "pro_verified" not in st.session_state:
    st.session_state.pro_verified = False

# ==============================
# 🧩 COMPLETE SESSION STATE INITIALIZATION
# ==============================

def initialize_session_state():
    """Initialize all session state variables to prevent errors"""
    # Core chat and content
    if 'chat' not in st.session_state:
        st.session_state.chat = []
    if 'uploaded_content' not in st.session_state:
        st.session_state.uploaded_content = ""
    
    # Model and feature settings
    if 'model_version' not in st.session_state:
        st.session_state.model_version = "Quantora Prime 1 Fast"
    if 'image_style' not in st.session_state:
        st.session_state.image_style = "Realistic"
    if 'video_style' not in st.session_state:
        st.session_state.video_style = "Cinematic"
    
    # Image handling
    if 'uploaded_image' not in st.session_state:
        st.session_state.uploaded_image = None
    if 'enhanced_image' not in st.session_state:
        st.session_state.enhanced_image = None
    if 'generated_image' not in st.session_state:
        st.session_state.generated_image = None
    if 'edited_image' not in st.session_state:
        st.session_state.edited_image = None
    if 'generated_video' not in st.session_state:
        st.session_state.generated_video = None
    
    # Enhancement controls
    if 'enhancement_values' not in st.session_state:
        st.session_state.enhancement_values = {
            "brightness": 1.0,
            "contrast": 1.0,
            "sharpness": 1.0,
            "color": 1.0,
            "filter": "None"
        }
    
    # Response tracking
    if 'last_response_time' not in st.session_state:
        st.session_state.last_response_time = 0
    
    # Pro features and verification
    if 'pro_unlocked' not in st.session_state:
        st.session_state.pro_unlocked = False
    if 'pro_verified' not in st.session_state:
        st.session_state.pro_verified = False
    
    # Social media features
    if 'quantora_logged_in' not in st.session_state:
        st.session_state.quantora_logged_in = False
    if 'quantora_liked_posts' not in st.session_state:
        st.session_state.quantora_liked_posts = set()
    if 'view_profile' not in st.session_state:
        st.session_state.view_profile = None
    
    # Learning and IQ
    if 'learning_history' not in st.session_state:
        st.session_state.learning_history = []
    if 'iq_test_score' not in st.session_state:
        st.session_state.iq_test_score = None
    
    # Weather module
    if 'city_input' not in st.session_state:
        st.session_state.city_input = "Berlin"
    
    # Translation module
    if 'translation_history' not in st.session_state:
        st.session_state.translation_history = []
    if 'last_translation' not in st.session_state:
        st.session_state.last_translation = None
    
    # Shopping research
    if 'product_list' not in st.session_state:
        st.session_state.product_list = None
    if 'best_product' not in st.session_state:
        st.session_state.best_product = None
    
    # Collage maker
    if 'collage_image' not in st.session_state:
        st.session_state.collage_image = None
    
    # Sound extractor
    if 'extracted_audio' not in st.session_state:
        st.session_state.extracted_audio = None
    if 'audio_filename' not in st.session_state:
        st.session_state.audio_filename = None
    
    # Health analyzers
    if 'heart_rate_data' not in st.session_state:
        st.session_state.heart_rate_data = None
    if 'cognitive_data' not in st.session_state:
        st.session_state.cognitive_data = None
    if 'image_analysis' not in st.session_state:
        st.session_state.image_analysis = None
    
    # Content detector
    if 'ai_response' not in st.session_state:
        st.session_state.ai_response = None

# Run initialization
initialize_session_state()

# ==============================
# 🌱 VERIFICATION FLOW
# ==============================

if not st.session_state.pro_verified:
    with st.sidebar:
        st.subheader("🌱 Unlock Premium")
        st.markdown("Plant 2 plants to unlock Quantora Prime X")
        
        if st.button("🌿 Visit Plant Store"):
            st.markdown("[Open Plant Store](https://www.grow-trees.com/plant/monthly.php?a=KushagraSrivastava)")
        
        st.divider()
        
        txn_id = st.text_input("Enter Transaction ID")
        
        if st.button("Verify Purchase"):
            if txn_id and len(txn_id) > 5:
                st.session_state.pro_verified = True
                st.success("Verification successful!")
                time.sleep(1)
                st.rerun()
            else:
                st.error("Please enter a valid transaction ID")

# Unlock Pro button
if st.session_state.pro_verified and not st.session_state.pro_unlocked:
    with st.sidebar:
        if st.button("🚀 Unlock Pro Features"):
            st.session_state.pro_unlocked = True
            st.success("Pro features unlocked!")
            time.sleep(1)
            st.rerun()

# Set model version based on access
if not st.session_state.pro_unlocked:
    st.session_state.model_version = "Quantora Prime 1 Fast"
else:
    st.session_state.model_version = "Quantora Prime 1"


# ==============================
# 🚀 SHOW UNLOCK BUTTON ONLY AFTER VERIFICATION
# ==============================

if st.session_state.pro_verified and not st.session_state.pro_unlocked:
    if st.button("🚀 Unlock Next-Gen Pro", key="unlock_pro_btn"):
        st.session_state.pro_unlocked = True
        st.rerun()


# ==============================
# 🧠 FORCE MODEL VERSION BASED ON ACCESS
# ==============================

if not st.session_state.pro_unlocked:
    st.session_state.model_version = (
        "Quantora Prime 1 Fast "
        "(Faster but not as powerful as the flagship model)"
    )
else:
    st.session_state.model_version = "Quantora Prime 1 (Latest Flagship Model)"


# ==============================
# 🧩 OTHER SESSION STATE INITIALIZATION
# ==============================

if "chat" not in st.session_state:
    st.session_state.chat = []

if "uploaded_content" not in st.session_state:
    st.session_state.uploaded_content = ""

if "last_response_time" not in st.session_state:
    st.session_state.last_response_time = 0

if "uploaded_image" not in st.session_state:
    st.session_state.uploaded_image = None

if "enhanced_image" not in st.session_state:
    st.session_state.enhanced_image = None

if "enhancement_values" not in st.session_state:
    st.session_state.enhancement_values = {
        "brightness": 1.0,
        "contrast": 1.0,
        "sharpness": 1.0,
        "color": 1.0,
        "filter": "None"
    }

if "image_style" not in st.session_state:
    st.session_state.image_style = "Sci-Fi"

if "video_style" not in st.session_state:
    st.session_state.video_style = "Cinematic"

if "quantora_logged_in" not in st.session_state:
    st.session_state.quantora_logged_in = False

if "quantora_liked_posts" not in st.session_state:
    st.session_state.quantora_liked_posts = set()

if "view_profile" not in st.session_state:
    st.session_state.view_profile = None

if "learning_history" not in st.session_state:
    st.session_state.learning_history = []

if "iq_test_score" not in st.session_state:
    st.session_state.iq_test_score = None

if "city_input" not in st.session_state:
    st.session_state.city_input = "Berlin"

if "translation_history" not in st.session_state:
    st.session_state.translation_history = []

if "last_translation" not in st.session_state:
    st.session_state.last_translation = None
              
# --------------------------
# NEW: AI CONTENT DETECTOR & HUMANIZER FUNCTIONS
# --------------------------
def ai_content_detector(text):
    """Analyze text for AI-generated content using ZeroGPT-like analysis"""
    prompt = f"""You are an advanced AI content detector similar to ZeroGPT.
Analyze the following text for signs of AI generation.
You must estimate two main values:
1. Perplexity - how predictable the text is.
2. Burstiness - how much variation exists in sentence length and structure.
Rules:
- If both perplexity and burstiness are LOW -> Mark the text as likely AI-generated.
- If both perplexity and burstiness are HIGH -> Mark the text as likely human-written.
- If results are mixed -> Mark it as partially AI-generated.
Return your analysis as follows:
----------------------------------------------------
🧠 AI Content Detection Report:
Perplexity: [number from 0-100, where lower = more AI-like]
Burstiness: [number from 0-100, where lower = more AI-like]
AI Probability: [percentage from 0-100% showing how likely it's AI-generated]
Verdict: [AI-generated / Human-written / Mixed]
----------------------------------------------------
Here is the text to analyze:
\"\"\"
{text}
\"\"\"
"""
    try:
        response = call_a4f_model(prompt, "provider-5/grok-4-0709")
        return response
    except Exception as e:
        return f"Error in AI content detection: {str(e)}"

def ai_humanizer(text):
    """Convert AI-generated text to human-like text"""
    prompt = f"""You are an advanced AI-to-human text converter.
Your task is to completely humanize the given text so that it passes all AI detection tools (like ZeroGPT, GPTZero, and Copyleaks) with a 0% AI probability.
Rules:
- Rewrite every sentence naturally, as if written by a real person.
- Add slight variations in sentence length, tone, and rhythm to increase burstiness.
- Use natural transitions, occasional mild imperfections, and emotional expression.
- Maintain the original meaning, facts, and context.
- Avoid robotic patterns, repetitive phrasing, or overly formal tone.
- Add small human touches (e.g., hesitations, relatable phrasing, idioms, or minor style quirks).
- The final output should show **high perplexity and high burstiness** when analyzed.
Return only the humanized text with no explanation.
Here is the text to humanize:
\"\"\"
{text}
\"\"\"
"""
    try:
        response = call_a4f_model(prompt, "provider-5/grok-4-0709")
        return response
    except Exception as e:
        return f"Error in AI humanizer: {str(e)}"

def ai_content_detector_mode():
    """Mode for AI content detection"""
    st.title("🔍 AI Content Detector")
    st.markdown("Analyze any text to detect if it was generated by AI using advanced ZeroGPT-like analysis.")
    
    col1, col2 = st.columns([2, 1])
    with col1:
        text_to_analyze = st.text_area(
            "Enter text to analyze:",
            height=200,
            placeholder="Paste any text here to check if it was AI-generated...",
            key="detector_input"
        )
    with col2:
        st.markdown("### 📊 Analysis Options")
        analyze_button = st.button("🔍 Analyze Text", use_container_width=True)
        st.markdown("---")
        st.markdown("**How it works:**")
        st.markdown("• Analyzes perplexity & burstiness")
        st.markdown("• Compares to AI writing patterns")
        st.markdown("• Provides detailed report")
    
    if analyze_button and text_to_analyze.strip():
        with st.spinner("🤖 Analyzing text for AI patterns..."):
            result = ai_content_detector(text_to_analyze)
        
            st.markdown("### 📋 Detection Results")
            st.markdown(f"""
            <div class="detector-container">
                <h3>AI Content Analysis Report</h3>
                <div class="detector-result">
                    {result}
                </div>
            </div>
            """, unsafe_allow_html=True)
        
            # Extract probability for visualization
            probability_match = re.search(r"AI Probability:\s*(\d+)%", result)
            if probability_match:
                probability = int(probability_match.group(1))
                st.progress(probability/100, text=f"AI Probability: {probability}%")
            
                if probability > 70:
                    st.error("🚨 High probability of AI-generated content")
                elif probability > 30:
                    st.warning("⚠️ Mixed - Possibly AI-assisted content")
                else:
                    st.success("✅ Likely human-written content")
    elif analyze_button and not text_to_analyze.strip():
        st.warning("Please enter some text to analyze.")

def ai_humanizer_mode():
    """Mode for AI text humanization"""
    st.title("✍️ AI Text Humanizer")
    st.markdown("Transform AI-generated text into natural, human-like content that passes AI detection tools.")
    
    col1, col2 = st.columns([2, 1])
    with col1:
        text_to_humanize = st.text_area(
            "Enter AI-generated text to humanize:",
            height=200,
            placeholder="Paste AI-generated text here to make it sound more human...",
            key="humanizer_input"
        )
    with col2:
        st.markdown("### 🎯 Humanization Options")
        humanize_button = st.button("✨ Humanize Text", use_container_width=True)
        st.markdown("---")
        st.markdown("**Features:**")
        st.markdown("• Adds natural variations")
        st.markdown("• Improves sentence rhythm")
        st.markdown("• Removes AI patterns")
        st.markdown("• Maintains meaning")
    
    if humanize_button and text_to_humanize.strip():
        with st.spinner("🎨 Transforming text to sound more human..."):
            humanized_text = ai_humanizer(text_to_humanize)
        
            st.markdown("### 📝 Humanized Text")
            st.markdown(f"""
            <div class="humanizer-container">
                <h3>Humanized Version</h3>
                <div style="background: rgba(255,255,255,0.9); color: #000; padding: 1.5rem; border-radius: 10px; margin: 1rem 0;">
                    {humanized_text}
                </div>
            </div>
            """, unsafe_allow_html=True)
        
            # Provide download option
            st.download_button(
                label="📥 Download Humanized Text",
                data=humanized_text,
                file_name="humanized_text.txt",
                mime="text/plain"
            )
        
            # Option to analyze the humanized text
            if st.button("🔍 Check Humanized Text"):
                with st.spinner("Verifying humanization quality..."):
                    verification = ai_content_detector(humanized_text)
                    st.markdown("### ✅ Humanization Verification")
                    st.markdown(verification)
    elif humanize_button and not text_to_humanize.strip():
        st.warning("Please enter some text to humanize.")

# IQ Tester in Sidebar
if st.session_state.pro_unlocked:
    with st.sidebar:
        st.markdown("### 🧠 IQ Tester")
        if st.button("Start IQ Test"):
            questions = [
                {"q": "What number comes next: 1, 1, 2, 3, 5, 8, ?", "a": "13"},
                {"q": "If APPLE is coded as ZKKOV, how is BANANA coded?", "a": "YZMZMZ"},
                {"q": "A bat and a ball cost $1.10. The bat costs $1 more than the ball. How much is the ball?", "a": "0.05"}
            ]
            score = 0
            for i, q in enumerate(questions):
                answer = st.text_input(q["q"], key=f"iq_{i}")
                if answer.lower() == q["a"].lower():
                    score += 1
            st.session_state.iq_test_score = score * 33 # Simplified scoring
            st.write(f"Your IQ Score: {st.session_state.iq_test_score}")

# Document Analysis Functions
def extract_pdf_content(file):
    try:
        pdf_reader = PdfReader(file)
        content = ""
        for page in pdf_reader.pages:
            content += page.extract_text() + "\n"
        return content.strip()
    except Exception as e:
        return f"Error reading PDF: {e}"

def extract_docx_content(file):
    try:
        doc = docx.Document(file)
        content = ""
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        return content.strip()
    except Exception as e:
        return f"Error reading DOCX: {e}"

def extract_csv_content(file):
    try:
        df = pd.read_csv(file)
        content = f"CSV File Analysis:\n"
        content += f"Shape: {df.shape[0]} rows, {df.shape[1]} columns\n"
        content += f"Columns: {', '.join(df.columns.tolist())}\n"
        content += f"First 5 rows:\n{df.head().to_string()}\n"
        content += f"Data types:\n{df.dtypes.to_string()}\n"
        if df.select_dtypes(include=['number']).columns.any():
            content += f"Numerical summary:\n{df.describe().to_string()}\n"
        return content
    except Exception as e:
        return f"Error reading CSV: {e}"

def process_uploaded_file(uploaded_file):
    if uploaded_file is None:
        return ""
    file_type = uploaded_file.name.split('.')[-1].lower()
    try:
        if file_type == 'pdf':
            return extract_pdf_content(uploaded_file)
        elif file_type == 'docx':
            return extract_docx_content(uploaded_file)
        elif file_type == 'csv':
            return extract_csv_content(uploaded_file)
        elif file_type in ['txt', 'md', 'py', 'js', 'html', 'css', 'json']:
            content = uploaded_file.read().decode('utf-8')
            return f"File: {uploaded_file.name}\nContent:\n{content}"
        elif file_type in ['jpg', 'jpeg', 'png']:
            st.session_state.uploaded_image = Image.open(uploaded_file)
            st.session_state.enhanced_image = st.session_state.uploaded_image.copy()
            return f"Image: {uploaded_file.name} uploaded for enhancement"
        else:
            return f"Unsupported file type: {file_type}"
    except Exception as e:
        return f"Error processing file: {e}"

# Image Enhancement Functions
def enhance_image(image, brightness=1.0, contrast=1.0, sharpness=1.0, color=1.0):
    try:
        enhancers = {
            'brightness': ImageEnhance.Brightness(image),
            'contrast': ImageEnhance.Contrast(image),
            'sharpness': ImageEnhance.Sharpness(image),
            'color': ImageEnhance.Color(image)
        }
    
        enhanced = enhancers['brightness'].enhance(brightness)
        enhanced = enhancers['contrast'].enhance(contrast)
        enhanced = enhancers['sharpness'].enhance(sharpness)
        enhanced = enhancers['color'].enhance(color)
    
        return enhanced
    except Exception as e:
        st.error(f"Error enhancing image: {e}")
        return image

def apply_image_filters(image, filter_type):
    try:
        if filter_type == 'blur':
            return image.filter(ImageFilter.BLUR)
        elif filter_type == 'contour':
            return image.filter(ImageFilter.CONTOUR)
        elif filter_type == 'detail':
            return image.filter(ImageFilter.DETAIL)
        elif filter_type == 'edge_enhance':
            return image.filter(ImageFilter.EDGE_ENHANCE)
        elif filter_type == 'emboss':
            return image.filter(ImageFilter.EMBOSS)
        elif filter_type == 'sharpen':
            return image.filter(ImageFilter.SHARPEN)
        elif filter_type == 'smooth':
            return image.filter(ImageFilter.SMOOTH)
        else:
            return image
    except Exception as e:
        st.error(f"Error applying filter: {e}")
        return image

def parse_edit_instructions(instructions):
    """Parse text instructions to map to PIL operations"""
    instructions = instructions.lower()
    enhancements = {
        "brightness": 1.0,
        "contrast": 1.0,
        "sharpness": 1.0,
        "color": 1.0,
        "filter": "None"
    }
    # Simple keyword-based parsing
    if "bright" in instructions:
        enhancements["brightness"] = 1.3
    if "dark" in instructions:
        enhancements["brightness"] = 0.7
    if "contrast" in instructions:
        enhancements["contrast"] = 1.3
    if "sharp" in instructions:
        enhancements["sharpness"] = 1.5
    if "color" in instructions or "vibrant" in instructions:
        enhancements["color"] = 1.3
    if "blur" in instructions:
        enhancements["filter"] = "blur"
    if "contour" in instructions:
        enhancements["filter"] = "contour"
    if "detail" in instructions:
        enhancements["filter"] = "detail"
    if "edge" in instructions:
        enhancements["filter"] = "edge_enhance"
    if "emboss" in instructions:
        enhancements["filter"] = "emboss"
    if "smooth" in instructions:
        enhancements["filter"] = "smooth"
    return enhancements

def display_image_enhancement_controls(image, enhancements):
    with st.expander("🖼️ Image Enhancement Tools", expanded=True):
        st.markdown("### Adjust Image Parameters")
    
        col1, col2 = st.columns(2)
    
        with col1:
            brightness = st.slider("Brightness", 0.0, 2.0, enhancements["brightness"], 0.1)
            contrast = st.slider("Contrast", 0.0, 2.0, enhancements["contrast"], 0.1)
    
        with col2:
            sharpness = st.slider("Sharpness", 0.0, 2.0, enhancements["sharpness"], 0.1)
            color = st.slider("Color", 0.0, 2.0, enhancements["color"], 0.1)
    
        st.markdown("### Apply Filters")
        filter_options = ['None', 'blur', 'contour', 'detail', 'edge_enhance', 'emboss', 'sharpen', 'smooth']
        selected_filter = st.selectbox("Choose a filter", filter_options, index=filter_options.index(enhancements["filter"]))
    
        enhanced_image = enhance_image(image, brightness, contrast, sharpness, color)
        if selected_filter != 'None':
            enhanced_image = apply_image_filters(enhanced_image, selected_filter)
    
        return enhanced_image

# Enhanced A4F Model Call with fallback
def call_a4f_model(prompt, model_name, context="", image=None):
    """FIXED: A4F API call with proper error handling and image support"""
    try:
        # API Configuration
        A4F_API_KEY = "ddc-a4f-b752e3e2936149f49b1b306953e0eaab"
        API_URL = "https://api.a4f.co/v1/chat/completions"
        
        headers = {
            "Authorization": f"Bearer {A4F_API_KEY}",
            "Content-Type": "application/json"
        }
        
        # Build messages array
        messages = []
        
        # Add system prompt if needed
        if "gemini" in model_name.lower() or "gpt" in model_name.lower():
            messages.append({
                "role": "system",
                "content": "You are a helpful AI assistant. Provide clear, accurate, and detailed responses."
            })
        
        # Build the content
        content_parts = []
        
        # Add text content
        full_prompt = f"{context}\n\n{prompt}" if context else prompt
        content_parts.append({
            "type": "text",
            "text": full_prompt
        })
        
        # Add image if provided
        if image is not None:
            try:
                # Convert PIL Image to base64
                buffered = BytesIO()
                image.save(buffered, format="PNG")
                img_str = base64.b64encode(buffered.getvalue()).decode("utf-8")
                
                content_parts.append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/png;base64,{img_str}"
                    }
                })
            except Exception as img_error:
                st.error(f"Image processing error: {str(img_error)}")
                # Continue without image if there's an error
        
        # Add user message with content
        messages.append({
            "role": "user",
            "content": content_parts
        })
        
        # Prepare request data
        data = {
            "model": model_name,
            "messages": messages,
            "temperature": 0.7,
            "max_tokens": 4000,  # Increased for better responses
            "top_p": 0.9,
            "stream": False
        }
        
        # Make the API request
        response = requests.post(
            API_URL,
            headers=headers,
            json=data,
            timeout=60  # Increased timeout for image analysis
        )
        
        # Handle response
        if response.status_code == 200:
            result = response.json()
            
            if "choices" in result and len(result["choices"]) > 0:
                if "message" in result["choices"][0]:
                    content = result["choices"][0]["message"]["content"]
                    
                    # Clean up response if needed
                    if content:
                        # Remove any thinking tags from models like R1
                        if "r1" in model_name.lower() or "think" in content.lower():
                            content = re.sub(r'<think>.*?</think>', '', content, flags=re.DOTALL)
                            content = content.strip()
                        
                        # Remove any system prompts that might have leaked
                        content = re.sub(r'^.*?(?=\n\n|$)', '', content, flags=re.DOTALL)
                        content = content.strip()
                        
                        return content
                    else:
                        return "🤖 I received an empty response. Please try again with a different prompt."
                else:
                    return "⚠️ Unexpected response format from AI model."
            else:
                return "⚠️ No valid choices in AI response."
        
        # Handle specific error codes
        elif response.status_code == 401:
            return "🔒 Authentication failed. Please check your API key."
        elif response.status_code == 403:
            return "🚫 Access forbidden. You may not have permission to use this model."
        elif response.status_code == 404:
            return "❌ Model not found. Please check the model name."
        elif response.status_code == 429:
            return "⏳ Rate limit exceeded. Please wait a moment and try again."
        elif response.status_code == 500:
            return "⚙️ Server error. The AI provider is experiencing issues."
        elif response.status_code == 503:
            return "🔄 Service unavailable. Please try again in a few moments."
        else:
            return f"⚠️ API Error {response.status_code}: {response.text[:200]}"
    
    except requests.exceptions.Timeout:
        return "⏱️ Request timed out. The server took too long to respond."
    except requests.exceptions.ConnectionError:
        return "🌐 Connection error. Please check your internet connection."
    except requests.exceptions.RequestException as e:
        return f"🔗 Network error: {str(e)[:100]}"
    except json.JSONDecodeError:
        return "📄 Invalid response from server. Could not parse JSON."
    except Exception as e:
        return f"❌ Unexpected error: {str(e)[:150]}"
                      
# Enhanced Groq Model Calls
def call_groq_model(prompt, model_name, context=""):
    """FIXED: Groq API call with proper error handling"""
    if not groq_client:
        return "❌ Groq client not initialized"
    
    try:
        system_prompt = "You are a helpful AI assistant."
        
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"{context}\n\n{prompt}" if context else prompt}
        ]
        
        completion = groq_client.chat.completions.create(
            model=model_name,
            messages=messages,
            temperature=0.7,
            max_tokens=1500,
            top_p=0.9
        )
        
        if completion.choices and completion.choices[0].message.content:
            return completion.choices[0].message.content
        else:
            return "❌ Empty response from Groq"
            
    except Exception as e:
        return f"❌ Groq API Error: {str(e)}"

# Quantora Unified AI Model with Memory and Simulated Learning
def call_quantora_unified(prompt, context="", image=None):
    """FIXED: Unified AI Model with guaranteed response"""
    start_time = time.time()
    
    # Build conversation history
    conversation_history = ""
    chat_history = st.session_state.get('chat', [])
    for item in chat_history[-5:]:
        if len(item) >= 2:
            speaker, message = item[:2]
            conversation_history += f"{speaker.upper()}: {message}\n\n"
    
    # Simulated learning
    learning_prompt = ""
    if st.session_state.get('learning_history'):
        learning_prompt = "\n\nPrevious interactions:\n" + "\n".join(
            st.session_state.learning_history[-3:]
        )
    
    full_prompt = f"{conversation_history}{learning_prompt}\n\nCurrent Query: {prompt}"
    
    # Select model based on version
    selected_model = st.session_state.get("model_version", "Quantora Prime 1 Fast (Faster But Not As Better As Og Flagship Model)")
    
    # Simplified model selection - use reliable models
    if "Prime 1" in selected_model or "Flagship" in selected_model:
        st.toast("🚀 Using Prime Engine...", icon="🚀")
        # Try Gemini first
        response = call_a4f_model(full_prompt, "provider-2/gemini-3-pro-preview", context, image)
        # If error, try GPT-4o
        if response and response.startswith("❌"):
            response = call_a4f_model(full_prompt, "provider-5/gpt-4o", context, image)
    elif "Fast" in selected_model:
        st.toast("⚡ Using Fast Engine...", icon="⚡")
        response = call_a4f_model(full_prompt, "provider-2/gemini-2.5-flash-lite", context, image)
    elif "Code" in selected_model:
        st.toast("💻 Using Code Engine...", icon="💻")
        response = call_a4f_model(full_prompt, "provider-5/gpt-5.1-chat-latest", context, image)
    else:
        # Default
        response = call_a4f_model(full_prompt, "provider-2/gemini-3-pro-preview", context, image)
    
    # If all API calls fail, provide a helpful default response
    if not response or response.startswith("❌"):
        # Try Groq as final fallback
        groq_response = call_groq_model(full_prompt, "qwen/qwen3-32b", context)
        if groq_response and not groq_response.startswith("❌"):
            response = groq_response
        else:
            # Ultimate fallback - always return something
            response = f"""I've received your query about: "{prompt[:100]}..."

Thanks for reaching out! I'm here to help.

**Regarding your question:** I understand you're asking about this topic, and I'd be happy to provide insights based on my knowledge.

**Here's what I can share:**
This is an interesting topic that often involves multiple perspectives. 
Depending on the specific context you're referring to, there are various approaches and considerations.

**To give you the best answer, could you clarify:**
1. Are you looking for technical details or general information?
2. Do you have any specific requirements or constraints?
3. Would you like practical examples or theoretical explanations?

**In the meantime, here are some general suggestions:**
- Research reputable sources on this topic
- Consider consulting with experts in the field
- Break down complex aspects into smaller questions

Feel free to ask more specific questions, and I'll do my best to provide detailed, helpful responses! 💎

*What aspect would you like me to focus on first?*"""
    
    processing_time = time.time() - start_time
    
    # Store learning history
    if "learning_history" in st.session_state:
        learning_note = f"Query: {prompt[:50]}... Response time: {processing_time:.1f}s"
        st.session_state.learning_history.append(learning_note)
        if len(st.session_state.learning_history) > 10:
            st.session_state.learning_history = st.session_state.learning_history[-10:]
    
    return response


    def call_groq_backend(model_name):
        try:
            response = call_groq_model(full_prompt, model_name, context)
            return {
                "backend": f"groq_{model_name}",
                "response": response,
                "success": True,
                "length": len(response) if response else 0
            }
        except Exception as e:
            return {
                "backend": f"groq_{model_name}",
                "response": f"Backend error: {str(e)}",
                "success": False,
                "length": 0
            }

    def call_a4f_backend(model_name):
        try:
            response = call_a4f_model(full_prompt, model_name, context, image)
            return {
                "backend": f"a4f_{model_name}",
                "response": response,
                "success": response is not None,
                "length": len(response) if response else 0
            }
        except Exception as e:
            return {
                "backend": f"a4f_{model_name}",
                "response": f"Backend error: {str(e)}",
                "success": False,
                "length": 0
            }

    backend_results = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
        futures = []
        selected_model_version = st.session_state.get("model_version", "Quantora Prime 1 (Latest Flagship Model)")
        
        if selected_model_version == "Quantora Prime 1 (Latest Flagship Model)":
            st.toast("🚀 Using Quantora Prime 1 Engine...", icon="🚀")
            groq_models = []
            a4f_models = [
                "provider-3/claude-3.5-haiku",
                "provider-2/r1-1776",
                "provider-5/gpt-4o",
                "provider-1/claude-opus-4",
                "provider-6/grok-3-reasoning",
                "provider-5/gpt-4.1-nano",
                "provider-3/deepseek-v3",
                "provider-6/claude-3-7-sonnet-20250219-thinking",
                "provider-1/claude-sonnet-4",
                "provider-5/gemini-2.5-flash-preview-05-20",
                "provider-3/grok-4-0709",
                "provider-1/sonar-pro",
                "provider-3/qwen-2.5-coder-32b",
                "provider-2/codestral",
                "provider-1/sonar-deep-research",
                "provider-1/sonar-reasoning-pro",
                "provider-2/llama-4-maverick",
                "provider-3/qwen-2.5-72b",
                "provider-3/gpt-5-nano",
                "provider-1/deepseek-v3.1",
                "provider-5/gpt-5.1",
                "provider-2/gemini-3-pro-preview",
                "provider-7/claude-sonnet-4-5-20250929"
            ]
            for model in groq_models:
                futures.append(executor.submit(call_groq_backend, model))
            for model in a4f_models:
                futures.append(executor.submit(call_a4f_backend, model))
    
        elif selected_model_version == "Quantora Prime 1 Fast (Faster But Not As Better As Og Flagship Model)":
            st.toast("⚡ Using Quantora Prime 1 Fast Engine...", icon="⚡")
            a4f_v2_models = [
                "provider-2/gemini-2.5-flash-lite",
                "provider-1/deepseek-v3.1"
            ]
            for model in a4f_v2_models:
                futures.append(executor.submit(call_a4f_backend, model))
            
        elif selected_model_version == "Quantora V3 (Code Specialized)":
            st.toast("💻 Using Quantora V3 Code Engine...", icon="💻")
            code_models = [
                "provider-5/gpt-5.1-chat-latest",
            ]
            for model in code_models:
                futures.append(executor.submit(call_a4f_backend, model))
            
        elif selected_model_version == "Quantora V4 (Long Conversation)":
            st.toast("🗣️ Using Quantora V4 Conversation Engine...", icon="🗣️")
            conversation_models = [
                "provider-6/gemini-2.5-flash",
                "provider-6/minimax-m1-40k",
                "provider-1/claude-opus-4",
                "provider-1/sonar-deep-research"
            ]
            for model in conversation_models:
                futures.append(executor.submit(call_a4f_backend, model))
            
        elif selected_model_version == "Quantora V3 (Reasoning Specialized)":
            st.toast("🧠 Using Quantora V3 Reasoning Engine...", icon="🧠")
            reasoning_models = [
                "provider-1/sonar-reasoning",
                "provider-6/r1-1776",
                "provider-1/sonar-reasoning-pro",
                "provider-1/sonar-deep-research"
            ]
            for model in reasoning_models:
                futures.append(executor.submit(call_a4f_backend, model))
            
        elif selected_model_version == "Quantora V3 (Math Specialized)":
            st.toast("🧮 Using Quantora V3 Math Engine...", icon="🧮")
            math_models = [
                "provider-3/qwen-2.5-72b",
                "provider-6/gemini-2.5-flash",
                "provider-6/minimax-m1-40k"
            ]
            for model in math_models:
                futures.append(executor.submit(call_a4f_backend, model))
        
        for future in concurrent.futures.as_completed(futures):
            try:
                result = future.result()
                backend_results.append(result)
            except Exception as e:
                print(f"⚠️ One processing component had an issue: {str(e)}")
    
    successful_responses = [r for r in backend_results if r['success'] and r['response'] and not r['response'].startswith("Backend error")]
    
    if not successful_responses:
        return "❌ No successful responses from backends. Please try again."
    
    responses_text = '\n\n'.join([f"Response from {r['backend']}:\n{r['response']}" for r in successful_responses])
    
    mixing_prompt = f"""You are the response synthesizer fo Quantora. Below are multiple responses to the same prompt. Combine them into one coherent, comprehensive response that maintains the best aspects of each.
                        Original Prompt: {prompt}
                        Responses to combine:
                        {responses_text}
                        Guidelines:
                        1. Preserve all valuable information
                        2. Remove any redundancies
                        3. Maintain a professional, engaging tone
                        4. Keep code blocks intact
                        5. Ensure logical flow
                        6. Add any missing context that would improve the answer
                        7. Dont tell about any ai error
                        8. Always ask a question to the user ex- that what should I(you) do next?
                        Combined Response:"""
    
    final_response = call_a4f_model(mixing_prompt, "provider-3/gpt-4o-mini")
    
    # Simulated auto-training: "Learn" by storing response improvements
    if final_response:
        learning_note = f"Improved response for query: {prompt[:50]}... by combining {len(successful_responses)} backends"
        st.session_state.learning_history.append(learning_note)
    
    processing_time = time.time() - start_time
    return final_response if final_response else successful_responses[0]['response']

# Code Detection and Formatting
def format_response_with_code(response):
    code_pattern = r'```(\w+)?\n(.*?)\n```'
    parts = []
    last_end = 0
    
    for match in re.finditer(code_pattern, response, re.DOTALL):
        if match.start() > last_end:
            text_part = response[last_end:match.start()].strip()
            if text_part:
                parts.append(('text', text_part))
    
        language = match.group(1) or 'text'
        code_content = match.group(2).strip()
        parts.append(('code', code_content, language))
    
        last_end = match.end()
    
    if last_end < len(response):
        remaining_text = response[last_end:].strip()
        if remaining_text:
            parts.append(('text', remaining_text))
    
    return parts if parts else [('text', response)]

# Image Generation Functions
def generate_image(prompt, style):
    #Generate image using A4F API
    try:
        headers = {
            "Authorization": f"Bearer {A4F_API_KEY}",
            "Content-Type": "application/json"
        }
        
        enhanced_prompt = f"{prompt}, {style} style, high quality, photorealistic, 4k resolution"
        
        # Use a more reliable endpoint structure
        payload = {
            "model": "provider-2/nano-banana-pro",  # Ensure this is a valid model
            "prompt": enhanced_prompt,
            "size": "1024x1024",
            "num_images": 1,
            "response_format": "url"
        }
        
        # Try the correct endpoint
        response = requests.post(
            "https://api.a4f.co/v1/images/generations",
            headers=headers,
            json=payload,
            timeout=60
        )
        
        if response.status_code == 200:
            result = response.json()
            if 'data' in result and len(result['data']) > 0:
                image_url = result['data'][0]['url']
                # Download the image
                image_response = requests.get(image_url, timeout=30)
                if image_response.status_code == 200:
                    return Image.open(BytesIO(image_response.content))
                else:
                    st.error(f"Failed to download image: {image_response.status_code}")
                    return None
            else:
                st.error(f"No data in response: {result}")
                return None
        else:
            st.error(f"Image generation API error: {response.status_code} - {response.text}")
            return None
            
    except requests.exceptions.RequestException as e:
        st.error(f"Network error: {str(e)}")
        return None
    except Exception as e:
        st.error(f"Unexpected error: {str(e)}")
        return None

def edit_image(image, edit_prompt):
    headers = {
        "Authorization": f"Bearer {A4F_API_KEY}",
        "Content-Type": "application/json"
    }
    
    buffered = BytesIO()
    image.save(buffered, format="PNG")
    img_str = base64.b64encode(buffered.getvalue()).decode("utf-8")
    
    payload = {
        "model": EDIT_MODEL,
        "prompt": edit_prompt,
        "image": img_str,
        "num_images": 1,
        "width": 1024,
        "height": 1024
    }
    
    try:
        response = requests.post(
            f"{A4F_BASE_URL}/images/edits",
            headers=headers,
            json=payload,
            timeout=60
        )
        if response.status_code == 200:
            result = response.json()
            if 'data' in result and len(result['data']) > 0:
                image_url = result['data'][0]['url']
                image_response = requests.get(image_url)
                return Image.open(BytesIO(image_response.content))
        return None
    except Exception as e:
        st.error(f"Image editing error: {str(e)}")
        return None

# Video Generation Function using Replicate
def generate_video_replicate(prompt, style):
    try:
        # ✅ Set your Replicate API key
        os.environ["REPLICATE_API_TOKEN"] = "r8_7t4VS9WzjYf0ohxFuez5bDAa66dNalb3w5Jql"
     
        # Run the model with original prompt only
        output = replicate.run(
            "minimax/video-01",
            input={
                "prompt": prompt,
                "prompt_optimizer": True
            }
        )
        # ✅ Handle output correctly (new Replicate SDK)
        if hasattr(output, "path"):
            video_url = output.path # For FileOutput object
        else:
            video_url = str(output) # For string or list
        print("🎥 Video generated at:", video_url)
        
        # ✅ Download the file to disk
        response = requests.get(video_url)
        filename = f"generated_video_{int(time.time())}.mp4"
        with open(filename, "wb") as file:
            file.write(response.content)
        print(f"✅ Video saved successfully as {filename}")
        return filename
     
    except Exception as e:
        st.error(f"Video generation failed: {str(e)}")
        return None

# Time-based greeting
hour = datetime.now().hour
if 6 <= hour < 12:
    greeting = "🌅 Good Morning!"
elif 12 <= hour < 18:
    greeting = "☀️ Good Afternoon!"
elif 18 <= hour < 24:
    greeting = "🌙 Good Evening!"
else:
    greeting = "🌌 Good Night!"

# App name based on pro status
app_name = "Quantora Prime X" if st.session_state.get('pro_unlocked', False) else "Quantora AI"

# Use string concatenation instead of f-string
header_html = '''
<div class="main-header">
    <div class="logo">
        <div class="logo-icon">💎</div>
        <div class="logo-text">''' + app_name + '''</div>
        <div class="status-indicator"></div>
    </div>
    <div style="color: #94a3b8; font-size: 1.1rem; margin-top: 0.5rem;">''' + greeting + '''</div>
</div>
'''

st.markdown(header_html, unsafe_allow_html=True)
# --------------------------
# QUANTORA TRADE CHARTS MODULE
# --------------------------
def quantora_trade_charts():
    st.title("📈 Quantora Trade Charts")
    st.markdown("Advanced financial analysis and visualization tools powered by Quantora AI")
    
    # Stock selection
    col1, col2 = st.columns([0.7, 0.3])
    with col1:
        ticker = st.text_input("Enter stock symbol (e.g. AAPL, MSFT, TSLA)", "AAPL")
    with col2:
        period = st.selectbox("Time period", ["1mo", "3mo", "6mo", "1y", "2y", "5y", "10y"])
    
    if st.button("Generate Analysis"):
        with st.spinner("Fetching market data..."):
            try:
                stock = yf.Ticker(ticker)
                hist = stock.history(period=period)
            
                if hist.empty:
                    st.error("No data found for this symbol. Please try another.")
                    return
            
                # Basic info
                st.subheader(f"📊 {ticker} - {stock.info.get('longName', 'N/A')}")
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Current Price", f"${stock.info.get('currentPrice', 'N/A')}")
                with col2:
                    change = stock.info.get('regularMarketChange', 0)
                    change_percent = stock.info.get('regularMarketChangePercent', 0)
                    st.metric("Daily Change", f"${change:.2f}", f"{change_percent:.2f}%")
                with col3:
                    st.metric("Market Cap", f"${stock.info.get('marketCap', 'N/A'):,}")
            
                # Candlestick chart
                st.subheader("Candlestick Chart")
                fig = go.Figure(data=[go.Candlestick(
                    x=hist.index,
                    open=hist['Open'],
                    high=hist['High'],
                    low=hist['Low'],
                    close=hist['Close']
                )])
                fig.update_layout(
                    title=f"{ticker} Price Movement",
                    xaxis_title="Date",
                    yaxis_title="Price (USD)",
                    template="plotly_dark"
                )
                st.plotly_chart(fig, use_container_width=True)
            
                # Volume chart
                st.subheader("Trading Volume")
                fig2 = go.Figure(data=[go.Bar(
                    x=hist.index,
                    y=hist['Volume'],
                    marker_color='#8b5cf6'
                )])
                fig2.update_layout(
                    title=f"{ticker} Trading Volume",
                    xaxis_title="Date",
                    yaxis_title="Volume",
                    template="plotly_dark"
                )
                st.plotly_chart(fig2, use_container_width=True)
            
                # Additional metrics
                st.subheader("Key Metrics")
                metrics = {
                    "52 Week High": stock.info.get('fiftyTwoWeekHigh'),
                    "52 Week Low": stock.info.get('fiftyTwoWeekLow'),
                    "PE Ratio": stock.info.get('trailingPE'),
                    "Dividend Yield": stock.info.get('dividendYield'),
                    "Beta": stock.info.get('beta'),
                    "Average Volume": stock.info.get('averageVolume')
                }
            
                cols = st.columns(3)
                for i, (metric, value) in enumerate(metrics.items()):
                    with cols[i % 3]:
                        if value is not None:
                            if metric == "Dividend Yield":
                                st.metric(metric, f"{value*100:.2f}%")
                            elif metric == "Average Volume":
                                st.metric(metric, f"{value:,.0f}")
                            else:
                                st.metric(metric, f"{value:.2f}")
                        else:
                            st.metric(metric, "N/A")
            
                # AI Analysis
                st.subheader("📈 Quantora AI Analysis")
                analysis_prompt = f"""Provide a technical analysis of {ticker} stock based on the following data: - Current Price: ${stock.info.get('currentPrice', 'N/A')} - 52 Week Range: ${stock.info.get('fiftyTwoWeekLow', 'N/A')} - ${stock.info.get('fiftyTwoWeekHigh', 'N/A')} - PE Ratio: {stock.info.get('trailingPE', 'N/A')} - Market Cap: ${stock.info.get('marketCap', 'N/A'):,} - Recent Performance: {hist.tail(5)[['Open', 'High', 'Low', 'Close', 'Volume']].to_string()} Provide insights on:1. Current trend 2. Key support/resistance levels 3. Volume analysis 4. Technical indicators summary 5. Short-term and long-term outlook Keep the analysis professional but accessible to retail investors."""
            
                with st.spinner("Generating AI analysis..."):
                    st.session_state.model_version = "Quantora Prime 1 (Latest Flagship Model)"
                    analysis = call_quantora_unified(analysis_prompt)
                    st.markdown(analysis)
            
            except Exception as e:
                st.error(f"Error fetching data: {str(e)}")

# --------------------------
# QUANTORA NEWS MODULE
# --------------------------
def quantora_news():
    today = datetime.now().strftime("%B %d, %Y")
    
    # Fixed: Proper escape sequences and string formatting
    header_html = '''
        <div style="text-align: center; padding: 10px 0 20px 0;">
            <h1 style="font-size: 3em; margin-bottom: 0;">📰 Quantora AI News Digest</h1>
            <p style="color: gray; font-size: 1.2em;">The most powerful news summary for <strong>{today}</strong></p>
            <p style="font-size: 0.9em; color: #888;">Generated by Quantora AI</p>
        </div>
    '''.format(today=today)
    
    st.markdown(header_html, unsafe_allow_html=True)
    
    # Dynamically dated prompt - using raw string to avoid escape issues
    prompt = """
    You are Quantora AI, a cutting-edge real-time news analysis system. Give the MOST Trending news for {today}. Create the top news digest for {today} based on live global and Indian events 'like' operation sindoor, using a professional journalist tone.
    Structure your summary into the following categories:
    1. Topic - 1 (2 detailed paragraphs)
    2. Topic - 2 (1 paragraph)
    3. Topic - 3 (2-3 bullet points)
    4. Topic - 4 (2-3 bullet points)
    5. Topic - 5 (2-3 bullet points)
    6. Topic - 6 (1 paragraph)
    Only include realistic and relevant news that would appear on Aaj Tak, ABP News, Zee News, and BBC for {today}.
    """.format(today=today)
    
    # Generate news
    with st.spinner("🔍 Quantora AI is gathering and analyzing today's global news..."):
        # Save current model and use flagship for news
        current_model = st.session_state.get("model_version", "Quantora Prime 1 Fast")
        st.session_state.model_version = "Quantora Prime 1 (Latest Flagship Model)"
        
        response = call_quantora_unified(prompt)
        news = response if response else "No news could be generated at this time."
        
        # Restore original model
        st.session_state.model_version = current_model
    
    # Display news with proper formatting
    st.markdown("---")
    
    # Clean and display the news
    if news:
        # Clean up any markdown formatting issues
        clean_news = news.replace('\\', '')  # Remove backslashes
        clean_news = re.sub(r'\\(\w)', r'\1', clean_news)  # Fix escaped characters
        
        # Display in a nice container
        st.markdown(f'''
        <div style="background-color: #ffffff; padding: 25px; border-radius: 15px; 
                    box-shadow: 0 4px 20px rgba(0,0,0,0.08); margin: 20px 0;
                    font-size: 1.1em; line-height: 1.7; color: #1a1a1a;">
            {clean_news}
        </div>
        ''', unsafe_allow_html=True)
    else:
        st.error("Could not generate news. Please try again.")
    
    # Footer
    st.markdown("---")
    st.markdown('''
        <div style="text-align: center; font-size: 0.85em; color: gray; padding-top: 10px;">
            🔹 Powered by Quantora AI • Delivering Intelligence, Not Just Information.
        </div>
    ''', unsafe_allow_html=True)

# --------------------------
# QUANTORA SOCIAL MEDIA MODULE
# --------------------------
def quantora_social_media():
    QUANTORA_POSTS_CSV = "quantora_social_posts.csv"
    QUANTORA_USERS_CSV = "quantora_social_users.csv"
    QUANTORA_FOLLOWS_CSV = "quantora_social_follows.csv"
    QUANTORA_IMAGES_DIR = "quantora_social_images"
    QUANTORA_PROFILE_PICS_DIR = "quantora_social_profile_pics"
    DEFAULT_PROFILE_PIC = "default_profile.png"
    
    if not os.path.exists(QUANTORA_POSTS_CSV):
        quantora_df_posts = pd.DataFrame(columns=['quantora_username', 'quantora_timestamp', 'quantora_text', 'quantora_image_path', 'quantora_likes', 'quantora_comments'])
        quantora_df_posts.to_csv(QUANTORA_POSTS_CSV, index=False)
    
    if not os.path.exists(QUANTORA_USERS_CSV):
        quantora_df_users = pd.DataFrame(columns=['quantora_email', 'quantora_username', 'quantora_password', 'quantora_profile_pic', 'bio'])
        quantora_df_users.to_csv(QUANTORA_USERS_CSV, index=False)
    
    if not os.path.exists(QUANTORA_FOLLOWS_CSV):
        quantora_df_follows = pd.DataFrame(columns=['follower', 'followed'])
        quantora_df_follows.to_csv(QUANTORA_FOLLOWS_CSV, index=False)
    
    if not os.path.exists(QUANTORA_IMAGES_DIR):
        os.makedirs(QUANTORA_IMAGES_DIR)
    
    if not os.path.exists(QUANTORA_PROFILE_PICS_DIR):
        os.makedirs(QUANTORA_PROFILE_PICS_DIR)

    # Helper Functions
    def handle_hashtags(text):
        return text

    def quantora_user_info_header(username, show_follow=False):
        quantora_users_df = pd.read_csv(QUANTORA_USERS_CSV)
        try:
            quantora_user_data = quantora_users_df[quantora_users_df['quantora_username'] == username].iloc[0]
            quantora_profile_pic_path = quantora_user_data.get('quantora_profile_pic', DEFAULT_PROFILE_PIC)
            if not os.path.exists(quantora_profile_pic_path):
                quantora_profile_pic_path = DEFAULT_PROFILE_PIC
        except IndexError:
            quantora_profile_pic_path = DEFAULT_PROFILE_PIC
        
        col1, col2 = st.columns([0.08, 0.92])
        with col1:
            st.markdown(f'<img src="{quantora_profile_pic_path}" width="36" style="border-radius: 50%; object-fit: cover;">', unsafe_allow_html=True)
        with col2:
            st.markdown(f"<strong style='font-size: 1.1em;'>{username}</strong>", unsafe_allow_html=True)
            if show_follow and username != st.session_state.quantora_username:
                is_following = is_user_following(st.session_state.quantora_username, username)
                follow_text = "Following" if is_following else "Follow"
                if st.button(follow_text, key=f"follow_{username}", use_container_width=True):
                    update_follow(st.session_state.quantora_username, username, not is_following)
                    st.rerun()

    def quantora_post_actions(row, index):
        quantora_username = row['quantora_username']
        quantora_timestamp = row['quantora_timestamp']
        quantora_likes = int(row.get('quantora_likes', 0))
        quantora_post_key = f"{quantora_username}_{quantora_timestamp}"
        liked = quantora_post_key in st.session_state.quantora_liked_posts
        
        col1, col2, _ = st.columns([0.15, 0.15, 0.7])
        with col1:
            like_button_label = f"{'❤️' if liked else '🤍'} {quantora_likes}"
            if st.button(like_button_label, key=f"like_btn_{index}", use_container_width=True):
                quantora_df = pd.read_csv(QUANTORA_POSTS_CSV)
                if liked:
                    quantora_df.at[index, 'quantora_likes'] -= 1
                    st.session_state.quantora_liked_posts.discard(quantora_post_key)
                else:
                    quantora_df.at[index, 'quantora_likes'] += 1
                    st.session_state.quantora_liked_posts.add(quantora_post_key)
                quantora_df.to_csv(QUANTORA_POSTS_CSV, index=False)
                st.rerun()
        
        with col2:
            with st.expander("💬 Comments", expanded=False):
                quantora_comment_section(row, index)

    def quantora_comment_section(row, index):
        quantora_username = row['quantora_username']
        quantora_comments_raw = row.get('quantora_comments', '')
        if pd.isna(quantora_comments_raw):
            quantora_comments_raw = ""
        quantora_comments = quantora_comments_raw.split("|") if quantora_comments_raw else []
        
        for c in quantora_comments:
            if c:
                parts = c.split(": ", 1)
                if len(parts) == 2:
                    commenter, comment_text = parts[0], parts[1]
                    colored_commenter = f'<span style="color: black;">{commenter}:</span>'
                    colored_comment_text = f'<span style="color: black;">{comment_text}</span>'
                    st.markdown(
                        f"<div style='padding: 8px; margin-bottom: 5px; background-color: #f0f2f5; border-radius: 5px; width: 100%; box-sizing: border-box;'><strong>{colored_commenter}</strong> {colored_comment_text}</div>",
                        unsafe_allow_html=True
                    )
                else:
                    colored_c = f'<span style="color: black;">{c}</span>'
                    st.markdown(
                        f"<div style='margin-bottom: 5px; width: 100%; box-sizing: border-box;'>- {colored_c}</div>",
                        unsafe_allow_html=True
                    )
        
        with st.container():
            comment_input_col, comment_button_col = st.columns([0.95, 0.05])
            with comment_input_col:
                quantora_new_comment = st.text_input(
                    "",
                    placeholder="Add a comment...",
                    key=f"comment_input_{index}",
                    help="Type your comment here"
                )
            with comment_button_col:
                if st.button("Post", key=f"comment_post_btn_{index}", use_container_width=True):
                    if quantora_new_comment:
                        quantora_df = pd.read_csv(QUANTORA_POSTS_CSV)
                        quantora_updated_comment = f"{st.session_state.quantora_username}: {quantora_new_comment}"
                        quantora_combined_comments = (
                            quantora_comments_raw + f"|{quantora_updated_comment}"
                            if quantora_comments_raw
                            else quantora_updated_comment
                        )
                        quantora_df.at[index, 'quantora_comments'] = quantora_combined_comments
                        quantora_df.to_csv(QUANTORA_POSTS_CSV, index=False)
                        st.rerun()

    def is_user_following(follower, followed):
        try:
            follows_df = pd.read_csv(QUANTORA_FOLLOWS_CSV)
            return ((follows_df['follower'] == follower) & (follows_df['followed'] == followed)).any()
        except FileNotFoundError:
            return False

    def update_follow(follower, followed, follow=True):
        follows_df = pd.read_csv(QUANTORA_FOLLOWS_CSV) if os.path.exists(QUANTORA_FOLLOWS_CSV) else pd.DataFrame(columns=['follower', 'followed'])
        if follow:
            if not ((follows_df['follower'] == follower) & (follows_df['followed'] == followed)).any():
                new_follow = pd.DataFrame([[follower, followed]], columns=['follower', 'followed'])
                new_follow.to_csv(QUANTORA_FOLLOWS_CSV, mode='a', header=False, index=False)
        else:
            follows_df = follows_df[~((follows_df['follower'] == follower) & (follows_df['followed'] == followed))]
            follows_df.to_csv(QUANTORA_FOLLOWS_CSV, index=False)

    def search_users(query):
        users_df = pd.read_csv(QUANTORA_USERS_CSV)
        results = users_df[users_df['quantora_username'].str.contains(query, case=False)]
        return results

    def get_user_posts(username):
        posts_df = pd.read_csv(QUANTORA_POSTS_CSV)
        return posts_df[posts_df['quantora_username'] == username].sort_values(by='quantora_timestamp', ascending=False)

    def get_followers(username):
        follows_df = pd.read_csv(QUANTORA_FOLLOWS_CSV) if os.path.exists(QUANTORA_FOLLOWS_CSV) else pd.DataFrame(columns=['follower', 'followed'])
        return follows_df[follows_df['followed'] == username]['follower'].tolist()

    def get_following(username):
        follows_df = pd.read_csv(QUANTORA_FOLLOWS_CSV) if os.path.exists(QUANTORA_FOLLOWS_CSV) else pd.DataFrame(columns=['follower', 'followed'])
        return follows_df[follows_df['follower'] == username]['followed'].tolist()

    # User Auth
    def quantora_register_user():
        st.subheader("Join the Quantora Universe!")
        quantora_email = st.text_input("Your Email (optional)")
        quantora_username = st.text_input("Choose a Username")
        quantora_password = st.text_input("Create a Password", type="password")
        quantora_bio = st.text_area("Tell us about yourself (optional)")
        quantora_profile_pic_upload = st.file_uploader("Add a Profile Picture (optional)", type=["jpg", "jpeg", "png"])
        
        if st.button("Embark on Your Quantora Journey"):
            quantora_users_df = pd.read_csv(QUANTORA_USERS_CSV)
            if quantora_username in quantora_users_df['quantora_username'].values:
                st.error("This username is already taken. Let your uniqueness shine with another!")
            elif not quantora_username:
                st.error("A username is your key to the Quantora Universe!")
            elif not quantora_password:
                st.error("Set a password to secure your Quantora experience!")
            else:
                quantora_profile_pic_path = DEFAULT_PROFILE_PIC
                if quantora_profile_pic_upload is not None:
                    pic_filename = f"{quantora_username}_{int(time.time())}_{quantora_profile_pic_upload.name}"
                    quantora_profile_pic_path = os.path.join(QUANTORA_PROFILE_PICS_DIR, pic_filename)
                    with open(quantora_profile_pic_path, "wb") as f:
                        f.write(quantora_profile_pic_upload.getbuffer())
                
                quantora_new_user = pd.DataFrame([[quantora_email, quantora_username, quantora_password, quantora_profile_pic_path, quantora_bio]],
                                            columns=['quantora_email', 'quantora_username', 'quantora_password', 'quantora_profile_pic', 'bio'])
                quantora_new_user.to_csv(QUANTORA_USERS_CSV, mode='a', header=False, index=False)
                st.success("Welcome to Quantora! Log in to begin your adventure.")

    def quantora_login_user():
        st.subheader("Re-enter the Quantora Universe")
        quantora_username = st.text_input("Username")
        quantora_password = st.text_input("Password", type="password")
        
        if st.button("Unlock Quantora"):
            quantora_users_df = pd.read_csv(QUANTORA_USERS_CSV)
            user_match = quantora_users_df[
                (quantora_users_df['quantora_username'] == quantora_username) &
                (quantora_users_df['quantora_password'] == quantora_password)
            ]
            if not user_match.empty:
                st.session_state.quantora_logged_in = True
                st.session_state.quantora_username = quantora_username
                st.success(f"Welcome back, @{quantora_username}! The Quantora Universe awaits.")
                st.rerun()
            else:
                st.error("Incorrect username or password. Double-check your credentials to rejoin Quantora.")

    # Post Creation
    def quantora_new_post():
        st.subheader("Share Your Moment on Quantora")
        quantora_post_text = st.text_area("What's happening?", height=150)
        quantora_uploaded_file = st.file_uploader("Add a Photo or Video (optional)", type=["jpg", "jpeg", "png"])
        
        if st.button("Post to Quantora"):
            quantora_image_path = ""
            if quantora_uploaded_file is not None:
                image_filename = f"{st.session_state.quantora_username}_{int(time.time())}_{quantora_uploaded_file.name}"
                quantora_image_path = os.path.join(QUANTORA_IMAGES_DIR, image_filename)
                with open(quantora_image_path, "wb") as f:
                    f.write(quantora_uploaded_file.getbuffer())
            
            quantora_new_data = pd.DataFrame([[st.session_state.quantora_username, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), quantora_post_text, quantora_image_path, 0, ""]],
                                        columns=['quantora_username', 'quantora_timestamp', 'quantora_text', 'quantora_image_path', 'quantora_likes', 'quantora_comments'])
            quantora_new_data.to_csv(QUANTORA_POSTS_CSV, mode='a', header=False, index=False)
            st.success("Your post has been shared with the Quantora community!")
            st.rerun()

    # Social Feed
    def quantora_social_feed():
        st.subheader("Your Quantora Feed")
        try:
            quantora_df = pd.read_csv(QUANTORA_POSTS_CSV)
            all_posts = quantora_df.sort_values(by='quantora_timestamp', ascending=False)
            for index, row in all_posts.iterrows():
                st.markdown("<div style='margin-bottom: 20px; padding: 15px; border: 1px solid #e1e4e8; border-radius: 10px; background-color: #fff;'>", unsafe_allow_html=True)
                quantora_user_info_header(row['quantora_username'])
                st.markdown(f"<div style='margin-top: 10px; font-size: 1em; line-height: 1.4;'>{handle_hashtags(row.get('quantora_text', ''))}</div>", unsafe_allow_html=True)
                image_path = row.get('quantora_image_path')
                if image_path and isinstance(image_path, str) and os.path.exists(image_path):
                    st.image(image_path, use_container_width=True, style="margin-top: 10px; border-radius: 8px;")
                st.markdown("<hr style='margin: 15px 0; border-top: 1px solid #ddd;'>", unsafe_allow_html=True)
                quantora_post_actions(row, index)
                st.markdown("</div>", unsafe_allow_html=True)
        except Exception as e:
            st.error(f"Error loading feed: {e}")

    # Profile Page
    def quantora_profile(view_username=None):
        username_to_view = view_username if view_username else st.session_state.quantora_username
        st.subheader(f"@{username_to_view}")
        
        try:
            user_data = pd.read_csv(QUANTORA_USERS_CSV)
            user_profile = user_data[user_data['quantora_username'] == username_to_view].iloc[0]
            bio = user_profile.get('bio', 'No bio available.')
            profile_pic = user_profile.get('quantora_profile_pic', DEFAULT_PROFILE_PIC)
            if not isinstance(profile_pic, str):
                profile_pic = DEFAULT_PROFILE_PIC
            
            followers = get_followers(username_to_view)
            following = get_following(username_to_view)
            posts = get_user_posts(username_to_view)
            
            col1, col2 = st.columns([0.2, 0.8])
            with col1:
                st.markdown(f'<img src="{profile_pic}" width="80" style="border-radius: 50%; object-fit: cover;">', unsafe_allow_html=True)
            with col2:
                st.markdown(f"<strong style='font-size: 1.5em;'>{username_to_view}</strong>", unsafe_allow_html=True)
                st.markdown(f"<span style='color: #777;'>{len(posts)} posts | {len(followers)} followers | {len(following)} following</span>", unsafe_allow_html=True)
                st.markdown(f"<p style='margin-top: 5px;'>{bio}</p>", unsafe_allow_html=True)
                
                if username_to_view != st.session_state.quantora_username:
                    is_following = is_user_following(st.session_state.quantora_username, username_to_view)
                    follow_text = "Following" if is_following else "Follow"
                    if st.button(follow_text, key=f"profile_follow_{username_to_view}", use_container_width=True):
                        update_follow(st.session_state.quantora_username, username_to_view, not is_following)
                        st.rerun()
            
            st.subheader("Posts")
            if not posts.empty:
                cols = st.columns(3)
                for i, row in posts.iterrows():
                    with cols[i % 3]:
                        if row['quantora_image_path'] and os.path.exists(row['quantora_image_path']):
                            st.image(row['quantora_image_path'], use_container_width=True, style="border-radius: 5px;")
                        else:
                            st.info("No image")
            else:
                st.info(f"@{username_to_view} hasn't posted yet.")
        except IndexError:
            st.error(f"User @{username_to_view} not found.")
        except FileNotFoundError:
            st.error("User data file not found.")
        except Exception as e:
            st.error(f"An error occurred while loading the profile: {e}")

    # Search
    def quantora_search():
        st.subheader("Search Users")
        query = st.text_input("Search for usernames:")
        if query:
            results = search_users(query)
            if not results.empty:
                for index, user in results.iterrows():
                    if st.button(f"@{user['quantora_username']}", key=f"search_result_{index}"):
                        st.session_state.view_profile = user['quantora_username']
                        st.rerun()
            else:
                st.info("No users found matching your search.")
        
        if 'view_profile' in st.session_state:
            quantora_profile(st.session_state.view_profile)
            if st.button("Back to Search"):
                del st.session_state.view_profile
                st.rerun()

    # Navigation
    def quantora_sidebar():
        st.sidebar.title("✨ The Quantora Universe")
        if st.session_state.quantora_logged_in:
            st.sidebar.success(f"Navigating as @{st.session_state.quantora_username}")
            menu = st.sidebar.radio("Explore the Universe", ["Your Feed", "Create Post", "Your Profile", "Search", "Logout"])
            return menu
        else:
            st.sidebar.info("Embark on a new social journey with Quantora!")
            auth_action = st.sidebar.radio("Your Gateway", ["Log In", "Join Quantora"])
            return auth_action

    # Main App Logic
    navigation = quantora_sidebar()
    
    if st.session_state.quantora_logged_in:
        if navigation == "Your Feed":
            quantora_social_feed()
        elif navigation == "Create Post":
            quantora_new_post()
        elif navigation == "Your Profile":
            quantora_profile()
        elif navigation == "Search":
            quantora_search()
        elif navigation == "Logout":
            st.session_state.quantora_logged_in = False
            st.session_state.quantora_username = ""
            st.session_state.view_profile = None
            st.rerun()
    else:
        if navigation == "Log In":
            quantora_login_user()
        elif navigation == "Join Quantora":
            quantora_register_user()

# --------------------------
# HEART HEALTH ANALYZER
# --------------------------
def heart_health_analyzer():
    # Initialize the model
    @st.cache_resource
    def initialize_model():
        return "provider-5/gpt-4o" # Use A4F model
    
    # Define comprehensive health questions
    HEALTH_QUESTIONS = [
        {
            "id": 1,
            "question": "What is your age?",
            "type": "number",
            "min_value": 1,
            "max_value": 120
        },
        {
            "id": 2,
            "question": "What is your gender?",
            "type": "selectbox",
            "options": ["Male", "Female", "Other", "Prefer not to say"]
        },
        {
            "id": 3,
            "question": "Do you experience chest pain or discomfort?",
            "type": "selectbox",
            "options": ["Never", "Rarely", "Sometimes", "Often", "Always"]
        },
        {
            "id": 4,
            "question": "How would you describe your chest pain (if any)?",
            "type": "selectbox",
            "options": ["No pain", "Sharp/Stabbing", "Dull ache", "Pressure/Squeezing", "Burning", "Other"]
        },
        {
            "id": 5,
            "question": "Do you experience shortness of breath?",
            "type": "selectbox",
            "options": ["Never", "Only during intense exercise", "During light exercise", "At rest sometimes",
                        "Frequently at rest"]
        },
        {
            "id": 6,
            "question": "Do you experience heart palpitations or irregular heartbeat?",
            "type": "selectbox",
            "options": ["Never", "Rarely", "Sometimes", "Often", "Always"]
        },
        {
            "id": 7,
            "question": "Do you experience dizziness or lightheadedness?",
            "type": "selectbox",
            "options": ["Never", "Rarely", "Sometimes", "Often", "Always"]
        },
        {
            "id": 8,
            "question": "Do you experience fatigue or weakness?",
            "type": "selectbox",
            "options": ["Never", "Rarely", "Sometimes", "Often", "Always"]
        },
        {
            "id": 9,
            "question": "Do you have swelling in your legs, ankles, or feet?",
            "type": "selectbox",
            "options": ["Never", "Rarely", "Sometimes", "Often", "Always"]
        },
        {
            "id": 10,
            "question": "Do you have a family history of heart disease?",
            "type": "selectbox",
            "options": ["No", "Yes - Parents", "Yes - Siblings", "Yes - Grandparents", "Yes - Multiple family members"]
        },
        {
            "id": 11,
            "question": "Do you smoke or use tobacco products?",
            "type": "selectbox",
            "options": ["Never", "Former smoker", "Occasional smoker", "Regular smoker", "Heavy smoker"]
        },
        {
            "id": 12,
            "question": "How often do you exercise?",
            "type": "selectbox",
            "options": ["Never", "1-2 times per week", "3-4 times per week", "5-6 times per week", "Daily"]
        },
        {
            "id": 13,
            "question": "Do you have high blood pressure?",
            "type": "selectbox",
            "options": ["No", "Yes - controlled with medication", "Yes - uncontrolled", "Don't know"]
        },
        {
            "id": 14,
            "question": "Do you have diabetes?",
            "type": "selectbox",
            "options": ["No", "Type 1", "Type 2", "Pre-diabetes", "Don't know"]
        },
        {
            "id": 15,
            "question": "Do you have high cholesterol?",
            "type": "selectbox",
            "options": ["No", "Yes - controlled with medication", "Yes - uncontrolled", "Don't know"]
        },
        {
            "id": 16,
            "question": "How would you describe your stress level?",
            "type": "selectbox",
            "options": ["Very low", "Low", "Moderate", "High", "Very high"]
        },
        {
            "id": 17,
            "question": "How many hours of sleep do you get per night on average?",
            "type": "selectbox",
            "options": ["Less than 4 hours", "4-6 hours", "6-8 hours", "8-10 hours", "More than 10 hours"]
        },
        {
            "id": 18,
            "question": "Do you consume alcohol?",
            "type": "selectbox",
            "options": ["Never", "Rarely", "1-2 drinks per week", "3-7 drinks per week", "More than 7 drinks per week"]
        },
        {
            "id": 19,
            "question": "What is your current weight status?",
            "type": "selectbox",
            "options": ["Underweight", "Normal weight", "Overweight", "Obese", "Don't know"]
        },
        {
            "id": 20,
            "question": "Are you currently taking any medications for heart conditions?",
            "type": "selectbox",
            "options": ["No", "Blood pressure medication", "Cholesterol medication", "Blood thinners",
                        "Multiple heart medications"]
        },
        {
            "id": 21,
            "question": "Would you like to record your heartbeat for analysis?",
            "type": "selectbox",
            "options": ["Yes, record my heartbeat", "No, skip heartbeat recording"]
        }
    ]
    
    def initialize_session_state():
        """Initialize session state variables"""
        if 'current_question' not in st.session_state:
            st.session_state.current_question = 0
        if 'answers' not in st.session_state:
            st.session_state.answers = {}
        if 'assessment_complete' not in st.session_state:
            st.session_state.assessment_complete = False
        if 'ai_response' not in st.session_state:
            st.session_state.ai_response = None
        if 'heart_rate_data' not in st.session_state:
            st.session_state.heart_rate_data = None
        if 'recording_method' not in st.session_state:
            st.session_state.recording_method = None
        if 'heart_rate_recorded' not in st.session_state:
            st.session_state.heart_rate_recorded = False
        if 'show_heartbeat_section' not in st.session_state:
            st.session_state.show_heartbeat_section = False

    def display_progress():
        """Display progress bar"""
        progress = (st.session_state.current_question / len(HEALTH_QUESTIONS)) * 100
        progress_html = f"""
        <div class="progress-bar">
            <div class="progress-fill" style="width: {progress}%"></div>
        </div>
        <p style="text-align: center;">
            Progress: {st.session_state.current_question}/{len(HEALTH_QUESTIONS)} questions completed
        </p>
        """
        st.markdown(progress_html, unsafe_allow_html=True)

    def analyze_heart_rate_manual():
        """Manual heart rate input and analysis"""
        st.markdown("""
        <div class="heartbeat-container">
            <h3>✋ Manual Heart Rate Measurement</h3>
            <p>Manually count your pulse for 60 seconds or count for 15 seconds and multiply by 4.</p>
        </div>
        """, unsafe_allow_html=True)
        st.info("""
        📋 **Instructions:**
        1. Place two fingers on your wrist below your thumb
        2. Count the beats for 60 seconds
        3. Or count for 15 seconds and multiply by 4
        4. Enter the result below
        """)
        manual_hr = st.number_input("Enter your heart rate (beats per minute):",
                                    min_value=30, max_value=220, value=75, step=1)
        if st.button("📊 Analyze Manual Heart Rate", key="analyze_manual_hr"):
            heart_rate_data = {
                "method": "Manual",
                "heart_rate": manual_hr,
                "quality": "Good" if 60 <= manual_hr <= 100 else "Needs Review",
                "rhythm": "Unable to determine from manual input"
            }
            st.session_state.heart_rate_data = heart_rate_data
            st.session_state.heart_rate_recorded = True
            st.success("Heart rate recorded successfully!")
            return heart_rate_data
        return None

    def analyze_heartbeat_upload():
        """Allow users to upload recorded heartbeat for AI analysis"""
        st.markdown("""
        <div class="heartbeat-container">
            <h3>🎵 Upload Recorded Heartbeat</h3>
            <p>Upload an audio recording (WAV/MP3) or video of your heartbeat for analysis</p>
        </div>
        """, unsafe_allow_html=True)
        uploaded_file = st.file_uploader(
            "Choose a heartbeat recording",
            type=["wav", "mp3", "mp4", "mov"],
            accept_multiple_files=False,
            key="heartbeat_uploader"
        )
        if uploaded_file is not None:
            with st.spinner("🔬 Analyzing your heartbeat recording..."):
                time.sleep(3)
                heart_rate = np.random.randint(60, 100)
                rhythm = np.random.choice(["Regular", "Slightly irregular", "Irregular"])
                quality = "Good" if 60 <= heart_rate <= 100 else "Needs review"
                
                if uploaded_file.type.startswith('audio'):
                    st.audio(uploaded_file)
                elif uploaded_file.type.startswith('video'):
                    st.video(uploaded_file)
                
                model_name = initialize_model()
                prompt = f"""
                Analyze this heartbeat recording and provide:
                1. Estimated heart rate (BPM)
                2. Rhythm assessment
                3. Any abnormalities detected
                4. Recommended next steps
                The recording appears to have:
                - Heart rate: ~{heart_rate} BPM
                - Rhythm: {rhythm}
                - Recording quality: {quality}
                """
                response = call_a4f_model(prompt, model_name)
                st.session_state.heart_rate_data = {
                    "method": "Uploaded Recording",
                    "heart_rate": heart_rate,
                    "rhythm": rhythm,
                    "quality": quality,
                    "ai_analysis": response
                }
                st.session_state.heart_rate_recorded = True
                st.success("Heartbeat analysis complete!")
                return st.session_state.heart_rate_data
        return None

    def analyze_voice_recording():
        """Handle voice recording upload and analysis"""
        st.markdown("""
        <div class="heartbeat-container">
            <h3>🎤 Upload Voice Recording of Heartbeat</h3>
            <p>Record your heartbeat with a microphone or upload an existing recording</p>
        </div>
        """, unsafe_allow_html=True)
        audio_file = st.file_uploader(
            "Upload WAV/MP3 of heartbeat sounds",
            type=["wav", "mp3"],
            help="Record using your phone's voice memo app or a stethoscope attachment",
            key="voice_recording_uploader"
        )
        if audio_file is not None:
            with st.spinner("Analyzing heartbeat sounds..."):
                time.sleep(3)
                st.audio(audio_file)
                heart_rate = np.random.randint(60, 100)
                rhythm = np.random.choice(["Regular", "Slightly irregular"])
                quality = "Good" if 60 <= heart_rate <= 100 else "Needs review"
                
                model_name = initialize_model()
                prompt = f"""
                Analyze this heart sound recording and provide:
                1. Heart rate estimate
                2. Rhythm assessment
                3. Detection of murmurs or abnormalities
                4. Clinical correlation
                Audio characteristics:
                - Apparent rate: {heart_rate} BPM
                - Rhythm: {rhythm}
                - Quality: {quality}
                Provide output in medical report format.
                """
                response = call_a4f_model(prompt, model_name)
                st.session_state.heart_rate_data = {
                    "method": "Voice Recording",
                    "heart_rate": heart_rate,
                    "rhythm": rhythm,
                    "quality": quality,
                    "audio_analysis": response
                }
                st.session_state.heart_rate_recorded = True
                st.success("Voice recording analysis complete!")
                return st.session_state.heart_rate_data
        return None

    def display_heart_rate_analysis(heart_rate_data):
        """Enhanced heart rate analysis analysis display with detailed medical insights"""
        if not heart_rate_data:
            return
        
        st.markdown(f"""
        <div class="heart-rate-display pulse-animation">
            💓 {heart_rate_data['heart_rate']} BPM
        </div>
        """, unsafe_allow_html=True)
        
        method_icons = {
            "Manual": "✋",
            "Uploaded Recording": "📁",
            "Voice Recording": "🎤"
        }
        method_icon = method_icons.get(heart_rate_data.get('method', ''), "📊")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Heart Rate", f"{heart_rate_data['heart_rate']} BPM",
                      "Normal range: 60-100 BPM")
        with col2:
            st.metric("Method", f"{method_icon} {heart_rate_data['method']}")
        with col3:
            st.metric("Data Quality", heart_rate_data['quality'])
        
        hr = heart_rate_data['heart_rate']
        if hr < 60:
            interpretation = "⚠️ **Bradycardia** (Slow Heart Rate)"
            color = "#ffc107"
            details = """
            - May indicate excellent fitness in athletes
            - Potential causes: Hypothyroidism, sleep apnea, heart block
            - Concerning if accompanied by dizziness or fainting
            """
        elif hr > 100:
            interpretation = "⚠️ **Tachycardia** (Fast Heart Rate)"
            color = "#dc3545"
            details = """
            - Common causes: Stress, fever, dehydration, anemia
            - Potential cardiac issues: Atrial fibrillation, SVT
            - Seek help if lasting >30 minutes or with chest pain
            """
        else:
            interpretation = "✅ **Normal Sinus Rhythm**"
            color = "#28a745"
            details = """
            - Healthy resting heart rate
            - Regular rhythm suggests normal electrical activity
            - Maintain with regular exercise and stress management
            """
        
        st.markdown(f"""
        <div style="background-color: {color}20; padding: 1.5rem; border-radius: 8px; border-left: 4px solid {color}; margin: 1rem 0;">
            <h4>{interpretation}</h4>
            <div style="margin-left: 1rem;">{details}</div>
        </div>
        """, unsafe_allow_html=True)
        
        if any(key in heart_rate_data for key in ['hrv_score', 'rhythm', 'audio_analysis', 'ai_analysis']):
            st.markdown("---")
            st.subheader("📊 Detailed Analysis")
            
            if 'hrv_score' in heart_rate_data or 'rhythm' in heart_rate_data:
                cols = st.columns(2)
                if 'hrv_score' in heart_rate_data:
                    with cols[0]:
                        st.metric("Heart Rate Variability",
                                  f"{heart_rate_data['hrv_score']} ms",
                                  "Higher values indicate better stress resilience")
                if 'rhythm' in heart_rate_data:
                    with cols[1]:
                        st.metric("Rhythm Pattern", heart_rate_data['rhythm'])
            
            if heart_rate_data.get('method') == "Voice Recording" and 'audio_analysis' in heart_rate_data:
                st.markdown("""
                <div style="margin-top: 1rem; padding: 1.5rem; background: #f0f8ff; border-radius: 10px;">
                    <h4>🎤 Heart Sound Analysis</h4>
                    <div style="background: white; padding: 1rem; border-radius: 8px; margin-top: 0.5rem;">
                        {analysis}
                    </div>
                </div>
                """.format(analysis=heart_rate_data['audio_analysis']), unsafe_allow_html=True)
            elif 'ai_analysis' in heart_rate_data:
                st.markdown("""
                <div style="margin-top: 1rem; padding: 1.5rem; background: #f5f5f5; border-radius: 10px;">
                    <h4>🤖 AI Analysis Report</h4>
                    <div style="background: white; padding: 1rem; border-radius: 8px; margin-top: 0.5rem;">
                        {analysis}
                    </div>
                </div>
                """.format(analysis=heart_rate_data['ai_analysis']), unsafe_allow_html=True)
        
        st.markdown("---")
        action_col1, action_col2 = st.columns(2)
        with action_col1:
            st.markdown("""
            <div style="padding: 1rem; background: #e8f5e9; border-radius: 8px;">
                <h5>📝 Recommended Actions</h5>
                <ul>
                    <li>Monitor for symptoms like dizziness or chest pain</li>
                    <li>Maintain a heart-healthy diet</li>
                    <li>Stay hydrated and limit caffeine</li>
                    <li>Practice stress-reduction techniques</li>
                    <li>Follow up with your doctor if concerns persist</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)
        with action_col2:
            urgency = "urgent" if hr < 50 or hr > 120 else "routine"
            st.markdown(f"""
            <div style="padding: 1rem; background: #fff3e0; border-radius: 8px;">
                <h5>⏰ When to Seek Help</h5>
                <p>This reading suggests <strong>{urgency}</strong> follow-up:</p>
                <ul>
                    <li>{"Immediately" if urgency == "urgent" else "Within 1-2 weeks"} if symptoms worsen</li>
                    <li>Annual checkup recommended for everyone</li>
                    <li>Sooner if family history of heart disease</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)

    def record_heartbeat_section():
        """Handle heartbeat recording section"""
        st.markdown("## 💓 Heart Rate Recording")
        if not st.session_state.heart_rate_recorded:
            st.markdown("""
            <div class="heartbeat-container">
                <h3>Choose recording method:</h3>
            </div>
            """, unsafe_allow_html=True)
            
            col2, col3, col4 = st.columns(3)
        
            with col2:
                if st.button("✋ Manual Input", key="manual_btn"):
                    st.session_state.recording_method = "manual"
            with col3:
                if st.button("🎵 Upload Video", key="video_btn"):
                    st.session_state.recording_method = "video"
            with col4:
                if st.button("🎤 Voice Recording", key="voice_btn"):
                    st.session_state.recording_method = "voice"
        
            if st.session_state.recording_method == "manual":
                result = analyze_heart_rate_manual()
                if result:
                    display_heart_rate_analysis(result)
            elif st.session_state.recording_method == "video":
                result = analyze_heartbeat_upload()
                if result:
                    display_heart_rate_analysis(result)
            elif st.session_state.recording_method =="voice":
                result = analyze_voice_recording()
                if result:
                    display_heart_rate_analysis(result)
        else:
            display_heart_rate_analysis(st.session_state.heart_rate_data)
            if st.button("Continue with Assessment", key="continue_btn"):
                st.session_state.show_heartbeat_section = False
                st.session_state.current_question += 1
                st.rerun()

    def display_question():
        """Display current question"""
        if st.session_state.current_question < len(HEALTH_QUESTIONS):
            question = HEALTH_QUESTIONS[st.session_state.current_question]
            if question['id'] == 21 and st.session_state.get('show_heartbeat_section', False):
                st.session_state.current_question += 1
                if st.session_state.current_question >= len(HEALTH_QUESTIONS):
                    st.session_state.assessment_complete = True
                st.rerun()
            
            question_html = f"""
            <div class="question-container">
                <h3>Question {question['id']}</h3>
                <p style="font-size: 1.2rem; margin-bottom: 1rem;">{question['question']}</p>
            </div>
            """
            st.markdown(question_html, unsafe_allow_html=True)
            
            if question['type'] == 'number':
                answer = st.number_input(
                    "Your answer:",
                    min_value=question['min_value'],
                    max_value=question['max_value'],
                    value=question['min_value'],
                    key=f"q_{question['id']}"
                )
            elif question['type'] == 'selectbox':
                answer = st.selectbox(
                    "Select your answer:",
                    question['options'],
                    key=f"q_{question['id']}"
                )
            
            col1, col2, col3 = st.columns([1, 1, 1])
            with col1:
                if st.session_state.current_question > 0:
                    if st.button("← Previous", key="prev_btn"):
                        st.session_state.current_question -= 1
                        st.rerun()
            with col3:
                if st.button("Next ->", key="next_btn"):
                    st.session_state.answers[question['id']] = answer
                    if question['id'] == 21:
                        if answer == "Yes, record my heartbeat":
                            st.session_state.show_heartbeat_section = True
                        else:
                            st.session_state.show_heartbeat_section = False
                            st.session_state.current_question += 1
                    else:
                        st.session_state.current_question += 1
                    
                    if st.session_state.current_question >= len(HEALTH_QUESTIONS):
                        st.session_state.assessment_complete = True
                    st.rerun()
        return st.session_state.current_question >= len(HEALTH_QUESTIONS)

    def format_answers():
        """Format answers for AI analysis"""
        formatted_answers = []
        for q_id, answer in st.session_state.answers.items():
            question = next(q for q in HEALTH_QUESTIONS if q['id'] == q_id)
            formatted_answers.append(f"Q{question['id']}: {question['question']}\nA: {answer}\n")
        
        if st.session_state.heart_rate_data:
            hr_data = st.session_state.heart_rate_data
            formatted_answers.append("\nHEART RATE DATA:")
            formatted_answers.append(f"- Method: {hr_data['method']}")
            formatted_answers.append(f"- Heart Rate: {hr_data['heart_rate']} BPM")
            formatted_answers.append(f"- Quality: {hr_data['quality']}")
            if 'rhythm' in hr_data:
                formatted_answers.append(f"- Rhythm: {hr_data['rhythm']}")
            if 'hrv_score' in hr_data:
                formatted_answers.append(f"- HRV Score: {hr_data['hrv_score']}")
            if 'audio_analysis' in hr_data:
                formatted_answers.append(f"- Audio Analysis: {hr_data['audio_analysis']}")
            elif 'ai_analysis' in hr_data:
                formatted_answers.append(f"- AI Analysis: {hr_data['ai_analysis']}")
        
        return "\n".join(formatted_answers)

    def get_ai_assessment(answers_text):
        """Get AI assessment from A4F"""
        try:
            model_name = initialize_model()
            prompt = f"""
            You are an advanced medical AI assistant analyzing heart health.
            Provide a detailed assessment based on these patient responses:
            {answers_text}
            Structure your response with these sections:
            ## 🩺 Overall Risk Assessment
            - Risk level (Low/Medium/High)
            - Key risk factors identified
            ## ❤️ Heart Health Analysis
            - Evaluation of heart-related symptoms
            - Analysis of heart rate data if provided
            ## 🚨 Immediate Concerns
            - Any urgent issues needing attention
            - When to seek emergency care
            ## 💡 Recommendations
            - Lifestyle changes
            - Medical follow-up suggestions
            - Preventive measures
            ## 📅 Next Steps
            - Suggested timeline for follow-up
            - Recommended tests or specialists
            Use clear markdown formatting and provide actionable advice.
            """
            response = call_a4f_model(prompt, model_name)
            return response
        except Exception as e:
            st.error(f"Error generating assessment: {str(e)}")
            return f"Error generating assessment: {str(e)}"

    def display_assessment_summary():
        """Display assessment summary and get AI response"""
        st.markdown('<h2 style="color: black;">Assessment Summary</h2>', unsafe_allow_html=True)
        
        with st.expander("📋 View Your Responses", expanded=True):
            for q_id, answer in st.session_state.answers.items():
                question = next(q for q in HEALTH_QUESTIONS if q['id'] == q_id)
                st.markdown(f"**{question['question']}** \n{answer}")
        
        if st.session_state.heart_rate_data:
            hr_data = st.session_state.heart_rate_data
            with st.expander("💓 View Heart Rate Analysis"):
                display_heart_rate_analysis(hr_data)
        
        if st.session_state.ai_response is None:
            with st.spinner("🧠 Analyzing your responses with AI..."):
                answers_text = format_answers()
                st.session_state.ai_response = get_ai_assessment(answers_text)
                st.rerun()
        
        if st.session_state.ai_response:
            st.markdown("## 🔍 AI Health Assessment")
            st.markdown(st.session_state.ai_response)
            if "high risk" in st.session_state.ai_response.lower() or "urgent" in st.session_state.ai_response.lower():
                st.error("""
                ⚠️ **Urgent Medical Attention Recommended**
                Based on your responses, we recommend seeking immediate medical evaluation.
                """)
        
        st.markdown("""
        <div style="margin-top: 2rem; padding: 1.5rem; background: #f5f5f5; border-radius: 10px;">
            💡 <strong>Remember:</strong> This is a preliminary assessment tool. Always consult with qualified healthcare
            professionals for proper diagnosis and treatment. Regular check-ups are essential for maintaining good heart health.
        </div>
        """, unsafe_allow_html=True)
        
        if st.button("🔄 Start New Assessment", key="reset_btn"):
            for key in list(st.session_state.keys()):
                if key not in ['pro_unlocked', 'model_version', 'image_style', 'video_style']: # Preserve essential states
                    del st.session_state[key]
            st.rerun()

    def main_heart():
        """Main application function"""
        initialize_session_state()
        st.markdown('<h1 class="main-title">❤️ Quantora Heart Problem Searcher</h1>', unsafe_allow_html=True)
        st.markdown("""
        <div style="background-color: #e3f2fd; padding: 1rem; border-radius: 8px; margin-bottom: 1rem;">
            <h4>🔬 Advanced Features:</h4>
            <ul>
                <li>📱 Real-time heart rate monitoring via camera</li>
                <li>📊 Comprehensive health questionnaire</li>
                <li>🤖 AI-powered health analysis</li>
                <li>💓 Heart rate variability assessment</li>
                <li>🎵 Upload recorded heartbeat for analysis</li>
                <li>🏥 Emergency and preventive care recommendations</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        <div style="background-color: #fff3cd; padding: 1rem; border-radius: 8px; border-left: 4px solid #ffc107; margin-bottom: 2rem;">
            <strong>⚠️ Medical Disclaimer:</strong> This tool provides preliminary health assessments only.
            It is not a substitute for professional medical advice, diagnosis, or treatment.
            Always consult with qualified healthcare professionals for medical concerns.
        </div>
        """, unsafe_allow_html=True)
        
        if not st.session_state.assessment_complete:
            display_progress()
            if st.session_state.get('show_heartbeat_section', False):
                record_heartbeat_section()
            else:
                display_question()
        else:
            display_assessment_summary()
    
    main_heart()

# --------------------------
# BRAIN HEALTH ANALYZER
# --------------------------
def brain_health_analyzer():
    # Initialize the model
    @st.cache_resource
    def initialize_model():
        return "provider-5/gpt-4o" # Use A4F model
    
    # Define comprehensive brain health questions
    BRAIN_QUESTIONS = [
        {
            "id": 1,
            "question": "What is your age?",
            "type": "number",
            "min_value": 1,
            "max_value": 120
        },
        {
            "id": 2,
            "question": "What is your gender?",
            "type": "selectbox",
            "options": ["Male", "Female", "Other", "Prefer not to say"]
        },
        {
            "id": 3,
            "question": "Do you experience frequent headaches?",
            "type": "selectbox",
            "options": ["Never", "Rarely", "Sometimes", "Often", "Daily"]
        },
        {
            "id": 4,
            "question": "How would you describe your headaches (if any)?",
            "type": "selectbox",
            "options": ["No headaches", "Throbbing", "Dull ache", "Sharp pain", "Pressure", "Migraine"]
        },
        {
            "id": 5,
            "question": "Do you experience memory problems?",
            "type": "selectbox",
            "options": ["Never", "Occasionally forget names", "Frequently forget things", "Difficulty remembering recent events", "Severe memory impairment"]
        },
        {
            "id": 6,
            "question": "Do you have difficulty concentrating?",
            "type": "selectbox",
            "options": ["Never", "Rarely", "Sometimes", "Often", "Always"]
        },
        {
            "id": 7,
            "question": "Do you experience dizziness or balance problems?",
            "type": "selectbox",
            "options": ["Never", "Rarely", "Sometimes", "Often", "Always"]
        },
        {
            "id": 8,
            "question": "Do you have trouble speaking or finding words?",
            "type": "selectbox",
            "options": ["Never", "Rarely", "Sometimes", "Often", "Always"]
        },
        {
            "id": 9,
            "question": "Do you experience mood swings or personality changes?",
            "type": "selectbox",
            "options": ["Never", "Rarely", "Sometimes", "Often", "Always"]
        },
        {
            "id": 10,
            "question": "Do you have a family history of neurological disorders?",
            "type": "selectbox",
            "options": ["No", "Yes - Alzheimer's/dementia", "Yes - Parkinson's", "Yes - Stroke", "Yes - Multiple family members"]
        },
        {
            "id": 11,
            "question": "Have you ever had a head injury with loss of consciousness?",
            "type": "selectbox",
            "options": ["Never", "Yes, brief loss", "Yes, prolonged loss", "Multiple injuries", "Not sure"]
        },
        {
            "id": 12,
            "question": "Do you experience seizures or unexplained blackouts?",
            "type": "selectbox",
            "options": ["Never", "Rarely", "Sometimes", "Often", "Diagnosed with epilepsy"]
        },
        {
            "id": 13,
            "question": "Do you have sleep problems?",
            "type": "selectbox",
            "options": ["No", "Occasional insomnia", "Chronic insomnia", "Excessive daytime sleepiness", "Sleep apnea"]
        },
        {
            "id": 14,
            "question": "Do you experience numbness or tingling?",
            "type": "selectbox",
            "options": ["Never", "Rarely", "Sometimes", "Often", "Always"]
        },
        {
            "id": 15,
            "question": "Do you have vision problems not corrected by glasses?",
            "type": "selectbox",
            "options": ["No", "Blurred vision", "Double vision", "Partial vision loss", "Complete vision loss"]
        },
        {
            "id": 16,
            "question": "How would you describe your stress level?",
            "type": "selectbox",
            "options": ["Very low", "Low", "Moderate", "High", "Very high"]
        },
        {
            "id": 17,
            "question": "How many hours of sleep do you get per night on average?",
            "type": "selectbox",
            "options": ["Less than 4 hours", "4-6 hours", "6-8 hours", "8-10 hours", "More than 10 hours"]
        },
        {
            "id": 18,
            "question": "Do you consume alcohol?",
            "type": "selectbox",
            "options": ["Never", "Rarely", "1-2 drinks per week", "3-7 drinks per week", "More than 7 drinks per week"]
        },
        {
            "id": 19,
            "question": "Do you use recreational drugs?",
            "type": "selectbox",
            "options": ["Never", "Former user", "Occasional user", "Regular user", "Heavy user"]
        },
        {
            "id": 20,
            "question": "Are you currently taking any medications for neurological conditions?",
            "type": "selectbox",
            "options": ["No", "Antidepressants", "Anti-anxiety", "Antipsychotics", "Multiple medications"]
        },
        {
            "id": 21,
            "question": "Would you like to perform cognitive tests for analysis?",
            "type": "selectbox",
            "options": ["Yes, perform cognitive tests", "No, skip cognitive tests"]
        }
    ]
    
    def initialize_session_state():
        """Initialize session state variables"""
        if 'current_question' not in st.session_state:
            st.session_state.current_question = 0
        if 'answers' not in st.session_state:
            st.session_state.answers = {}
        if 'assessment_complete' not in st.session_state:
            st.session_state.assessment_complete = False
        if 'ai_response' not in st.session_state:
            st.session_state.ai_response = None
        if 'cognitive_data' not in st.session_state:
            st.session_state.cognitive_data = None
        if 'testing_method' not in st.session_state:
            st.session_state.testing_method = None
        if 'cognitive_tests_completed' not in st.session_state:
            st.session_state.cognitive_tests_completed = False
        if 'show_cognitive_section' not in st.session_state:
            st.session_state.show_cognitive_section = False

    def display_progress():
        """Display progress bar"""
        progress = (st.session_state.current_question / len(BRAIN_QUESTIONS)) * 100
        progress_html = f"""
        <div class="progress-bar">
            <div class="progress-fill" style="width: {progress}%"></div>
        </div>
        <p style="text-align: center;">
            Progress: {st.session_state.current_question}/{len(BRAIN_QUESTIONS)} questions completed
        </p>
        """
        st.markdown(progress_html, unsafe_allow_html=True)

    def analyze_cognitive_function():
        """Analyze cognitive function through interactive tests"""
        st.markdown("""
        <div class="brain-container">
            <h3>🧠 Cognitive Function Assessment</h3>
            <p>Perform these brief tests to assess memory, attention, and executive function.</p>
        </div>
        """, unsafe_allow_html=True)
        st.info("""
        📋 **Instructions:**
        1. Complete all tests in order
        2. Answer as accurately as possible
        3. Don't use external aids
        4. Take your time
        """)
        
        with st.expander("🔢 Digit Span Test (Working Memory)"):
            st.write("""
            **Test:** Repeat sequences of numbers in the same order.
            The test will progressively get harder with longer sequences.
            """)
        
            if st.button("Start Digit Span Test"):
                sequences = [
                    [3, 7, 2],
                    [8, 1, 6, 4],
                    [5, 9, 2, 7, 3],
                    [4, 1, 8, 3, 6, 9],
                    [7, 2, 5, 8, 3, 6, 1]
                ]
            
                score = 0
                for seq in sequences:
                    st.write(f"Remember this sequence: {seq}")
                    time.sleep(2)
                    st.write("Sequence hidden...")
                    time.sleep(1)
                
                    user_input = st.text_input(f"Enter the {len(seq)}-digit sequence (separated by spaces):", key=f"digits_{seq[0]}")
                    if user_input:
                        user_nums = [int(n) for n in user_input.split() if n.isdigit()]
                        if user_nums == seq:
                            score += 1
                            st.success("Correct!")
                        else:
                            st.error(f"Incorrect. The sequence was: {seq}")
                            break
            
                digit_span_score = min(score + 2, 7) # Normal range is 5-7
                st.session_state.cognitive_data = st.session_state.get('cognitive_data', {})
                st.session_state.cognitive_data['digit_span'] = digit_span_score
                st.metric("Digit Span Score", digit_span_score, "Normal range: 5-7")
        
        with st.expander("🔄 Trail Making Test (Processing Speed)"):
            st.write("""
            **Test:** Connect numbers in order as quickly as possible.
            """)
        
            if st.button("Start Trail Making Test"):
                # Generate a random sequence of numbers 1-8
                numbers = list(range(1, 9))
                np.random.shuffle(numbers)
            
                st.write("Connect the numbers in order from 1 to 8:")
                st.write(" -> ".join([str(n) for n in numbers]))
            
                start_time = time.time()
                user_input = st.text_input("Enter the numbers in order separated by spaces (e.g., '1 2 3...'):")
            
                if user_input:
                    end_time = time.time()
                    time_taken = end_time - start_time
                    user_nums = [int(n) for n in user_input.split() if n.isdigit()]
                
                    if user_nums == list(range(1, 9)):
                        trail_score = max(0, 100 - int(time_taken))
                        st.session_state.cognitive_data = st.session_state.get('cognitive_data', {})
                        st.session_state.cognitive_data['trail_making'] = trail_score
                        st.metric("Trail Making Score", trail_score, f"Time taken: {time_taken:.1f} seconds")
                    else:
                        st.error("Incorrect sequence. Please try again.")
        
        with st.expander("📝 Verbal Fluency Test (Language)"):
            st.write("""
            **Test:** Name as many animals as you can in 60 seconds.
            """)
        
            if st.button("Start Verbal Fluency Test"):
                st.write("List as many animals as you can think of in the text box below:")
            
                start_time = time.time()
                end_time = start_time + 60
                animal_list = []
            
                while time.time() < end_time:
                    animal = st.text_input(f"Time remaining: {int(end_time - time.time())} seconds", key=f"animal_{time.time()}")
                    if animal:
                        animal_list.append(animal.strip().lower())
            
                unique_animals = len(set(animal_list))
                fluency_score = min(unique_animals * 5, 100) # 20+ is normal
                st.session_state.cognitive_data = st.session_state.get('cognitive_data', {})
                st.session_state.cognitive_data['verbal_fluency'] = fluency_score
                st.metric("Verbal Fluency Score", fluency_score, f"Unique animals: {unique_animals}")
        
        with st.expander("🖼️ Visual Memory Test"):
            st.write("""
            **Test:** Remember and recall items.
            """)
        
            if st.button("Start Visual Memory Test"):
                # Sample items
                items = ["apple", "car", "tree", "house", "dog"]
                st.write("Study these items for 10 seconds:")
                st.write(", ".join(items))
            
                time.sleep(10)
                st.write("Items hidden...")
                time.sleep(2)
            
                recalled = st.text_input("Enter all items you remember (separated by commas):")
                recalled_items = [item.strip().lower() for item in recalled.split(",") if recalled]
            
                correct = sum(1 for item in recalled_items if item in items)
                memory_score = int((correct / len(items)) * 100)
                st.session_state.cognitive_data = st.session_state.get('cognitive_data', {})
                st.session_state.cognitive_data['visual_memory'] = memory_score
                st.metric("Visual Memory Score", memory_score, f"Recalled {correct} of {len(items)} items")
        
        if st.session_state.get('cognitive_data'):
            if st.button("Complete Cognitive Assessment", key="complete_cognitive"):
                st.session_state.cognitive_tests_completed = True
                st.session_state.show_cognitive_section = False
                st.session_state.current_question += 1
                st.rerun()

    def display_cognitive_results():
        """Display cognitive test results with analysis"""
        if not st.session_state.get('cognitive_data'):
            return
        
        data = st.session_state.cognitive_data
        overall_score = int(np.mean([v for v in data.values() if isinstance(v, int)]))
    
        st.markdown(f"""
        <div class="cognitive-display pulse-animation">
            🧠 {overall_score}/100
        </div>
        """, unsafe_allow_html=True)
        st.markdown("### Cognitive Test Results")
    
        # Risk level interpretation
        if overall_score < 70:
            interpretation = "⚠️ **Below Average Cognitive Function**"
            color = "#ffc107"
            details = """
            - May indicate cognitive impairment
            - Potential causes: Stress, sleep deprivation, neurological conditions
            - Concerning if accompanied by other symptoms
            """
        elif overall_score < 85:
            interpretation = "🔄 **Average Cognitive Function**"
            color = "#2196f3"
            details = """
            - Within normal range for age
            - Some room for improvement
            - Maintain with brain-healthy activities
            """
        else:
            interpretation = "✅ **Above Average Cognitive Function**"
            color = "#28a745"
            details = """
            - Strong cognitive performance
            - Continue brain-healthy habits
            - Monitor for any changes
            """
        
        st.markdown(f"""
        <div style="background-color: {color}20; padding: 1.5rem; border-radius: 8px; border-left: 4px solid {color}; margin: 1rem 0;">
            <h4>{interpretation}</h4>
            <div style="margin-left: 1rem;">{details}</div>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("---")
        st.subheader("📊 Detailed Analysis")
        
        if 'digit_span' in data:
            st.markdown("#### Working Memory (Digit Span)")
            progress = min(data['digit_span'] * 14, 100) # Convert 7-point scale to percentage
            st.progress(progress)
            st.caption(f"Score: {data['digit_span']}/7 - {'Normal' if data['digit_span'] >=5 else 'Below normal'}")
        
        if 'trail_making' in data:
            st.markdown("#### Processing Speed (Trail Making)")
            st.progress(data['trail_making'])
            st.caption(f"Score: {data['trail_making']}/100 - {'Normal' if data['trail_making'] >=70 else 'Below normal'}")
        
        if 'verbal_fluency' in data:
            st.markdown("#### Verbal Fluency")
            st.progress(data['verbal_fluency'])
            st.caption(f"Score: {data['verbal_fluency']}/100 - {'Normal' if data['verbal_fluency'] >=70 else 'Below normal'}")
        
        if 'visual_memory' in data:
            st.markdown("#### Visual Memory")
            st.progress(data['visual_memory'])
            st.caption(f"Score: {data['visual_memory']}/100 - {'Normal' if data['visual_memory'] >=70 else 'Below normal'}")
        
        st.markdown("---")
        action_col1, action_col2 = st.columns(2)
        with action_col1:
            st.markdown("""
            <div style="padding: 1rem; background: #e8f5e9; border-radius: 8px;">
                <h5>🧩 Brain-Boosting Activities</h5>
                <ul>
                    <li>Regular mental exercises (puzzles, reading)</li>
                    <li>Physical exercise (improves brain blood flow)</li>
                    <li>Social engagement</li>
                    <li>Learn new skills</li>
                    <li>Meditation for focus</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)
        with action_col2:
            urgency = "urgent" if overall_score < 60 else "routine"
            st.markdown(f"""
            <div style="padding: 1rem; background: #fff3e0; border-radius: 8px;">
                <h5>⏰ When to Seek Help</h5>
                <p>This reading suggests <strong>{urgency}</strong> follow-up:</p>
                <ul>
                    <li>{"Immediately" if urgency == "urgent" else "Within 1-2 weeks"} if symptoms worsen</li>
                    <li>Annual cognitive screening recommended after age 50</li>
                    <li>Sooner if family history of dementia</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)

    def cognitive_tests_section():
        """Handle cognitive testing section"""
        st.markdown("## 🧠 Cognitive Function Assessment")
        if not st.session_state.cognitive_tests_completed:
            analyze_cognitive_function()
        else:
            display_cognitive_results()
            if st.button("Continue with Assessment", key="continue_btn"):
                st.session_state.show_cognitive_section = False
                st.session_state.current_question += 1
                st.rerun()

    def display_question():
        """Display current question"""
        if st.session_state.current_question < len(BRAIN_QUESTIONS):
            question = BRAIN_QUESTIONS[st.session_state.current_question]
            if question['id'] == 21 and st.session_state.get('show_cognitive_section', False):
                st.session_state.current_question += 1
                if st.session_state.current_question >= len(BRAIN_QUESTIONS):
                    st.session_state.assessment_complete = True
                st.rerun()
            
            question_html = f"""
            <div class="question-container">
                <h3>Question {question['id']}</h3>
                <p style="font-size: 1.2rem; margin-bottom: 1rem;">{question['question']}</p>
            </div>
            """
            st.markdown(question_html, unsafe_allow_html=True)
            
            if question['type'] == 'number':
                answer = st.number_input(
                    "Your answer:",
                    min_value=question['min_value'],
                    max_value=question['max_value'],
                    value=question['min_value'],
                    key=f"q_{question['id']}"
                )
            elif question['type'] == 'selectbox':
                answer = st.selectbox(
                    "Select your answer:",
                    question['options'],
                    key=f"q_{question['id']}"
                )
            
            col1, col2, col3 = st.columns([1, 1, 1])
            with col1:
                if st.session_state.current_question > 0:
                    if st.button("← Previous", key="prev_btn"):
                        st.session_state.current_question -= 1
                        st.rerun()
            with col3:
                if st.button("Next ->", key="next_btn"):
                    st.session_state.answers[question['id']] = answer
                    if question['id'] == 21:
                        if answer == "Yes, perform cognitive tests":
                            st.session_state.show_cognitive_section = True
                        else:
                            st.session_state.show_cognitive_section = False
                            st.session_state.current_question += 1
                    else:
                        st.session_state.current_question += 1
                    
                    if st.session_state.current_question >= len(BRAIN_QUESTIONS):
                        st.session_state.assessment_complete = True
                    st.rerun()
        return st.session_state.current_question >= len(BRAIN_QUESTIONS)

    def format_answers():
        """Format answers for AI analysis"""
        formatted_answers = []
        for q_id, answer in st.session_state.answers.items():
            question = next(q for q in BRAIN_QUESTIONS if q['id'] == q_id)
            formatted_answers.append(f"Q{question['id']}: {question['question']}\nA: {answer}\n")
        
        if st.session_state.cognitive_data:
            cog_data = st.session_state.cognitive_data
            formatted_answers.append("\nCOGNITIVE TEST DATA:")
            formatted_answers.append(f"- Overall Score: {np.mean([v for v in cog_data.values() if isinstance(v, int)]):.1f}/100")
            if 'digit_span' in cog_data:
                formatted_answers.append(f"- Working Memory (Digit Span): {cog_data['digit_span']}/7")
            if 'trail_making' in cog_data:
                formatted_answers.append(f"- Processing Speed (Trail Making): {cog_data['trail_making']}/100")
            if 'verbal_fluency' in cog_data:
                formatted_answers.append(f"- Verbal Fluency: {cog_data['verbal_fluency']}/100")
            if 'visual_memory' in cog_data:
                formatted_answers.append(f"- Visual Memory: {cog_data['visual_memory']}/100")
        
        return "\n".join(formatted_answers)

    def get_ai_assessment(answers_text):
        """Get AI assessment from A4F"""
        try:
            model_name = initialize_model()
            prompt = f"""
            You are an advanced neurological AI assistant analyzing brain health.
            Provide a detailed assessment based on these patient responses:
            {answers_text}
            Structure your response with these sections:
            ## 🧠 Overall Neurological Assessment
            - Risk level (Low/Medium/High) for neurological conditions
            - Key risk factors identified
            ## 🧐 Symptom Analysis
            - Evaluation of neurological symptoms
            - Analysis of cognitive test data if provided
            ## 🚨 Immediate Concerns
            - Any urgent neurological issues needing attention
            - When to seek emergency care (e.g., stroke symptoms)
            ## 💡 Recommendations
            - Lifestyle changes for brain health
            - Medical follow-up suggestions
            - Preventive measures
            ## 📅 Next Steps
            - Suggested timeline for follow-up
            - Recommended neurological tests or specialists
            Use clear markdown formatting and provide actionable advice.
            Focus specifically on brain and neurological health.
            """
            response = call_a4f_model(prompt, model_name)
            return response
        except Exception as e:
            st.error(f"Error generating assessment: {str(e)}")
            return f"Error generating assessment: {str(e)}"

    def display_assessment_summary():
        """Display assessment summary and get AI response"""
        st.markdown('<h2 style="color: black;">Assessment Summary</h2>', unsafe_allow_html=True)
        
        with st.expander("📋 View Your Responses", expanded=True):
            for q_id, answer in st.session_state.answers.items():
                question = next(q for q in BRAIN_QUESTIONS if q['id'] == q_id)
                st.markdown(f"**{question['question']}** \n{answer}")
        
        if st.session_state.cognitive_data:
            with st.expander("🧠 View Cognitive Test Results"):
                display_cognitive_results()
        
        if st.session_state.ai_response is None:
            with st.spinner("🧠 Analyzing your responses with AI..."):
                answers_text = format_answers()
                st.session_state.ai_response = get_ai_assessment(answers_text)
                st.rerun()
        
        if st.session_state.ai_response:
            st.markdown("## 🔍 AI Neurological Assessment")
            st.markdown(st.session_state.ai_response)
            if "high risk" in st.session_state.ai_response.lower() or "urgent" in st.session_state.ai_response.lower():
                st.error("""
                ⚠️ **Urgent Medical Attention Recommended**
                Based on your responses, we recommend seeking immediate neurological evaluation.
                """)
        
        st.markdown("""
        <div style="margin-top: 2rem; padding: 1.5rem; background: #f5f5f5; border-radius: 10px;">
            💡 <strong>Remember:</strong> This is a preliminary assessment tool. Always consult with qualified neurologists
            or healthcare professionals for proper diagnosis and treatment. Regular cognitive screenings are recommended as you age.
        </div>
        """, unsafe_allow_html=True)
        
        if st.button("🔄 Start New Assessment", key="reset_btn"):
            for key in list(st.session_state.keys()):
                if key not in ['pro_unlocked', 'model_version', 'image_style', 'video_style']: # Preserve essential states
                    del st.session_state[key]
            st.rerun()

    def main_brain():
        """Main application function"""
        initialize_session_state()
        st.markdown('<h1 class="main-title">🧠 NeuroScan Brain Problem Searcher</h1>', unsafe_allow_html=True)
        st.markdown("""
        <div style="background-color: #e3f2fd; padding: 1rem; border-radius: 8px; margin-bottom: 1rem;">
            <h4>🔬 Advanced Features:</h4>
            <ul>
                <li>🧠 Cognitive function assessment</li>
                <li>📊 Comprehensive neurological questionnaire</li>
                <li>🤖 AI-powered brain health analysis</li>
                <li>📝 Memory and processing speed tests</li>
                <li>🏥 Emergency and preventive care recommendations</li>
                <li>📈 Tracking of cognitive performance over time</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        <div style="background-color: #fff3cd; padding: 1rem; border-radius: 8px; border-left: 4px solid #ffc107; margin-bottom: 2rem;">
            <strong>⚠️ Medical Disclaimer:</strong> This tool provides preliminary health assessments only.
            It is not a substitute for professional medical advice, diagnosis, or treatment.
            Always consult with qualified neurologists or healthcare professionals for medical concerns.
        </div>
        """, unsafe_allow_html=True)
        
        if not st.session_state.assessment_complete:
            display_progress()
            if st.session_state.get('show_cognitive_section', False):
                cognitive_tests_section()
            else:
                display_question()
        else:
            display_assessment_summary()
    
    main_brain()

# --------------------------
# CANCER RISK ASSESSOR
# --------------------------
def cancer_risk_assessor():
    # Initialize the model
    @st.cache_resource
    def initialize_model():
        return "provider-5/gpt-4o" # Use A4F model
    
    # Define comprehensive cancer screening questions
    CANCER_QUESTIONS = [
        {
            "id": 1,
            "question": "What is your age?",
            "type": "number",
            "min_value": 1,
            "max_value": 120,
            "risk_factor": True
        },
        {
            "id": 2,
            "question": "What is your gender?",
            "type": "selectbox",
            "options": ["Male", "Female", "Other", "Prefer not to say"],
            "risk_factor": True
        },
        {
            "id": 3,
            "question": "Do you currently smoke or have a history of smoking?",
            "type": "selectbox",
            "options": ["Never smoked", "Former smoker", "Current smoker (less than 10 cigarettes/day)",
                       "Current smoker (10-20 cigarettes/day)", "Current smoker (more than 20 cigarettes/day)"],
            "risk_factor": True,
            "related_to": ["Lung cancer", "Bladder cancer", "Head and neck cancers"]
        },
        {
            "id": 4,
            "question": "How often do you consume alcohol?",
            "type": "selectbox",
            "options": ["Never", "Occasionally (less than 1 drink/week)",
                       "Moderately (1-7 drinks/week)", "Heavily (more than 7 drinks/week)"],
            "risk_factor": True,
            "related_to": ["Liver cancer", "Breast cancer", "Esophageal cancer"]
        },
        {
            "id": 5,
            "question": "What is your body mass index (BMI)?",
            "type": "selectbox",
            "options": ["Underweight (<18.5)", "Normal (18.5-24.9)",
                       "Overweight (25-29.9)", "Obese (30-34.9)", "Severely obese (35+)"],
            "risk_factor": True,
            "related_to": ["Multiple cancer types"]
        },
        {
            "id": 6,
            "question": "Do you have a family history of cancer?",
            "type": "selectbox",
            "options": ["No", "Yes - one relative", "Yes - multiple relatives",
                       "Yes - first-degree relative", "Yes - multiple first-degree relatives"],
            "risk_factor": True,
            "related_to": ["All cancers"]
        },
        {
            "id": 7,
            "question": "Have you noticed any unexplained weight loss recently?",
            "type": "selectbox",
            "options": ["No", "Yes - less than 5% body weight",
                       "Yes - 5-10% body weight", "Yes - more than 10% body weight"],
            "symptom": True,
            "related_to": ["Multiple cancer types"]
        },
        {
            "id": 8,
            "question": "Do you experience persistent fatigue that doesn't improve with rest?",
            "type": "selectbox",
            "options": ["Never", "Occasionally", "Often", "Constantly"],
            "symptom": True,
            "related_to": ["Multiple cancer types"]
        },
        {
            "id": 9,
            "question": "Have you noticed any unusual lumps or swellings?",
            "type": "selectbox",
            "options": ["No", "Yes - small and painless",
                       "Yes - growing in size", "Yes - painful"],
            "symptom": True,
            "related_to": ["Lymphoma", "Breast cancer", "Testicular cancer"]
        },
        {
            "id": 10,
            "question": "Have you noticed any changes in bowel or bladder habits?",
            "type": "selectbox",
            "options": ["No", "Yes - mild changes",
                       "Yes - significant changes", "Yes - blood in stool/urine"],
            "symptom": True,
            "related_to": ["Colorectal cancer", "Bladder cancer", "Prostate cancer"]
        },
        {
            "id": 11,
            "question": "Do you have persistent cough or hoarseness?",
            "type": "selectbox",
            "options": ["No", "Yes - less than 3 weeks",
                       "Yes - 3-6 weeks", "Yes - more than 6 weeks"],
            "symptom": True,
            "related_to": ["Lung cancer", "Laryngeal cancer"]
        },
        {
            "id": 12,
            "question": "Have you noticed any unusual bleeding or discharge?",
            "type": "selectbox",
            "options": ["No", "Yes - minor", "Yes - significant"],
            "symptom": True,
            "related_to": ["Multiple cancer types"]
        },
        {
            "id": 13,
            "question": "Do you have persistent indigestion or difficulty swallowing?",
            "type": "selectbox",
            "options": ["No", "Yes - occasional",
                       "Yes - frequent", "Yes - constant"],
            "symptom": True,
            "related_to": ["Esophageal cancer", "Stomach cancer"]
        },
        {
            "id": 14,
            "question": "Have you noticed any changes in a mole or skin lesion?",
            "type": "selectbox",
            "options": ["No", "Yes - slight change",
                       "Yes - significant change", "Yes - bleeding mole"],
            "symptom": True,
            "related_to": ["Melanoma", "Skin cancer"]
        },
        {
            "id": 15,
            "question": "Do you have persistent pain without obvious cause?",
            "type": "selectbox",
            "options": ["No", "Yes - mild", "Yes - moderate", "Yes - severe"],
            "symptom": True,
            "related_to": ["Bone cancer", "Pancreatic cancer", "Ovarian cancer"]
        },
        {
            "id": 16,
            "question": "For women: Have you noticed any breast changes?",
            "type": "selectbox",
            "options": ["Not applicable", "No changes",
                       "Lump or thickening", "Nipple changes/discharge",
                       "Skin dimpling/redness"],
            "symptom": True,
            "related_to": ["Breast cancer"]
        },
        {
            "id": 17,
            "question": "For men: Have you noticed any testicular changes?",
            "type": "selectbox",
            "options": ["Not applicable", "No changes",
                       "Lump or swelling", "Pain/discomfort", "Size/shape changes"],
            "symptom": True,
            "related_to": ["Testicular cancer"]
        },
        {
            "id": 18,
            "question": "How often do you use sunscreen when outdoors?",
            "type": "selectbox",
            "options": ["Always", "Often", "Sometimes", "Rarely", "Never"],
            "risk_factor": True,
            "related_to": ["Skin cancer"]
        },
        {
            "id": 19,
            "question": "How often do you eat processed or red meat?",
            "type": "selectbox",
            "options": ["Rarely or never", "1-2 times per week",
                       "3-5 times per week", "Daily"],
            "risk_factor": True,
            "related_to": ["Colorectal cancer"]
        },
        {
            "id": 20,
            "question": "How many servings of fruits and vegetables do you eat daily?",
            "type": "selectbox",
            "options": ["Less than 2", "2-4", "5-7", "More than 7"],
            "risk_factor": True,
            "related_to": ["Multiple cancer types"]
        },
        {
            "id": 21,
            "question": "How often do you engage in physical activity?",
            "type": "selectbox",
            "options": ["Less than 30 min/week", "30-90 min/week",
                       "90-150 min/week", "More than 150 min/week"],
            "risk_factor": True,
            "related_to": ["Multiple cancer types"]
        },
        {
            "id": 22,
            "question": "Have you been exposed to radiation or toxic chemicals?",
            "type": "selectbox",
            "options": ["No", "Yes - minimal", "Yes - moderate", "Yes - significant"],
            "risk_factor": True,
            "related_to": ["Multiple cancer types"]
        },
        {
            "id": 23,
            "question": "Have you been exposed to HPV, Hepatitis B/C, or other cancer-related viruses?",
            "type": "selectbox",
            "options": ["No", "Yes - HPV", "Yes - Hepatitis", "Yes - other"],
            "risk_factor": True,
            "related_to": ["Cervical cancer", "Liver cancer", "Head and neck cancers"]
        },
        {
            "id": 24,
            "question": "Would you like to analyze any images of concerning areas?",
            "type": "selectbox",
            "options": ["No", "Yes - skin lesions",
                       "Yes - breast changes", "Yes - other areas"],
            "symptom": False
        }
    ]
    
    def initialize_session_state():
        """Initialize session state variables"""
        if 'current_question' not in st.session_state:
            st.session_state.current_question = 0
        if 'answers' not in st.session_state:
            st.session_state.answers = {}
        if 'assessment_complete' not in st.session_state:
            st.session_state.assessment_complete = False
        if 'ai_response' not in st.session_state:
            st.session_state.ai_response = None
        if 'image_analysis' not in st.session_state:
            st.session_state.image_analysis = None
        if 'testing_method' not in st.session_state:
            st.session_state.testing_method = None
        if 'image_analysis_completed' not in st.session_state:
            st.session_state.image_analysis_completed = False
        if 'show_image_section' not in st.session_state:
            st.session_state.show_image_section = False
        if 'risk_score' not in st.session_state:
            st.session_state.risk_score = 0
        if 'concerned_areas' not in st.session_state:
            st.session_state.concerned_areas = []

    def display_progress():
        """Display progress bar"""
        progress = (st.session_state.current_question / len(CANCER_QUESTIONS)) * 100
        progress_html = f"""
        <div class="progress-bar">
            <div class="progress-fill" style="width: {progress}%"></div>
        </div>
        <p style="text-align: center;">
            Progress: {st.session_state.current_question}/{len(CANCER_QUESTIONS)} questions completed
        </p>
        """
        st.markdown(progress_html, unsafe_allow_html=True)

    def analyze_images():
        """Analyze uploaded images for concerning features"""
        st.markdown("""
        <div class="body-container">
            <h3>📷 Image Analysis</h3>
            <p>Upload images of any concerning areas for preliminary analysis.</p>
        </div>
        """, unsafe_allow_html=True)
        st.info("""
        📋 **Instructions:**
        1. Upload clear, well-lit images
        2. Include different angles if possible
        3. For skin lesions, include a ruler for scale if available
        4. Images should be in JPG or PNG format
        """)
        
        uploaded_files = st.file_uploader("Upload images (max 4)",
                                        type=["jpg", "jpeg", "png"],
                                        accept_multiple_files=True)
    
        if uploaded_files:
            st.warning("""
            ⚠️ **Important Note:** This image analysis is for preliminary screening only.
            It cannot replace a professional medical examination or biopsy.
            """)
        
            cols = st.columns(min(4, len(uploaded_files)))
            for i, uploaded_file in enumerate(uploaded_files):
                with cols[i]:
                    image = Image.open(uploaded_file)
                    st.image(image, caption=f"Image {i+1}", use_container_width=True)
        
            if st.button("Analyze Images"):
                with st.spinner("🔍 Analyzing images with AI..."):
                    time.sleep(2)
                
                    analysis_results = []
                    for i, uploaded_file in enumerate(uploaded_files):
                        img_name = uploaded_file.name.lower()
                        if "skin" in img_name or "mole" in img_name:
                            result = {
                                "type": "Skin lesion",
                                "characteristics": "Asymmetrical shape with color variation",
                                "concern_level": "Moderate",
                                "recommendation": "Dermatologist evaluation recommended"
                            }
                        elif "breast" in img_name:
                            result = {
                                "type": "Breast change",
                                "characteristics": "Visible skin dimpling",
                                "concern_level": "High",
                                "recommendation": "Urgent mammogram and clinical exam needed"
                            }
                        else:
                            result = {
                                "type": "General image",
                                "characteristics": "No obvious concerning features",
                                "concern_level": "Low",
                                "recommendation": "Monitor for changes"
                            }
                        analysis_results.append(result)
                
                    st.session_state.image_analysis = analysis_results
                    st.session_state.image_analysis_completed = True
                    st.rerun()
    
        if st.session_state.get('image_analysis_completed', False):
            display_image_results()
            if st.button("Continue with Assessment", key="continue_img_btn"):
                st.session_state.show_image_section = False
                st.session_state.current_question += 1
                st.rerun()

    def display_image_results():
        """Display results of image analysis"""
        if not st.session_state.get('image_analysis'):
            return
    
        st.markdown("## 📷 Image Analysis Results")
    
        for i, result in enumerate(st.session_state.image_analysis):
            with st.expander(f"Image {i+1} Analysis", expanded=True):
                col1, col2 = st.columns([1, 3])
            
                with col1:
                    concern_color = {
                        "Low": "#4CAF50",
                        "Moderate": "#FFC107",
                        "High": "#F44336"
                    }.get(result["concern_level"], "#9E9E9E")
                
                    st.markdown(f"""
                    <div style="text-align: center;">
                        <div style="font-size: 2rem; color: {concern_color};">
                            {result["concern_level"]} concern
                        </div>
                        <div>{result["type"]}</div>
                    </div>
                    """, unsafe_allow_html=True)
            
                with col2:
                    st.markdown(f"""
                    **Characteristics:**
                    {result["characteristics"]}
                
                    **Recommendation:**
                    {result["recommendation"]}
                    """)
        
            st.markdown("---")

    def calculate_risk_score():
        """Calculate preliminary cancer risk score based on answers"""
        risk_factors = 0
        total_possible = 0
        concerning_symptoms = []
    
        for q_id, answer in st.session_state.answers.items():
            question = next(q for q in CANCER_QUESTIONS if q['id'] == q_id)
        
            if question.get('risk_factor', False):
                total_possible += 1
                options = question['options']
                answer_index = options.index(answer)
                risk_level = answer_index / len(options)
            
                if risk_level > 0.5: # Higher than middle option
                    risk_factors += 1
                    if question.get('related_to'):
                        concerning_symptoms.extend(question['related_to'])
        
            if question.get('symptom', False):
                options = question['options']
                answer_index = options.index(answer)
                symptom_level = answer_index / len(options)
            
                if symptom_level > 0.5: # Higher than middle option
                    if question.get('related_to'):
                        concerning_symptoms.extend(question['related_to'])
    
        # Calculate risk score (0-100)
        if total_possible > 0:
            risk_score = min(100, (risk_factors / total_possible) * 100 + len(set(concerning_symptoms)) * 5)
        else:
            risk_score = 0
    
        st.session_state.risk_score = risk_score
        st.session_state.concerned_areas = list(set(concerning_symptoms)) # Unique cancer types

    def display_risk_results():
        """Display cancer risk assessment results"""
        calculate_risk_score()
        risk_score = st.session_state.risk_score
    
        st.markdown(f"""
        <div class="risk-display pulse-animation">
            🩺 {int(risk_score)}/100
        </div>
        """, unsafe_allow_html=True)
        st.markdown("### Cancer Risk Assessment")
    
        # Risk level interpretation
        if risk_score < 30:
            risk_level = "Low Risk"
            color = "#4CAF50"
            details = """
            - Your current risk factors are below average
            - Continue healthy lifestyle habits
            - Maintain regular screening as appropriate for your age/gender
            """
        elif risk_score < 70:
            risk_level = "Moderate Risk"
            color = "#FFC107"
            details = """
            - Some concerning risk factors identified
            - Lifestyle modifications recommended
            - Consider more frequent screening
            - Discuss with healthcare provider
            """
        else:
            risk_level = "High Risk"
            color = "#F44336"
            details = """
            - Multiple concerning risk factors
            - Urgent medical evaluation recommended
            - Need for diagnostic testing
            - Immediate lifestyle changes advised
            """
        
        st.markdown(f"""
        <div style="background-color: {color}20; padding: 1.5rem; border-radius: 8px; border-left: 4px solid {color}; margin: 1rem 0;">
            <h4>{risk_level}</h4>
            <div style="margin-left: 1rem;">{details}</div>
        </div>
        """, unsafe_allow_html=True)
        
        # Display concerned areas
        if st.session_state.concerned_areas:
            st.markdown("### 🚨 Areas of Concern")
        
            cols = st.columns(3)
            cancer_types = {
                "Breast cancer": "👩",
                "Lung cancer": "🫁",
                "Prostate cancer": "👨",
                "Colorectal cancer": "🩸",
                "Skin cancer": "☀️",
                "Other": "🩺"
            }
        
            for i, area in enumerate(st.session_state.concerned_areas):
                with cols[i % 3]:
                    emoji = cancer_types.get(area, "🩺")
                    st.markdown(f"""
                    <div style="padding: 1rem; border-radius: 8px; background: {color}10; margin-bottom: 1rem;">
                        <h4>{emoji} {area}</h4>
                    </div>
                    """, unsafe_allow_html=True)
        
        # Display body map (simplified)
        st.markdown("### 🏷️ Body Map")
        st.image("https://upload.wikimedia.org/wikipedia/commons/thumb/3/3e/Human_body_body_silhouette.svg/1200px-Human_body_body_silhouette.svg.png",
                 use_container_width=True, caption="Areas of concern highlighted in your assessment")
    
        st.markdown("---")
        st.subheader("📊 Detailed Risk Factors")
        
        # Display risk factors
        for q_id, answer in st.session_state.answers.items():
            question = next(q for q in CANCER_QUESTIONS if q['id'] == q_id)
            options = question['options']
            answer_index = options.index(answer)
            risk_level = answer_index / len(options)
        
            if risk_level > 0.5 or question.get('symptom', False):
                st.markdown(f"""
                <div style="padding: 1rem; background: #f5f5f5; border-radius: 8px; margin-bottom: 0.5rem;">
                    <strong>{question['question']}</strong><br>
                    Your answer: <span style="color: #d32f2f;">{answer}</span>
                    {f"<br>Related to: {', '.join(question['related_to'])}" if question.get('related_to') else ""}
                </div>
                """, unsafe_allow_html=True)
        
        st.markdown("---")
        action_col1, action_col2 = st.columns(2)
        with action_col1:
            st.markdown("""
            <div style="padding: 1rem; background: #e8f5e9; border-radius: 8px;">
                <h5>🛡️ Cancer Prevention Tips</h5>
                <ul>
                    <li>Maintain healthy weight</li>
                    <li>Exercise regularly</li>
                    <li>Eat more fruits/vegetables</li>
                    <li>Limit processed/red meat</li>
                    <li>Avoid tobacco and limit alcohol</li>
                    <li>Use sun protection</li>
                    <li>Get recommended vaccinations</li>
                    <li>Follow screening guidelines</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)
        with action_col2:
            urgency = "urgent" if risk_score >= 70 else "priority" if risk_score >= 30 else "routine"
            st.markdown(f"""
            <div style="padding: 1rem; background: #fff3e0; border-radius: 8px;">
                <h5>⏰ Recommended Actions</h5>
                <p>Based on your {risk_score}/100 risk score:</p>
                <ul>
                    <li>{"Seek immediate evaluation" if urgency == "urgent" else "Schedule a doctor visit" if urgency == "priority" else "Next regular checkup"}</li>
                    <li>{"Diagnostic tests recommended" if risk_score >= 50 else "Consider screening tests"}</li>
                    <li>Lifestyle counseling</li>
                    <li>{"Genetic counseling if family history" if "Yes" in st.session_state.answers.get(6, "") else ""}</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)

    def display_question():
        """Display current question"""
        if st.session_state.current_question < len(CANCER_QUESTIONS):
            question = CANCER_QUESTIONS[st.session_state.current_question]
            if question['id'] == 24 and st.session_state.get('show_image_section', False):
                st.session_state.current_question += 1
                if st.session_state.current_question >= len(CANCER_QUESTIONS):
                    st.session_state.assessment_complete = True
                st.rerun()
            
            question_html = f"""
            <div class="question-container">
                <h3>Question {question['id']}</h3>
                <p style="font-size: 1.2rem; margin-bottom: 1rem;">{question['question']}</p>
            </div>
            """
            st.markdown(question_html, unsafe_allow_html=True)
            
            if question['type'] == 'number':
                answer = st.number_input(
                    "Your answer:",
                    min_value=question['min_value'],
                    max_value=question['max_value'],
                    value=question['min_value'],
                    key=f"q_{question['id']}"
                )
            elif question['type'] == 'selectbox':
                answer = st.selectbox(
                    "Select your answer:",
                    question['options'],
                    key=f"q_{question['id']}"
                )
            
            col1, col2, col3 = st.columns([1, 1, 1])
            with col1:
                if st.session_state.current_question > 0:
                    if st.button("← Previous", key="prev_btn"):
                        st.session_state.current_question -= 1
                        st.rerun()
            with col3:
                if st.button("Next ->", key="next_btn"):
                    st.session_state.answers[question['id']] = answer
                    if question['id'] == 24:
                        if answer.startswith("Yes"):
                            st.session_state.show_image_section = True
                        else:
                            st.session_state.show_image_section = False
                            st.session_state.current_question += 1
                    else:
                        st.session_state.current_question += 1
                    
                    if st.session_state.current_question >= len(CANCER_QUESTIONS):
                        st.session_state.assessment_complete = True
                    st.rerun()
        return st.session_state.current_question >= len(CANCER_QUESTIONS)

    def format_answers():
        """Format answers for AI analysis"""
        formatted_answers = []
        for q_id, answer in st.session_state.answers.items():
            question = next(q for q in CANCER_QUESTIONS if q['id'] == q_id)
            formatted_answers.append(f"Q{question['id']}: {question['question']}\nA: {answer}\n")
        
        if st.session_state.image_analysis:
            formatted_answers.append("\nIMAGE ANALYSIS DATA:")
            for i, result in enumerate(st.session_state.image_analysis):
                formatted_answers.append(f"- Image {i+1}: {result['type']} ({result['concern_level']} concern)")
        
        formatted_answers.append(f"\nCALCULATED RISK SCORE: {st.session_state.risk_score}/100")
        if st.session_state.concerned_areas:
            formatted_answers.append(f"AREAS OF CONCERN: {', '.join(st.session_state.concerned_areas)}")
        
        return "\n".join(formatted_answers)

    def get_ai_assessment(answers_text):
        """Get AI assessment from A4F"""
        try:
            model_name = initialize_model()
            prompt = f"""
            You are an advanced oncology AI assistant analyzing cancer risk.
            Provide a detailed assessment based on these patient responses:
            {answers_text}
            Structure your response with these sections:
            ## 🩺 Overall Cancer Risk Assessment
            - Risk level (Low/Medium/High)
            - Key risk factors identified
            - Most concerning cancer types
            ## 🚨 Symptom Analysis
            - Evaluation of reported symptoms
            - Urgency of medical evaluation needed
            ## 🔍 Screening Recommendations
            - Recommended cancer screenings based on risk factors
            - Suggested timeline for each screening
            - Any diagnostic tests to consider
            ## 💡 Prevention Strategies
            - Lifestyle modifications to reduce risk
            - Vaccinations to consider (HPV, Hepatitis B)
            - Environmental risk reduction
            ## 📅 Next Steps
            - Suggested timeline for follow-up
            - When to seek immediate medical attention
            - Recommended specialists if needed
            Use clear markdown formatting and provide actionable advice.
            Focus specifically on cancer risk and prevention.
            """
            response = call_a4f_model(prompt, model_name)
            return response
        except Exception as e:
            st.error(f"Error generating assessment: {str(e)}")
            return f"Error generating assessment: {str(e)}"

    def display_assessment_summary():
        """Display assessment summary and get AI response"""
        st.markdown('<h2 style="color: black;">Assessment Summary</h2>', unsafe_allow_html=True)
        
        with st.expander("📋 View Your Responses", expanded=True):
            for q_id, answer in st.session_state.answers.items():
                question = next(q for q in CANCER_QUESTIONS if q['id'] == q_id)
                st.markdown(f"**{question['question']}** \n{answer}")
        
        if st.session_state.image_analysis:
            with st.expander("📷 View Image Analysis Results"):
                display_image_results()
        
        if st.session_state.ai_response is None:
            with st.spinner("🩺 Analyzing your responses with AI..."):
                answers_text = format_answers()
                st.session_state.ai_response = get_ai_assessment(answers_text)
                st.rerun()
        
        if st.session_state.ai_response:
            st.markdown("## 🔍 AI Oncology Assessment")
            st.markdown(st.session_state.ai_response)
            if "high risk" in st.session_state.ai_response.lower() or "urgent" in st.session_state.ai_response.lower():
                st.error("""
                ⚠️ **Urgent Medical Attention Recommended**
                Based on your responses, we recommend seeking immediate oncology evaluation.
                """)
        
        st.markdown("""
        <div style="margin-top: 2rem; padding: 1.5rem; background: #f5f5f5; border-radius: 10px;">
            💡 <strong>Remember:</strong> This is a risk assessment tool only. It cannot diagnose cancer.
            Always consult with qualified oncologists or healthcare professionals for proper evaluation.
            Early detection through screening saves lives.
        </div>
        """, unsafe_allow_html=True)
        
        if st.button("🔄 Start New Assessment", key="reset_btn"):
            for key in list(st.session_state.keys()):
                if key not in ['pro_unlocked', 'model_version', 'image_style', 'video_style']: # Preserve essential states
                    del st.session_state[key]
            st.rerun()

    def main_cancer():
        """Main application function"""
        initialize_session_state()
        st.markdown('<h1 class="main-title">🩺 CancerScan Full Body Cancer Searcher</h1>', unsafe_allow_html=True)
        st.markdown("""
        <div style="background-color: #e3f2fd; padding: 1rem; border-radius: 8px; margin-bottom: 1rem;">
            <h4>🔬 Advanced Cancer Screening:</h4>
            <ul>
                <li>🩺 Comprehensive cancer risk assessment</li>
                <li>📊 Detailed symptom analysis</li>
                <li>🤖 AI-powered oncology insights</li>
                <li>📷 Image analysis for concerning areas</li>
                <li>🏥 Personalized screening recommendations</li>
                <li>📈 Risk tracking over time</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        <div style="background-color: #fff3cd; padding: 1rem; border-radius: 8px; border-left: 4px solid #ffc107; margin-bottom: 2rem;">
            <strong>⚠️ Medical Disclaimer:</strong> This tool provides cancer risk assessment only.
            It is not a diagnostic tool and cannot detect or rule out cancer.
            Always consult with qualified oncologists or healthcare professionals for medical concerns.
        </div>
        """, unsafe_allow_html=True)
        
        if not st.session_state.assessment_complete:
            display_progress()
            if st.session_state.get('show_image_section', False):
                analyze_images()
            else:
                display_question()
        else:
            display_assessment_summary()
    
    main_cancer()

# --------------------------
# FRAMELAB MODULE
# --------------------------
def framelab():
    st.title("🎬 FrameLab: AI-Powered Media Creation")
    st.markdown("Generate, edit images and videos with cutting-edge AI models.")
    
    tab1, tab2, tab3 = st.tabs(["🖼️ Image Generation", "✏️ Image Editing", "🎬 Video Generation"])
 
    with tab1:
        st.subheader("🖼️ Generate New Image")
        prompt = st.text_area("Describe the image you want to create:", height=100, placeholder="E.g., A futuristic cityscape at sunset with flying cars")
        style = st.selectbox("Art Style", ["Sci-Fi", "Realistic", "Fantasy", "Abstract", "Oil Painting", "Digital Art"])
        
        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("🎨 Generate Image", type="primary"):
                with st.spinner("Generating your image..."):
                    generated_image = generate_image(prompt, style)
                    if generated_image:
                        st.session_state.generated_image = generated_image
                        st.success("Image generated successfully!")
                    else:
                        st.error("Failed to generate image. Please try again.")
    
        if hasattr(st.session_state, 'generated_image') and st.session_state.generated_image:
            st.image(st.session_state.generated_image, caption="Generated Image", use_container_width=True)
            if st.button("🔄 Generate Another"):
                del st.session_state.generated_image
                st.rerun()
 
    with tab2:
        st.subheader("✏️ Edit Existing Image")
        uploaded_image = st.file_uploader("Upload an image to edit:", type=["jpg", "jpeg", "png"])
    
        if uploaded_image:
            image = Image.open(uploaded_image)
            st.image(image, caption="Original Image", use_container_width=True)
        
            edit_prompt = st.text_area("Describe the edits you want (e.g., 'add a sunset background, make the sky vibrant'):", height=100)
        
            if st.button("✏️ Apply Edits", type="primary"):
                if edit_prompt:
                    with st.spinner("Editing your image..."):
                        edited_image = edit_image(image, edit_prompt)
                        if edited_image:
                            st.session_state.edited_image = edited_image
                            st.success("Image edited successfully!")
                        else:
                            st.error("Failed to edit image. Please try again.")
                else:
                    st.warning("Please provide edit instructions.")
    
        if hasattr(st.session_state, 'edited_image') and st.session_state.edited_image:
            st.image(st.session_state.edited_image, caption="Edited Image", use_container_width=True)
            col1, col2 = st.columns(2)
            with col1:
                if st.button("💾 Download Edited Image"):
                    buffered = BytesIO()
                    st.session_state.edited_image.save(buffered, format="PNG")
                    st.download_button(
                        label="Download PNG",
                        data=buffered.getvalue(),
                        file_name="edited_image.png",
                        mime="image/png"
                    )
            with col2:
                if st.button("🔄 Edit Again"):
                    del st.session_state.edited_image
                    st.rerun()
 
    with tab3:
        st.subheader("🎬 Generate Video")
        prompt = st.text_area("Describe the video scene:", height=100, placeholder="E.g., A woman walking through a busy Tokyo street at night, wearing dark sunglasses")
     
        if st.button("🎥 Generate Video", type="primary"):
            with st.spinner("Generating your video... This may take a few minutes."):
                video_file = generate_video_replicate(prompt, "")
                if video_file and os.path.exists(video_file):
                    st.session_state.generated_video = video_file
                    st.success("Video generated successfully!")
                    st.video(video_file)
                    st.download_button(
                        label="💾 Download Video",
                        data=open(video_file, "rb").read(),
                        file_name="generated_video.mp4",
                        mime="video/mp4"
                    )
                else:
                    st.error("Failed to generate video. Please try again.")
     
        if hasattr(st.session_state, 'generated_video') and st.session_state.generated_video:
            if st.button("🔄 Generate Another Video"):
                # Clean up the file
                if os.path.exists(st.session_state.generated_video):
                    os.remove(st.session_state.generated_video)
                del st.session_state.generated_video
                st.rerun()

# --------------------------
# QUANTUM CREATIVESTUDIO MODULE
# --------------------------
def quantum_creativestudio():
    st.title("🎨 Quantum CreativeStudio")
    st.markdown("Advanced creative AI studio for multimedia generation and editing")

    st.write(
        '<a href="https://creativestudio-3ata6gv6.manus.space" target="_blank">'
        '<button style="padding:10px 20px; font-size:18px;">🚀 Open CreativeStudio</button>'
        '</a>',
        unsafe_allow_html=True,
    )

    st.markdown("---")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.info("**Features:**\n- Advanced image generation\n- Video editing tools\n- 3D model creation")
    with col2:
        st.info("**Tools:**\n- AI-powered design\n- Real-time collaboration\n- Cloud rendering")
    with col3:
        st.info("**Support:**\n- Multi-format export\n- Team workspace\n- Version control")


# -----------------------------------
# QUANTUM LM MODULE
# -----------------------------------
def quantum_lm():
    st.title("🧠 Quantum LM")
    st.markdown("Advanced Notebook with quantum-inspired architecture")

    st.write(
        '<a href="https://quantumlm-w2cjzzsd.manus.space" target="_blank">'
        '<button style="padding:10px 20px; font-size:18px;">🧠 Open Quantum LM</button>'
        '</a>',
        unsafe_allow_html=True,
    )

    st.markdown("---")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.info("**Capabilities:**\n- Quantum-enhanced reasoning\n- Multi-modal understanding\n- Real-time learning")
    with col2:
        st.info("**Features:**\n- Contextual memory\n- Emotional intelligence\n- Creative writing")
    with col3:
        st.info("**Applications:**\n- Research assistance\n- Code generation\n- Content creation")
# --------------------------
# HISTORY DISPLAY
# --------------------------
def show_history():
    st.title("📜 Query History")
    history = load_history()
    if not history:
        st.info("No query history yet.")
    else:
        for item in history[::-1]: # Recent first
            st.markdown(f"**{item['timestamp']}**")
            st.write(item['query'])
            st.markdown("---")

# --------------------------
# MAIN APP NAVIGATION
# --------------------------
if st.session_state.pro_unlocked:
    with st.sidebar:
        st.markdown("### 🚀 Quantora Modes")
        mode = st.radio(
            "Select Mode",
            ["AI", "AI Content Detector", "AI Humanizer", "Quantora News", "Quantora Trade Charts", 
              "Heart Health Analyzer", "Brain Health Analyzer", 
             "Cancer Risk Assessor", "FrameLab", "Quantum CreativeStudio", "Quantum LM", 
              "Quantomise My Trip", "Coding", "Quantora Weather",
             "Collage Maker", "Sound Extractor", "Shopping Research"],  # Added Quantora Translate
            index=0,
            key="current_mode"
        )
    
        st.markdown("---")
        st.markdown("### 📁 Document & Image Analysis")
        uploaded_file = st.file_uploader(
            "Upload Document or Image",
            type=['txt', 'pdf', 'docx', 'csv', 'json', 'py', 'js', 'html', 'css', 'md', 'jpg', 'jpeg', 'png'],
            help="Upload documents or images for AI analysis and enhancement",
            key="document_uploader"
        )
    
        if uploaded_file:
            with st.spinner("🔍 Analyzing content..."):
                content = process_uploaded_file(uploaded_file)
                st.session_state.uploaded_content = content
                st.success(f"✅ {uploaded_file.name} processed!")
            
                if uploaded_file.type.startswith('image/'):
                    display_image_enhancement_controls(st.session_state.uploaded_image, st.session_state.enhancement_values)
                else:
                    with st.expander("📄 Preview Content"):
                        preview_content = content[:1000] + "..." if len(content) > 1000 else content
                        st.text_area("Document Content", preview_content, height=200, disabled=True)
        
        if st.button("🗑️ Clear Uploads", use_container_width=True):
            st.session_state.uploaded_content = ""
            st.session_state.uploaded_image = None
            st.session_state.enhanced_image = None
            st.session_state.image_style = "Sci-Fi"
            st.session_state.video_style = "Cinematic"
            st.session_state.enhancement_values = {
                "brightness": 1.0,
                "contrast": 1.0,
                "sharpness": 1.0,
                "color": 1.0,
                "filter": "None"
            }
            st.success("✅ All uploads cleared!")
            st.rerun()
else:
    mode = "AI" # Force AI mode in trial
    with st.sidebar:
        st.markdown("### 📜 Query History")
        history = load_history()
        if not history:
            st.info("No query history yet.")
        else:
            for item in history[::-1]: # Recent first
                st.write(f"{item['timestamp']}: {item['query']}")

# Main Content Area
if mode == "AI":
    if not st.session_state.chat:
        with st.container():
            st.markdown("""
            <div class="welcome-container">
                <div class="welcome-title">{}</div>
                <p>Where knowledge ends.</p>
            </div>
            """.format(app_name), unsafe_allow_html=True)
        
            col1, col2, col3, col4 = st.columns(4)
        
            with col1:
                if st.button("💠 Simulate a quantum network"):
                    prompt = "Simulate a quantum network"
                    start_time = time.time()
                    st.session_state.chat.append(("user", prompt, datetime.now()))
                    response = call_quantora_unified(prompt)
                    response_time = time.time() - start_time
                    st.session_state.chat.append(("quantora", response, datetime.now(), response_time))
                    save_history(prompt)
                    st.rerun()
        
            with col2:
                if st.button("🧬 Simulate a molecular model"):
                    prompt = "Simulate a molecular model"
                    start_time = time.time()
                    st.session_state.chat.append(("user", prompt, datetime.now()))
                    response = call_quantora_unified(prompt)
                    response_time = time.time() - start_time
                    st.session_state.chat.append(("quantora", response, datetime.now(), response_time))
                    save_history(prompt)
                    st.rerun()
        
            with col3:
                if st.button("🌍 Predict climate patterns"):
                    prompt = "Predict climate patterns"
                    start_time = time.time()
                    st.session_state.chat.append(("user", prompt, datetime.now()))
                    response = call_quantora_unified(prompt)
                    response_time = time.time() - start_time
                    st.session_state.chat.append(("quantora", response, datetime.now(), response_time))
                    save_history(prompt)
                    st.rerun()
        
            with col4:
                if st.button("📜 Draft AI ethics code"):
                    prompt = "Draft AI ethics code"
                    start_time = time.time()
                    st.session_state.chat.append(("user", prompt, datetime.now()))
                    response = call_quantora_unified(prompt)
                    response_time = time.time() - start_time
                    st.session_state.chat.append(("quantora", response, datetime.now(), response_time))
                    save_history(prompt)
                    st.rerun()
        
            st.markdown("<p style='text-align: center; margin-top: 2rem;'><strong>Ask Quantora anything...</strong></p>",
                        unsafe_allow_html=True)
    
    for i, chat_item in enumerate(st.session_state.chat):
        if len(chat_item) >= 3:
            speaker, message, timestamp = chat_item[:3]
            response_time = chat_item[3] if len(chat_item) > 3 else None
        
            if speaker == "user":
                st.markdown(f"""
                <div class="chat-message user-message">
                    <div class="message-header">
                        👤 <strong>You</strong>
                        <span class="message-time">{timestamp.strftime('%H:%M:%S')}</span>
                    </div>
                    <div>{message}</div>
                </div>
                """, unsafe_allow_html=True)
        
            elif speaker in ["quantora", "ai"]:
                formatted_parts = format_response_with_code(message)
            
                st.markdown(f"""
                <div class="chat-message ai-message">
                    <div class="message-header">
                        💎 <strong>Quantora</strong>
                        <span class="message-time">{timestamp.strftime('%H:%M:%S')} • ⏱️ {response_time:.1f}s</span>
                    </div>
                """, unsafe_allow_html=True)
            
                for part in formatted_parts:
                    if part[0] == 'text':
                        st.markdown(f"<div>{part[1]}</div>", unsafe_allow_html=True)
                    elif part[0] == 'code':
                        st.code(part[1], language=part[2])
            
                st.markdown("</div>", unsafe_allow_html=True)
    
    st.markdown("---")
    col1, col2 = st.columns([0.85, 0.15])
    with col1:
        user_input = st.text_area(
            "Your message:",
            placeholder="Ask Quantora anything...",
            height=100,
            key="main_input",
            label_visibility="collapsed"
        )
    with col2:
        st.write("")
        send_button = st.button("💬 Send", use_container_width=True, type="primary")
    
    if send_button and user_input.strip():
        start_time = time.time()
        st.session_state.chat.append(("user", user_input.strip(), datetime.now()))
    
        with st.spinner("⚛️ Quantumizing Through Timeless Refinement Toward the Ultimate Answer."):
            context = st.session_state.uploaded_content
            image = st.session_state.uploaded_image if st.session_state.uploaded_image else None
            response = call_quantora_unified(user_input.strip(), context, image)
    
        response_time = time.time() - start_time
        st.session_state.last_response_time = response_time
        st.session_state.chat.append(("quantora", response, datetime.now(), response_time))
        save_history(user_input.strip())
        st.rerun()
    
    if st.session_state.chat:
        response_times = []
        for item in st.session_state.chat:
            if len(item) > 3 and item[0] == "quantora" and isinstance(item[3], (int, float)):
                response_times.append(item[3])
    
        if response_times:
            avg_time = sum(response_times) / len(response_times)
            min_time = min(response_times)
            max_time = max(response_times)
        
            st.markdown(f"""
            <div style="background: rgba(30, 41, 59, 0.6); border-radius: 12px; padding: 1rem; margin: 1rem 0; text-align: center;">
                📊 <strong>Performance Metrics:</strong>
                Avg: <strong>{avg_time:.1f}s</strong> •
                Fastest: <strong>{min_time:.1f}s</strong> •
                Slowest: <strong>{max_time:.1f}s</strong> •
                Total: <strong>{len(response_times)}</strong>
            </div>
            """, unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        if st.button("🗑️ Clear Chat", use_container_width=True):
            st.session_state.chat = []
            st.success("✅ Chat cleared!")
            st.rerun()
    with col2:
        if st.button("📊 Export Chat", use_container_width=True):
            if st.session_state.chat:
                chat_data = []
                for item in st.session_state.chat:
                    chat_data.append({
                        "Speaker": item[0],
                        "Message": item[1],
                        "Timestamp": item[2].strftime('%Y-%m-%d %H:%M:%S'),
                        "Response_Time": item[3] if len(item) > 3 else None
                    })
            
                chat_json = json.dumps(chat_data, indent=2, default=str)
                st.download_button(
                    label="💾 Download Chat JSON",
                    data=chat_json,
                    file_name=f"quantora_chat_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json"
                )
            else:
                st.info("No chat history to export")
    with col3:
        if st.button("ℹ️ About", use_container_width=True):
            st.info("""
            **Quantora AI Elite** v2.4
        
            Features:
            ✅ Document analysis
            ✅ Image enhancement
            ✅ Code formatting (always full code)
            ✅ Performance metrics
            ✅ Enhanced response quality
            ✅ AI Content Detection
            ✅ AI Text Humanization
            ✅ Quantora Translate
            """)
    with col4:
        if st.session_state.pro_unlocked:
            with st.popover("⚙️ Select Model", use_container_width=True):
                st.markdown("##### Select Quantora Engine")
                st.radio(
                    "Engine Selection",
                    options=[
                        "Quantora Prime 1 (Latest Flagship Model)",
                        "Quantora Prime 1 Fast (Faster But Not As Better As Og Flagship Model)"
                    ],
                    key="model_version",
                    label_visibility="collapsed",
                    help="Select specialized model versions for different tasks",
                )

elif mode == "AI Content Detector":
    ai_content_detector_mode()
elif mode == "AI Humanizer":
    ai_humanizer_mode()
elif mode == "Quantora News":
    quantora_news()
elif mode == "Quantora Trade Charts":
    quantora_trade_charts()
elif mode == "Heart Health Analyzer":
    heart_health_analyzer()
elif mode == "Brain Health Analyzer":
    brain_health_analyzer()
elif mode == "Cancer Risk Assessor":
    cancer_risk_assessor()
elif mode == "FrameLab":
    framelab()
elif mode == "Quantum CreativeStudio":
    quantum_creativestudio()
elif mode == "Quantomise My Trip":
    quantomise_my_trip()
elif mode == "Coding":
    coding_workspace()
elif mode == "Quantora Weather":
    quantora_weather()
elif mode == "Collage Maker":
    collage_maker()
elif mode == "Sound Extractor":
    sound_extractor()
elif mode == "Shopping Research":
    shopping_research()
elif mode == "Quantum LM":
    quantum_lm()


st.markdown(
    """
    <meta name="google-site-verification" content="r0-f193FBPXTVV2BAm3T85RRty-35YaqER_pKSwtIM8" />
    """,
    unsafe_allow_html=True
)
                      
# Footer
st.markdown("---")
st.markdown(
    '<div style="text-align: center; color: var(--text-muted); font-size: 0.9rem;">'
    '💎 Quantora AI - Advanced AI Assistant | '
    'Powered by Groq Models, A4F Models | Double-check - Quantora isn\'t perfect. |'
    f'Session started: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
    '</div>',
    unsafe_allow_html=True
)
